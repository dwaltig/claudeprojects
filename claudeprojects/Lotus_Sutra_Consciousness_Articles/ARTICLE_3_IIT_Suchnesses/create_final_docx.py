#!/usr/bin/env python3
"""
Generate final Word document for "Lotus Sutra, IIT, and Universal Consciousness"
Converts HUMANIZED_DRAFT.md to submission-ready DOCX
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
import os

# Get the directory where this script is located
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "JCS_SUBMISSION_BLIND.md")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "JCS_SUBMISSION_BLIND.docx")


def setup_document():
    """Create and configure the Word document."""
    doc = Document()
    
    # Set up page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Configure Normal style (body text)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    
    return doc


def process_inline_formatting(paragraph, text):
    """Handle inline formatting like *italics*, **bold**, and $math$."""
    import re
    
    # Pattern to find formatting markers - order matters (bold before italic)
    # Using word boundary patterns to avoid matching partial asterisks
    patterns = [
        (r'\*\*([^*]+)\*\*', 'bold'),      # **bold** - non-greedy, no asterisks inside
        (r'(?<!\*)\*([^*]+)\*(?!\*)', 'italic'),  # *italic* - not adjacent to other asterisks
        (r'\$([^$]+)\$', 'italic'),         # $math$ (render as italic)
    ]
    
    # Find all matches for all patterns
    all_matches = []
    for pattern, style_type in patterns:
        for match in re.finditer(pattern, text):
            all_matches.append((match.start(), match.end(), match.group(1), style_type))
    
    # Sort by position
    all_matches.sort(key=lambda x: x[0])
    
    # Remove overlapping matches (keep first)
    filtered_matches = []
    last_end = 0
    for start, end, content, style_type in all_matches:
        if start >= last_end:
            filtered_matches.append((start, end, content, style_type))
            last_end = end
    
    # If no matches, just add the text as-is
    if not filtered_matches:
        paragraph.add_run(text)
        return
    
    # Build runs
    current_pos = 0
    for start, end, content, style_type in filtered_matches:
        if start > current_pos:
            paragraph.add_run(text[current_pos:start])
        run = paragraph.add_run(content)
        if style_type == 'bold':
            run.bold = True
        elif style_type == 'italic':
            run.italic = True
        current_pos = end
    
    # Add remaining text
    if current_pos < len(text):
        paragraph.add_run(text[current_pos:])


def convert_markdown_to_docx(md_content, doc):
    """Convert markdown content to Word document."""
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i]
        
        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            i += 1
            continue
            
        if line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
            i += 1
            continue
        
        # Handle blockquotes
        if line.startswith('>'):
            quote_text = line[1:].strip()
            while quote_text.startswith('>'):
                quote_text = quote_text[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            process_inline_formatting(p, quote_text)
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.strip()[2:]
            process_inline_formatting(p, text)
            i += 1
            continue
        
        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if match:
            list_num = int(match.group(1))
            p = doc.add_paragraph()
            # Format as numbered list manually to control restart
            p.style = 'List Number'
            # If this is item 1, we need to restart numbering
            if list_num == 1:
                # Use XML to restart numbering
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                pPr = p._p.get_or_add_pPr()
                numPr = pPr.get_or_add_numPr()
                # Set numId to create new list instance
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), '0')  # Will be reset by Word
                numPr.append(numId)
            process_inline_formatting(p, match.group(2))
            i += 1
            continue
        
        # Handle tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            
            # Parse row
            cells = [c.strip() for c in line.strip().split('|')]
            cells = [c for c in cells if c]  # Remove empty strings
            
            # Skip separator rows
            if cells and all(set(c) <= set('-|: ') for c in cells):
                i += 1
                continue
                
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table, create it
            if table_rows:
                num_cols = max(len(row) for row in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # Clear default text and use inline formatting
                            cell.text = ""
                            process_inline_formatting(cell.paragraphs[0], cell_text)
                            # Bold header row
                            if row_idx == 0:
                                for run in cell.paragraphs[0].runs:
                                    run.bold = True
                
                doc.add_paragraph()  # Space after table
            in_table = False
            table_rows = []
            continue
        
        # Handle empty lines
        if not line.strip():
            i += 1
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        process_inline_formatting(p, line.strip())
        i += 1
    
    # Handle any remaining table
    if in_table and table_rows:
        num_cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    table.rows[row_idx].cells[col_idx].text = cell_text


def main():
    print(f"Reading: {INPUT_FILE}")
    
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    print("Creating Word document...")
    doc = setup_document()
    
    convert_markdown_to_docx(md_content, doc)
    
    print(f"Saving: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    
    print("\n✅ Document created successfully!")
    print(f"   Output: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
