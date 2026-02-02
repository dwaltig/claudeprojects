#!/usr/bin/env python3
"""
Create professionally formatted Word document for Consc-4 (Bodhisattva Field).
Converts Draft_Consc_4_Bodhisattva_Field.md to .docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import os

def create_formatted_document(input_file, output_file):
    """Create a professionally formatted Word document from text file."""

    print(f"Reading from: {input_file}")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Set style to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        
        if not line:
            continue
            
        # Check for headers
        is_header = False
        is_title = False
        
        if line.startswith('# '):
            is_title = True
            clean_line = line[2:].strip()
        elif line.startswith('## '):
            is_header = True
            clean_line = line[3:].strip()
        elif line.startswith('### '):
            is_header = True
            clean_line = line[4:].strip()
        else:
            clean_line = line

        p = doc.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        if is_title:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(clean_line)
            run.bold = True
            run.font.size = Pt(14) # Slightly larger for main title
            continue # Skip normal processing for title
            
        elif is_header:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Run header logic below
        else:
            # Body paragraph
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.first_line_indent = Inches(0.5)

        process_line_formatting(p, clean_line, is_header)

    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def process_line_formatting(paragraph, text, is_header=False):
    # Handle list bullets first - strip them and add regular text
    if text.startswith('*   '):
        # This is a list item, not italics
        run = paragraph.add_run('\u2022 ')  # Unicode bullet
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        text = text[4:]  # Remove the "* " prefix

    # Now handle bold (**text**) and italics (*text*) 
    # Use a pattern that requires non-space after opening and before closing
    # to avoid matching list bullets
    
    # First pass: handle **bold**
    bold_pattern = r'\*\*([^*]+)\*\*'
    # Second pass: handle *italics* (single asterisks, but not list bullets)
    italic_pattern = r'(?<!\*)\*([^*\s][^*]*[^*\s]|[^*\s])\*(?!\*)'
    
    # Combined approach: split and process
    parts = []
    last_end = 0
    
    # Find all bold matches first
    for match in re.finditer(bold_pattern, text):
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    # Now process each part for italics
    final_parts = []
    for style, content in parts:
        if style == 'bold':
            final_parts.append(('bold', content))
        else:
            # Look for italics within normal text
            italic_last = 0
            for match in re.finditer(italic_pattern, content):
                if match.start() > italic_last:
                    final_parts.append(('normal', content[italic_last:match.start()]))
                final_parts.append(('italic', match.group(1)))
                italic_last = match.end()
            if italic_last < len(content):
                final_parts.append(('normal', content[italic_last:]))
    
    # Render all parts
    for style, content in final_parts:
        if not content:
            continue
        run = paragraph.add_run(content)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if style == 'bold' or is_header:
            run.bold = True
        if style == 'italic':
            run.italic = True

if __name__ == '__main__':
    base_dir = '/Users/williamaltig/claudeprojects/Lotus_Sutra_Consciousness_Articles'
    input_file = os.path.join(base_dir, 'Draft_Consc_4_Bodhisattva_Field.md')
    output_file = os.path.join(base_dir, 'Draft_Consc_4_Bodhisattva_Field.docx')
    
    create_formatted_document(input_file, output_file)
