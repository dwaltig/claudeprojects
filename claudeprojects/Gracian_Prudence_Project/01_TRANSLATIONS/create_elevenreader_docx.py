#!/usr/bin/env python3
"""
Create DOCX version of Gracián ElevenReader edition
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_docx():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(12)
    
    # Read source
    with open('Oraculo_Manual_ELEVENREADER.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            # Chapter/Title heading
            heading_text = line[2:]
            p = doc.add_heading(heading_text, level=1)
        elif line.startswith('## '):
            # Aphorism heading
            heading_text = line[3:]
            p = doc.add_heading(heading_text, level=2)
        elif line.startswith('*') and line.endswith('*'):
            # Italic subtitle
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line == '---':
            # Separator - skip
            continue
        else:
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Save
    doc.save('Oraculo_Manual_ELEVENREADER.docx')
    print("Created: Oraculo_Manual_ELEVENREADER.docx")

if __name__ == '__main__':
    create_docx()
