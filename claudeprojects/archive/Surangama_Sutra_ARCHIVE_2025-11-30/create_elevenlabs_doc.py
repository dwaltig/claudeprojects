#!/usr/bin/env python3
"""
Create Word document for ElevenLabs with NATURAL prose flow
Smart AI voices read punctuation - they don't need artificial line breaks
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the clean audio script
with open('SURANGAMA_SUTRA_AUDIO_SCRIPT.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set up styles
styles = doc.styles
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Arial'
normal_font.size = Pt(12)
normal_paragraph = normal_style.paragraph_format
normal_paragraph.space_after = Pt(6)
normal_paragraph.line_spacing = 1.15

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Process content
lines = content.split('\n')
prose_buffer = []

def is_prose_line(stripped):
    """Check if this is regular prose that should be combined"""
    if not stripped:
        return False
    # Not prose if it's a header, speaker, stage direction, or special line
    if stripped.isupper():
        return False
    if stripped.startswith('(') and stripped.endswith(')'):
        return False
    if ':' in stripped and stripped.split(':')[0].isupper():
        return False
    if stripped.startswith('Voice:') or stripped.startswith('Voices:'):
        return False
    if stripped.startswith('---'):
        return False
    if 'SURANGAMA SUTRA' in stripped:
        return False
    if stripped.startswith('PART ') or stripped.startswith('Location '):
        return False
    if stripped in ['CONTENTS', 'FRONT MATTER', 'BACK MATTER']:
        return False
    # Otherwise it's prose
    return True

def flush_prose():
    """Combine accumulated prose into one flowing paragraph"""
    global prose_buffer
    if prose_buffer:
        combined = ' '.join(prose_buffer)
        doc.add_paragraph(combined)
        prose_buffer = []

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip empty lines but flush prose buffer
    if not stripped:
        flush_prose()
        i += 1
        continue

    # Title page - centered
    if 'THE SURANGAMA SUTRA' in stripped:
        flush_prose()
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        i += 1
        continue

    if 'Volumes 1-4' in stripped or 'Multi-Vocal' in stripped or 'Classical Chinese' in stripped:
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(12)
        if 'Multi-Vocal' in stripped or 'Classical' in stripped:
            run.font.italic = True
        i += 1
        continue

    if 'Translated by' in stripped or 'Houston, Texas' in stripped:
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        i += 1
        continue

    # Major sections
    if stripped in ['CONTENTS', 'FRONT MATTER', 'VOLUME ONE', 'VOLUME TWO',
                    'VOLUME THREE', 'VOLUME FOUR', 'BACK MATTER']:
        flush_prose()
        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_before = Pt(18)
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        i += 1
        continue

    # All caps headings
    if stripped.isupper() and len(stripped) > 5:
        flush_prose()
        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_before = Pt(12)
        run = p.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        i += 1
        continue

    # Speaker labels
    if ':' in stripped and stripped.split(':')[0].isupper():
        flush_prose()
        parts = stripped.split(':', 1)
        speaker = parts[0]
        text = parts[1].strip() if len(parts) > 1 else ''

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)

        # Bold speaker
        run = p.add_run(speaker + ':')
        run.bold = True

        # Text after speaker
        if text:
            p.add_run(' ' + text)
        i += 1
        continue

    # Stage directions
    if stripped.startswith('(') and stripped.endswith(')'):
        flush_prose()
        p = doc.add_paragraph(stripped)
        run = p.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        i += 1
        continue

    # Voice annotations
    if stripped.startswith('Voice:') or stripped.startswith('Voices:'):
        flush_prose()
        p = doc.add_paragraph(stripped)
        run = p.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(0, 100, 200)
        i += 1
        continue

    # Section breaks
    if stripped.startswith('---'):
        flush_prose()
        doc.add_paragraph()
        i += 1
        continue

    # Part/Location headings
    if (stripped.startswith('PART ') or stripped.startswith('Location ') or
        stripped.startswith('The ') and len(stripped.split()) <= 5):
        flush_prose()
        p = doc.add_paragraph(stripped)
        run = p.runs[0]
        run.bold = True
        i += 1
        continue

    # Regular prose - accumulate for flowing paragraphs
    if is_prose_line(stripped):
        prose_buffer.append(stripped)
        i += 1
        continue

    # Anything else - add as is
    flush_prose()
    doc.add_paragraph(stripped)
    i += 1

# Flush any remaining prose
flush_prose()

# Save
output_file = 'SURANGAMA_SUTRA_ELEVENLABS_NATURAL.docx'
doc.save(output_file)

print(f"✅ ElevenLabs document created: {output_file}")
print(f"📢 Natural prose flow for smart AI voices:")
print(f"   - Prose combines into flowing paragraphs")
print(f"   - AI voices read punctuation naturally")
print(f"   - Speaker changes = paragraph breaks")
print(f"   - Stage directions = separate paragraphs")
print(f"   - NO artificial pauses from line breaks")
