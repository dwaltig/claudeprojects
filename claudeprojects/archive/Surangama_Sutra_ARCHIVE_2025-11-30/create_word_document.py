#!/usr/bin/env python3
"""
Create professionally formatted Word document from Surangama Sutra text
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def clean_markdown(text):
    """Remove all markdown formatting from text"""
    # Remove header markers
    text = re.sub(r'#{1,6}\s*', '', text)
    # Remove bold markers
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    return text.strip()

def parse_markdown_to_runs(paragraph, text):
    """Parse markdown formatting and add formatted runs to paragraph"""
    # First remove any ## symbols
    text = re.sub(r'#{1,6}\s*', '', text)

    # Pattern to find **bold**, *italic*, or plain text
    pattern = r'(\*\*.*?\*\*|\*.*?\*|[^*]+|\*)'

    parts = re.findall(pattern, text)

    for part in parts:
        if not part or part == '*':
            continue

        if part.startswith('**') and part.endswith('**'):
            # Bold text
            clean_text = part.strip('*')
            if clean_text:
                run = paragraph.add_run(clean_text)
                run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            clean_text = part.strip('*')
            if clean_text:
                run = paragraph.add_run(clean_text)
                run.italic = True
        else:
            # Plain text
            if part:
                paragraph.add_run(part)

def is_special_line(stripped):
    """Check if line requires special formatting (not regular prose)"""
    if not stripped:
        return True
    if stripped.startswith('===='):
        return True
    if stripped.startswith('##'):
        return True
    if stripped.startswith('- ') or stripped.startswith('* '):
        return True
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
        return True
    if stripped.startswith('**') and '**' in stripped[2:]:
        return True
    if stripped.startswith('(') and stripped.endswith(')'):
        return True
    if stripped.isupper() and len(stripped) > 3:
        return True
    if stripped in ['FRONT MATTER', 'VOLUME ONE', 'VOLUME TWO', 'VOLUME THREE', 'VOLUME FOUR', 'BACK MATTER']:
        return True
    return False

# Read the source text
with open('SURANGAMA_SUTRA_VOLUMES_1-4_COMPLETE.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set up styles
styles = doc.styles

# Configure Normal style
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Garamond'
normal_font.size = Pt(12)
normal_paragraph = normal_style.paragraph_format
normal_paragraph.space_after = Pt(6)
normal_paragraph.line_spacing = 1.15

# Configure Heading 1
heading1_style = styles['Heading 1']
heading1_font = heading1_style.font
heading1_font.name = 'Garamond'
heading1_font.size = Pt(18)
heading1_font.bold = True
heading1_font.color.rgb = RGBColor(0, 0, 0)

# Configure Heading 2
heading2_style = styles['Heading 2']
heading2_font = heading2_style.font
heading2_font.name = 'Garamond'
heading2_font.size = Pt(14)
heading2_font.bold = True

# Configure Heading 3
heading3_style = styles['Heading 3']
heading3_font = heading3_style.font
heading3_font.name = 'Garamond'
heading3_font.size = Pt(12)
heading3_font.bold = True
heading3_font.italic = True

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Process content line by line
lines = content.split('\n')

in_title_page = False
current_volume = None
prose_buffer = []  # Buffer for accumulating prose lines

def flush_prose_buffer():
    """Add accumulated prose lines as a single paragraph"""
    global prose_buffer
    if prose_buffer:
        combined_text = ' '.join(prose_buffer)
        p = doc.add_paragraph()
        parse_markdown_to_runs(p, combined_text)
        prose_buffer = []

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip separator lines
    if stripped.startswith('===='):
        flush_prose_buffer()
        i += 1
        continue

    # Handle title page (centered, larger)
    if 'THE SURANGAMA SUTRA' in stripped and 'Volumes 1-4' in stripped:
        flush_prose_buffer()
        in_title_page = True
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_format = p.paragraph_format
        p_format.space_before = Pt(72)
        p_format.space_after = Pt(12)
        run = p.runs[0]
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.name = 'Garamond'
        i += 1
        continue

    # Handle title page subtitle lines
    if in_title_page and stripped and not stripped.startswith('CONTENTS'):
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Garamond'
        if 'Classical Chinese' in stripped or 'Translated by' in stripped:
            run.font.size = Pt(12)
            run.font.italic = True
        else:
            run.font.size = Pt(14)
        i += 1
        continue

    # End of title page, add page break
    if stripped.startswith('CONTENTS'):
        flush_prose_buffer()
        in_title_page = False
        doc.add_page_break()
        doc.add_heading('CONTENTS', level=1)
        i += 1
        continue

    # Handle major sections as Heading 1
    if stripped in ['FRONT MATTER', 'VOLUME ONE', 'VOLUME TWO', 'VOLUME THREE',
                    'VOLUME FOUR', 'BACK MATTER']:
        flush_prose_buffer()
        if current_volume is not None:
            doc.add_page_break()
        doc.add_heading(stripped, level=1)
        current_volume = stripped
        i += 1
        continue

    # Handle section headers (lines with all caps or starting with ##)
    if stripped and (stripped.isupper() or stripped.startswith('##')):
        header_text = stripped.replace('##', '').strip()
        header_text = clean_markdown(header_text)
        if len(header_text) > 3 and header_text not in ['THE TEACHER', 'ANANDA']:
            flush_prose_buffer()
            doc.add_heading(header_text, level=2)
            i += 1
            continue

    # Handle subsection headers (starting with ###)
    if stripped.startswith('###'):
        flush_prose_buffer()
        header_text = stripped.replace('###', '').strip()
        header_text = clean_markdown(header_text)
        doc.add_heading(header_text, level=3)
        i += 1
        continue

    # Handle bulleted lists
    if stripped.startswith('- ') or stripped.startswith('* '):
        flush_prose_buffer()
        list_text = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        parse_markdown_to_runs(p, list_text)
        i += 1
        continue

    # Handle numbered lists
    if len(stripped) > 2 and stripped[0].isdigit() and stripped[1] in '.):':
        flush_prose_buffer()
        list_text = stripped[2:].strip()
        p = doc.add_paragraph(style='List Number')
        parse_markdown_to_runs(p, list_text)
        i += 1
        continue

    # Handle speaker labels
    if stripped.startswith('**') and '**' in stripped[2:]:
        flush_prose_buffer()
        match = re.match(r'\*\*(.*?)\*\*:?(.*)', stripped)
        if match:
            speaker = match.group(1)
            rest = match.group(2).strip()

            p = doc.add_paragraph()
            run = p.add_run(speaker + ':')
            run.bold = True
            run.font.size = Pt(12)

            if rest:
                p.add_run(' ')
                parse_markdown_to_runs(p, rest)

            p.paragraph_format.space_before = Pt(6)
            i += 1
            continue

    # Handle stage directions
    if stripped.startswith('(') and stripped.endswith(')'):
        flush_prose_buffer()
        p = doc.add_paragraph()
        content_text = stripped[1:-1]
        run = p.add_run('(')
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

        parse_markdown_to_runs(p, content_text)

        run = p.add_run(')')
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)

        p.paragraph_format.left_indent = Inches(0.5)
        i += 1
        continue

    # Handle quoted speech (keep as separate paragraphs with indent)
    if stripped.startswith('"') or stripped.startswith('"'):
        flush_prose_buffer()
        p = doc.add_paragraph()
        parse_markdown_to_runs(p, stripped)
        p.paragraph_format.left_indent = Inches(0.25)
        i += 1
        continue

    # Handle blank lines - flush prose buffer and add paragraph break
    if not stripped:
        flush_prose_buffer()
        doc.add_paragraph()
        i += 1
        continue

    # Regular prose - accumulate in buffer
    prose_buffer.append(stripped)
    i += 1

# Flush any remaining prose
flush_prose_buffer()

# Add page numbers to footer
for section in doc.sections:
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '\t\t'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# Save document
output_file = 'SURANGAMA_SUTRA_VOLUMES_1-4_FORMATTED.docx'
doc.save(output_file)

print(f"✅ Professional Word document created: {output_file}")
print(f"📄 Font: Garamond 12pt")
print(f"📐 Margins: 1.25\" left/right, 1\" top/bottom")
print(f"📑 Page numbers in footer")
print(f"🎨 Complete formatting:")
print(f"   - NO ## symbols anywhere (all stripped)")
print(f"   - Prose flows naturally")
print(f"   - Paragraph breaks only at blank lines")
print(f"   - No markdown syntax visible")
