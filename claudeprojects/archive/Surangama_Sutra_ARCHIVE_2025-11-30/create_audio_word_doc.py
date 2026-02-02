#!/usr/bin/env python3
"""
Create Word document from audio script for ElevenLabs TTS
Clean formatting optimized for text-to-speech processing
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Read the clean audio script
with open('SURANGAMA_SUTRA_AUDIO_SCRIPT.txt', 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set up styles for clean TTS reading
styles = doc.styles

# Configure Normal style - simple and clean for TTS
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Arial'
normal_font.size = Pt(12)
normal_paragraph = normal_style.paragraph_format
normal_paragraph.space_after = Pt(8)
normal_paragraph.line_spacing = 1.5

# Set page margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Process content line by line
lines = content.split('\n')

for line in lines:
    stripped = line.strip()

    # Skip completely empty lines
    if not stripped:
        doc.add_paragraph()
        continue

    # Title page elements (centered, larger)
    if 'THE SURANGAMA SUTRA' in stripped or 'Volumes 1-4' in stripped:
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(18)
        run.font.bold = True
        continue

    if 'Multi-Vocal American Translation' in stripped or 'Classical Chinese Text' in stripped:
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(12)
        run.font.italic = True
        continue

    if 'Translated by' in stripped or 'Houston, Texas' in stripped:
        p = doc.add_paragraph(stripped)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(11)
        continue

    # Major headings (VOLUME ONE, VOLUME TWO, etc.)
    if stripped in ['CONTENTS', 'FRONT MATTER', 'VOLUME ONE', 'VOLUME TWO',
                    'VOLUME THREE', 'VOLUME FOUR', 'BACK MATTER']:
        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        run = p.runs[0]
        run.font.size = Pt(16)
        run.font.bold = True
        continue

    # Section titles (all caps, longer phrases)
    if stripped.isupper() and len(stripped) > 10:
        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_before = Pt(12)
        run = p.runs[0]
        run.font.size = Pt(14)
        run.font.bold = True
        continue

    # Speaker labels (THE TEACHER:, ANANDA:, etc.)
    if ':' in stripped and stripped.split(':')[0].isupper():
        parts = stripped.split(':', 1)
        speaker = parts[0]
        text = parts[1].strip() if len(parts) > 1 else ''

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)

        # Bold speaker name
        run = p.add_run(speaker + ':')
        run.bold = True
        run.font.size = Pt(12)

        # Normal text after colon
        if text:
            run = p.add_run(' ' + text)
            run.font.size = Pt(12)
        continue

    # Stage directions (parenthetical)
    if stripped.startswith('(') and stripped.endswith(')'):
        p = doc.add_paragraph(stripped)
        run = p.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        p.paragraph_format.left_indent = Inches(0.5)
        continue

    # Voice annotations
    if stripped.startswith('Voice:') or stripped.startswith('Voices:'):
        p = doc.add_paragraph(stripped)
        run = p.runs[0]
        run.italic = True
        run.font.color.rgb = RGBColor(0, 100, 200)
        continue

    # Section dividers (---)
    if stripped.startswith('---'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        continue

    # Subsection headings (title case, shorter)
    if stripped[0].isupper() and not stripped.isupper() and len(stripped.split()) < 8:
        # Check if it looks like a heading
        if any(keyword in stripped for keyword in ['Part', 'Location', 'The ', 'Volume']):
            p = doc.add_paragraph(stripped)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.bold = True
            continue

    # Regular text - each line is its own paragraph for TTS pause points
    p = doc.add_paragraph(stripped)
    p.paragraph_format.space_after = Pt(6)

# Save document
output_file = 'SURANGAMA_SUTRA_AUDIO_ELEVENLABS.docx'
doc.save(output_file)

print(f"✅ ElevenLabs Word document created: {output_file}")
print(f"📢 Features:")
print(f"   - Clean Arial 12pt font (TTS-friendly)")
print(f"   - Each line = separate paragraph (natural pauses)")
print(f"   - Speaker labels in bold")
print(f"   - Stage directions in italic gray")
print(f"   - Voice annotations in blue italic")
print(f"   - 1.5 line spacing for readability")
print(f"   - Ready to upload to ElevenLabs")
