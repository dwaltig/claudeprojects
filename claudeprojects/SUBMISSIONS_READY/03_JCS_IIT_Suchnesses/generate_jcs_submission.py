#!/usr/bin/env python3
"""
Create professionally formatted Word document for JCS Submission.
Converts JCS_SUBMISSION_BLIND.md to .docx.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import os

def create_formatted_document(input_file, output_file):
    """Create a professionally formatted Word document from text file."""

    print(f"Reading from: {input_file}")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Set style to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')

    for line in lines:
        line = line.strip()
        
        # Skip empty lines (doc.add_paragraph() adds spacing already, but double spacing handles it?)
        # Actually, double spacing means we don't need extra empty paragraphs between logical paragraphs usually.
        # But markdown uses blank lines. Let's just skip empty lines to avoid huge gaps.
        if not line:
            continue
            
        # Check for headers
        is_header = False
        is_title = False
        
        if line.startswith('# '):
            is_title = True
            clean_line = line[2:].strip()
        elif line.startswith('## '):
            is_header = True
            clean_line = line[3:].strip()
        elif line.startswith('### '):
            is_header = True
            clean_line = line[4:].strip()
        else:
            clean_line = line

        p = doc.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        if is_title:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_header:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Optional: Bold headers? JCS usually standard.
        else:
            # Body paragraph
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_format.first_line_indent = Inches(0.5)

        # Handle simple bold/italics
        # Regex for **bold** and *italics*
        # We need a robust parser. For simplicity, let's adapt the previous one and add bold.
        
        process_line_formatting(p, clean_line, is_header or is_title)

    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def process_line_formatting(paragraph, text, is_header=False):
    # Split by bold/italic markers
    # This is a simple parser, might break on complex nested stuff but good enough for this file.
    # We'll prioritize italics as that's most common in this text.
    
    # Simple tokenization by *
    
    # Check for bold first ?? No, let's just stick to the script's italics support for now,
    # as the JCS text relies heavily on italics for terms (phi, samadhi).
    
    parts = re.split(r'(\*.*?\*)', text)
    
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
            if is_header:
                run.bold = True
        else:
            if not part: continue
            run = paragraph.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_header:
                run.bold = True

if __name__ == '__main__':
    base_dir = '/Users/williamaltig/claudeprojects/SUBMISSIONS_READY/03_JCS_IIT_Suchnesses'
    input_file = os.path.join(base_dir, 'JCS_SUBMISSION_BLIND.md')
    output_file = os.path.join(base_dir, 'JCS_SUBMISSION_REVISED_DEC30.docx')
    
    create_formatted_document(input_file, output_file)
