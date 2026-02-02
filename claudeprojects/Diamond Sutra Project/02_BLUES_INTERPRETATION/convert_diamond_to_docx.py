#!/usr/bin/env python3
"""
Convert The Diamond Sutra Reborn to DOCX format for ElevenReader
Clean formatting - keeps performance notes as italic text
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_docx():
    # Read the source file (TTS-ready: no --- lines, no performance notes)
    with open('DIAMOND_SUTRA_BLUES_TTS_READY.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content
    lines = content.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines but add paragraph break
        if not line:
            doc.add_paragraph()
            continue
        
        # Skip markdown decorative lines
        if line.startswith('═') or line.startswith('━'):
            continue
        
        # Section dividers
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('* * *')
            continue
        
        # Detect markdown H1 headers (# Title)
        if line.startswith('# '):
            title_text = line[2:]  # Remove "# "
            heading = doc.add_heading(title_text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Detect markdown H2 headers (## Title)
        if line.startswith('## '):
            title_text = line[3:]  # Remove "## "
            heading = doc.add_heading(title_text, level=2)
            continue
        
        # Detect markdown H3 headers (### Title)
        if line.startswith('### '):
            title_text = line[4:]  # Remove "### "
            heading = doc.add_heading(title_text, level=3)
            continue
        
        # Clean markdown bold
        clean_line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        
        # Performance notes in parentheses - render as italic paragraph
        if re.match(r'^\*?\(.*\)\*?$', clean_line):
            note_text = re.sub(r'^\*|\*$', '', clean_line)
            p = doc.add_paragraph()
            run = p.add_run(note_text)
            run.italic = True
            continue
        
        # Regular paragraphs
        p = doc.add_paragraph(clean_line)
    
    # Save
    output_file = 'The_Diamond_Sutra_Reborn_ELEVENREADER.docx'
    doc.save(output_file)
    print(f"Created: {output_file}")
    print("Headers: Detected from markdown # ## ### syntax")
    print("Performance notes: Italic text")

if __name__ == '__main__':
    create_docx()
