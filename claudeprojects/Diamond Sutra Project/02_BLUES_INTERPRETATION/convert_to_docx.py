#!/usr/bin/env python3
"""
Convert Diamond Sutra Blues Complete markdown to professionally formatted Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Create document with custom styles"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Georgia'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    
    # Heading 1 - Chapter titles
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Georgia'
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 - Subtitles
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Georgia'
    h2_font.size = Pt(14)
    h2_font.italic = True
    h2_font.color.rgb = RGBColor(51, 51, 51)
    
    # Heading 3 - Section headers
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Georgia'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(102, 102, 102)
    
    # Normal text
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Garamond'
    normal_font.size = Pt(12)
    
    # Quote style for verses
    quote_style = styles['Quote']
    quote_font = quote_style.font
    quote_font.name = 'Garamond'
    quote_font.size = Pt(11)
    quote_font.italic = True
    
    return doc

def add_markdown_content(doc, markdown_file):
    """Parse markdown and add to document with formatting"""
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_verse = False
    in_prose = False
    skip_next_blank = False
    
    while i < len(lines):
        line = lines[i]
        
        # Main title
        if line.startswith('# THE DIAMOND SUTRA'):
            p = doc.add_paragraph(line[2:], style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            skip_next_blank = True
        
        # Subtitle
        elif line.startswith('## ') and 'Complete Translation' in line:
            p = doc.add_paragraph(line[3:], style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            skip_next_blank = True
        
        # Chapter headings (## CHAPTER)
        elif line.startswith('## CHAPTER') or line.startswith('## PART') or line.startswith('## CLOSING'):
            doc.add_page_break()
            doc.add_paragraph(line[3:], style='Heading 1')
            skip_next_blank = True
        
        # Subsection headings (###)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 2')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            skip_next_blank = True
        
        # Bold metadata
        elif line.startswith('**') and line.endswith('**') and len(line) < 100:
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            if 'Classical Chinese' in line or 'Translated' in line or 'Houston' in line:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section markers
        elif line.startswith('**[PROSE RENDERING]**'):
            doc.add_paragraph('[PROSE RENDERING]', style='Heading 3')
            in_prose = True
            in_verse = False
        
        elif line.startswith('**[VERSE RENDERING]**'):
            doc.add_paragraph('[VERSE RENDERING]', style='Heading 3')
            in_verse = True
            in_prose = False
        
        # Performance notes in italics
        elif line.startswith('*(') and line.endswith(')*'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()  # Just add space
        
        # Blank lines
        elif line.strip() == '':
            if not skip_next_blank:
                doc.add_paragraph()
            skip_next_blank = False
        
        # Regular content
        elif line.strip():
            # Verse lines (end with two spaces or are short poetic lines)
            if in_verse and not line.startswith('#'):
                p = doc.add_paragraph(line, style='Quote')
                p.paragraph_format.left_indent = Inches(0.5)
            # Prose
            else:
                # Handle bold text within paragraphs
                if '**' in line:
                    p = doc.add_paragraph()
                    parts = re.split(r'(\*\*.*?\*\*)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            run = p.add_run(part.strip('*'))
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    doc.add_paragraph(line)
        
        i += 1
    
    return doc

def main():
    """Main conversion function"""
    print("Creating Diamond Sutra Blues Complete Word document...")
    
    # Create document with styles
    doc = create_styled_document()
    
    # Add content
    markdown_file = 'DIAMOND_SUTRA_BLUES_COMPLETE.md'
    doc = add_markdown_content(doc, markdown_file)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Save document
    output_file = 'DIAMOND_SUTRA_BLUES_COMPLETE.docx'
    doc.save(output_file)
    
    print(f"✓ Document created: {output_file}")
    print(f"  - Professional formatting applied")
    print(f"  - Chapter breaks added")
    print(f"  - Verse sections styled")
    print(f"  - Ready for printing or distribution")

if __name__ == '__main__':
    main()
