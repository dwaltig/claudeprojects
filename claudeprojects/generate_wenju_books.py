import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_markdown_file(file_path):
    """Reads a markdown file and returns its content."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            # Remove navigation footers
            content = re.sub(r'\n---+\n\*\*Navigation:\*\*.*$', '', content, flags=re.DOTALL)
            content = re.sub(r'\[« Previous Fascicle\].*$', '', content, flags=re.DOTALL)
            # Remove source reference blocks if desired, or keep them. Keeping for now.
            return content
    except FileNotFoundError:
        print(f"Warning: File not found: {file_path}")
        return ""

def format_text(paragraph, text):
    """Parses simple markdown (**bold**, *italic*) and adds runs to the paragraph."""
    # Split by bold markers first
    # This regex splits by **...**
    parts_bold = re.split(r'(\*\*.*?\*\*)', text)
    
    for part_b in parts_bold:
        if part_b.startswith('**') and part_b.endswith('**'):
            # This is bold text
            clean_text = part_b[2:-2]
            # Handle italics inside bold? Or just add bold run. 
            # Simple version: just bold.
            run = paragraph.add_run(clean_text)
            run.bold = True
        else:
            # Check for italics in the non-bold parts
            parts_italic = re.split(r'(\*.*?\*)', part_b)
            for part_i in parts_italic:
                if part_i.startswith('*') and part_i.endswith('*'):
                    # This is italic text
                    clean_text = part_i[1:-1]
                    run = paragraph.add_run(clean_text)
                    run.italic = True
                else:
                    # Regular text
                    if part_i: # Avoid empty strings
                        paragraph.add_run(part_i)

def add_markdown_to_docx(doc, md_content):
    """Parses simple markdown and adds it to the docx object."""
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('---'):
            doc.add_paragraph('________________________________________________________________')
        elif line.startswith('> '):
            # Quotes
            p = doc.add_paragraph()
            p.style = 'Quote' 
            # Manually formatting quote content
            format_text(p, line[2:])
        elif line.startswith('* ') or line.startswith('- '):
            # Bullet list
            p = doc.add_paragraph(style='List Bullet')
            format_text(p, line[2:])
        elif re.match(r'^\d+\. ', line):
            # Ordered list
            match = re.match(r'^\d+\. ', line)
            prefix_len = len(match.group(0))
            content = line[prefix_len:]
            p = doc.add_paragraph(style='List Number')
            format_text(p, content)
        elif line.startswith('|'):
            # Table row - treating as monospace text
            if '---' in line:
                continue
            p = doc.add_paragraph(style='No Spacing')
            p.add_run(line) # No formatting in tables for simplicity/alignment preservation
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            format_text(p, line)

def generate_book(title, output_filename, file_list):
    doc = Document()
    
    # Title Page
    doc.add_heading(title, 0)
    doc.add_paragraph('\n\nGenerated: January 26, 2026\n\n')
    doc.add_page_break()

    for file_path in file_list:
        print(f"Processing: {file_path}")
        content = read_markdown_file(file_path)
        if content:
            add_markdown_to_docx(doc, content)
            doc.add_page_break()
    
    doc.save(output_filename)
    print(f"Successfully created: {output_filename}")

# --- Configuration ---

base_dir = "Tiantai_Teachings_Project"
scholarly_dir = os.path.join(base_dir, "01_TRANSLATIONS/The_Words_and_Phrases/Scholarly")
practitioner_dir = os.path.join(base_dir, "01_TRANSLATIONS/The_Words_and_Phrases/Practitioner")
docs_dir = os.path.join(base_dir, "03_DOCUMENTATION")

scholarly_files = [
    os.path.join(scholarly_dir, "WENJU_PUBLICATION_PREFACE.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_01_Scholarly_Complete.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_02_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_03_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_04_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_05_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_06_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_07_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_08_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_09_Scholarly.md"),
    os.path.join(scholarly_dir, "Wenju_Fascicle_10_Scholarly.md"),
    os.path.join(docs_dir, "WENJU_SCHOLARLY_GLOSSARY.md")
]

practitioner_files = [
    os.path.join(practitioner_dir, f"Wenju_Fascicle_{i:02d}_Practitioner.md") for i in range(1, 11)
]

# --- Execution ---

print("Generating Scholarly Edition...")
generate_book(
    "The Words and Phrases of the Lotus Sutra\n(Scholarly Study Edition)",
    os.path.join(scholarly_dir, "Fahua_Wenju_Scholarly_Complete.docx"),
    scholarly_files
)

print("\nGenerating Practitioner Manual...")
generate_book(
    "The Practitioner's Manual to the Fahua Wenju\n(The Blues Track)",
    os.path.join(practitioner_dir, "Fahua_Wenju_Practitioner_Complete.docx"),
    practitioner_files
)