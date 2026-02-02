#!/usr/bin/env python3
"""
Convert The Immeasurable Meanings Sutra Reborn to DOCX format for ElevenReader
Uses proper H1/H2 heading styles for chapter recognition
Source file already has Title Case headers
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx():
    # Read the source file
    with open('The Immeasurable Meanings Sutra Reborn - ELEVENREADER.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content
    lines = content.split('\n')
    
    # H1 keywords (main title and chapters)
    h1_keywords = [
        'The Immeasurable Meanings Sutra Reborn',
        'Chapter One:',
        'Chapter Two:',
        'Chapter Three:',
        'Preface to the Immeasurable Meanings Sutra'
    ]
    
    # H2 keywords (subsections)
    h2_keywords = [
        'Publication Information',
        'Dedication',
        'Table of Contents',
        'Introduction',
        'Front Matter',
        'The Sutra Text',
        'A Meeting of Two Sacred Traditions',
        'Why Blues?',
        'The Source Text',
        'How to Read This Translation',
        'A Note on Sacred Translation',
        'An Invitation',
        'Translation and Editorial Notes:',
        'An American Blues Interpretation'
    ]
    
    for line in lines:
        line = line.rstrip()
        
        # Skip empty lines but add paragraph break
        if not line:
            doc.add_paragraph()
            continue
        
        # Section dividers - replace with visual separator
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('* * *')
            continue
        
        # Check for H1 (major chapters/sections)
        is_h1 = False
        for keyword in h1_keywords:
            if line.startswith(keyword) or line == keyword:
                is_h1 = True
                break
        
        if is_h1:
            # Use Heading 1 style (already in Title Case)
            heading = doc.add_heading(line, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        # Check for H2 (subsections)
        is_h2 = False
        for keyword in h2_keywords:
            if line == keyword or line.startswith(keyword):
                is_h2 = True
                break
        
        if is_h2:
            # Use Heading 2 style (already in Title Case)
            heading = doc.add_heading(line, level=2)
            continue
        
        # Regular paragraphs
        p = doc.add_paragraph(line)
    
    # Save
    output_file = 'The_Immeasurable_Meanings_Sutra_Reborn_ELEVENREADER.docx'
    doc.save(output_file)
    print(f"Created: {output_file}")
    print("Headers: Title Case (from source)")
    print("Heading structure: H1 for chapters, H2 for subsections")

if __name__ == '__main__':
    create_docx()
