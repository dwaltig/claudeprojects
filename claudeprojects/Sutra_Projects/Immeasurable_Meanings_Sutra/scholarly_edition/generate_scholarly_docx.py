#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def create_scholarly_docx(input_md, output_docx):
    """
    Converts IMS Scholarly Markdown to a professionally formatted Word document.
    Handles:
    - Headers (# and ##)
    - Blockquotes (>)
    - Italics (*text*)
    - Standard body text
    - Academic formatting (Margins, Font, Spacing)
    """
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(input_md, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        
        # Skip empty lines but add spacing
        if not line:
            doc.add_paragraph()
            continue

        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE # Single for poetry/verses, maybe adjust for prose?
        # User usually prefers double for scholarship? Let's check common project styles.
        # Actually, let's go with 1.15 or single for sutra text to keep it readable.
        pf.space_after = Pt(6)

        # Handle Headers
        if line.startswith('# '):
            p.text = line[2:]
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = 'Times New Roman'
        elif line.startswith('## '):
            p.text = line[3:]
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
        elif line.startswith('### '):
            p.text = line[4:]
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
        # Handle Blockquotes (Verses)
        elif line.startswith('>'):
            content = line[1:].strip()
            pf.left_indent = Inches(0.5)
            # pf.alignment = WD_ALIGN_PARAGRAPH.CENTER # Often verses are centered?
            process_formatting(p, content, italic_default=True)
        # Regular Paragraph
        else:
            process_formatting(p, line)

    doc.save(output_docx)
    print(f"Successfully generated: {output_docx}")

def process_formatting(paragraph, text, italic_default=False):
    """Processes markdown italics and set font."""
    # Simple regex for *italic*
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if italic_default:
                run.italic = True
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

if __name__ == "__main__":
    import os
    input_file = "/Users/williamaltig/claudeprojects/Sutra_Projects/Immeasurable_Meanings_Sutra/scholarly_edition/CHAPTER_01_VIRTUOUS_CONDUCT_SCHOLARLY.md"
    output_file = "/Users/williamaltig/claudeprojects/Sutra_Projects/Immeasurable_Meanings_Sutra/scholarly_edition/CHAPTER_01_VIRTUOUS_CONDUCT_SCHOLARLY.docx"
    
    if os.path.exists(input_file):
        create_scholarly_docx(input_file, output_file)
    else:
        print(f"Input file not found: {input_file}")
