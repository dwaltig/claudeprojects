#!/usr/bin/env python3
"""
Combine all 10 Blues chapters into a single Word document for ElevenReader TTS.
ElevenReader Formatting Expert Rules Applied:
- No Table of Contents (TTS will read it awkwardly)
- Chapter titles as Heading 1
- Section titles as Heading 2
- Endnotes/Notes moved to APPENDIX at the end
- Clean paragraphs for natural reading
- Removal of footnote superscript markers
- Removal of Markdown artifacts
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import re

# Paths
BLUES_DIR = "/Users/williamaltig/claudeprojects/Sutra_Projects/Bodhicaryavatara/01_TRANSLATIONS/Blues"
OUTPUT_FILE = "/Users/williamaltig/claudeprojects/Sutra_Projects/Bodhicaryavatara/Bodhicaryavatara_Blues_Complete.docx"

# Chapter files in order
CHAPTER_FILES = [
    "Chapter_01_Bodhicitta_Benefits_Blues.md",
    "Chapter_02_Confession_Blues.md",
    "Chapter_03_Embracing_Bodhicitta_Blues.md",
    "Chapter_04_Heedfulness_Blues.md",
    "Chapter_05_Vigilance_Blues.md",
    "Chapter_06_Patience_Blues.md",
    "Chapter_07_Heroic_Effort_Blues.md",
    "Chapter_08_Meditation_Blues.md",
    "Chapter_09_Wisdom_Blues.md",
    "Chapter_10_Dedication_Blues.md",
]

def clean_for_tts(text):
    """Remove all TTS-unfriendly elements."""
    # Remove footnote superscript markers (¹²³⁴⁵⁶⁷⁸⁹⁰)
    text = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰]+', '', text)
    # Remove bold markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic markers
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Remove horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    # Clean up multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def process_chapter(doc, filepath, chapter_num):
    """Process a single chapter file, add content to doc, and return extracted notes."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Clean for TTS
    content = clean_for_tts(content)
    lines = content.split('\n')
    
    chapter_notes = []
    capture_endnotes = False
    skip_rest = False
    
    chapter_title = f"Chapter {chapter_num}" # Default
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Skip everything after Blues interpretation completed line
        if line.startswith("Blues interpretation completed") or line.startswith("**Blues interpretation"):
            skip_rest = True
            continue
        if skip_rest:
            continue
            
        # Detect start of Endnotes
        if (line == '# Endnotes' or line == 'Endnotes' or line == '## Endnotes'):
            capture_endnotes = True
            continue # Don't add header to doc, don't add to notes list yet (header itself isn't needed in list if we use dict key)
            
        # Detect start of Theological Notes - STOP capturing endnotes, START adding to doc
        if ('# Theological Notes' in line or 'Theological Notes' in line or line == '## Theological Notes'):
            capture_endnotes = False # Stop capturing endnotes
            doc.add_heading('Theological Notes', level=2)
            continue

        if capture_endnotes:
            # Add to notes buffer for Appendix
            chapter_notes.append(line)
            continue
            
        # Capture Chapter Title for the Appendix Header
        if line.startswith('# Chapter'):
            chapter_title = line.replace('# ', '').replace(':', ' -')
            
            # Add page break before chapters (except first)
            if chapter_num > 1:
                doc.add_page_break()
            doc.add_heading(line.replace('# ', ''), level=1)
            continue
        
        # Subtitle (## Vīryapāramitā — ...)
        if line.startswith('## ') and ('—' in line or 'Pāramitā' in line.lower() or 'parigraha' in line.lower() or 'deśanā' in line.lower() or 'anuśaṃsa' in line.lower()):
            # This is the Sanskrit subtitle - add as italic paragraph
            p = doc.add_paragraph()
            run = p.add_run(line.replace('## ', ''))
            run.italic = True
            continue
        
        # Author line - SKIP (remove AI agent credits)
        if line.startswith('A Blues Interpretation by') or line.startswith('*A Blues Interpretation'):
            continue
        
        # Section headers (## Something) -> Heading 2
        if line.startswith('## '):
            section_title = line.replace('## ', '')
            doc.add_heading(section_title, level=2)
            continue
        
        # Skip standalone hash symbols
        if line == '#':
            continue
            

        
        # Regular paragraph (includes Theological Notes content)
        doc.add_paragraph(line)
        
    return chapter_title, chapter_notes

def create_combined_docx():
    """Create the combined Word document optimized for ElevenReader."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title page (simple, TTS-friendly)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("A Guide to the Bodhisattva's Way of Life")
    run.bold = True
    run.font.size = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Bodhicaryāvatāra")
    run.italic = True
    run.font.size = Pt(18)
    
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("by Śāntideva")
    run.font.size = Pt(14)
    
    translator = doc.add_paragraph()
    translator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = translator.add_run("\n\nBlues Vernacular Interpretation")
    run.font.size = Pt(14)
    
    # NO TABLE OF CONTENTS - TTS doesn't need it
    
    all_notes = []
    
    # Add each chapter
    for i, filename in enumerate(CHAPTER_FILES):
        filepath = os.path.join(BLUES_DIR, filename)
        
        if not os.path.exists(filepath):
            print(f"Warning: {filepath} not found, skipping...")
            continue
        
        ch_title, ch_notes = process_chapter(doc, filepath, i + 1)
        all_notes.append((ch_title, ch_notes))
        print(f"Processed: Chapter {i + 1} (Extracted {len(ch_notes)} note lines)")
    
    # ==========================================
    # APPENDIX: COMPREHENSIVE ENDNOTES
    # ==========================================
    if any(notes for _, notes in all_notes):
        doc.add_page_break()
        doc.add_heading("Appendix: Comprehensive Endnotes", level=1)
        doc.add_paragraph("Identified by Chapter for Reference.")
        
        for ch_title, notes in all_notes:
            if not notes:
                continue
            
            doc.add_heading(ch_title, level=2)
            
            for line in notes:
                # Handle Note Subheaders (Heading 3)
                if line.startswith('# Endnotes') or line == 'Endnotes' or line.startswith('## Endnotes'):
                     # We skip the main 'Endnotes' header since we have the Chapter Title
                     continue
                
                if 'Theological Notes' in line:
                    doc.add_heading("Theological Notes", level=3)
                    continue
                
                # Format actual notes (bold term -> bold run)
                # Example: "¹ **"Term"**: Definition..." -> "Term: Definition..." (cleaner)
                # But we already stripped ¹ symbols in clean_for_tts
                
                # Bold Logic: If line starts with "**Term**:", allow it to be bold
                # Check for "**Term**:" pattern
                match = re.match(r'^\*\*(.+?)\*\*:(.*)', line)
                if match:
                    # Note Text
                    term = match.group(1)
                    definition = match.group(2)
                    p = doc.add_paragraph()
                    run = p.add_run(term + ":")
                    run.bold = True
                    p.add_run(definition)
                else:
                    doc.add_paragraph(line)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"\nCreated: {OUTPUT_FILE}")
    print("ElevenReader formatting applied:")
    print("  - No Table of Contents")
    print("  - Chapter titles as Heading 1")
    print("  - Section titles as Heading 2")
    print("  - Endnotes moved to APPENDIX")
    print("  - Footnote markers removed")
    print("  - Clean paragraphs for natural TTS reading")

if __name__ == "__main__":
    create_combined_docx()
