
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE = '01_TRANSLATIONS/Shikshasamucchaya_Scholarly_Edition.md'
OUTPUT_FILE = 'Shikshasamucchaya_Scholarly_Edition_Publication_Ready.docx'


def add_markdown_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    
    # Very basic Markdown parser for *italic* and **bold**
    # We will split strictly by ** first, then *
    
    # 1. Handle Bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold content
            content = part[2:-2]
            # Check for nested italic inside bold? (Simpler to ignore nesting for now or do 2nd pass)
            # Let's do a simple 2nd pass for italic inside bold
            sub_parts = re.split(r'(\*.*?\*)', content)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*'):
                    run = p.add_run(sub[1:-1])
                    run.bold = True
                    run.italic = True
                else:
                    if sub:
                        run = p.add_run(sub)
                        run.bold = True
        else:
            # Non-bold content, check for Italic (*text*)
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*'):
                    # Italic content
                    run = p.add_run(sub[1:-1])
                    run.italic = True
                else:
                    # Normal text
                    if sub:
                         # Remove Links [text](url) -> text (visual only)
                         clean = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', sub)
                         # Remove HTML
                         clean = re.sub(r'<.*?>', '', clean)
                         p.add_run(clean)

def create_scholarly_docx():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    skip_mode = False
    skip_toc = False
    title_page_done = False
    
    # Add Title Page
    doc.add_heading('Śikṣāsamuccaya: The Compendium of Training', 0)
    p = doc.add_paragraph('A Scholarly Translation')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    p = doc.add_paragraph('Translation and Commentary by Dr. Rajesh Patel')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    title_page_done = True

    for line in lines:
        stripped_line = line.strip()
        
        if not stripped_line:
            continue

        # Skip TOC
        if stripped_line.startswith("## Table of Contents"):
            skip_toc = True
            continue
        if skip_toc:
            if stripped_line.startswith('---') or (stripped_line.startswith('#') and "Chapter" in stripped_line):
                skip_toc = False
            else:
                continue

        # Skip Agent Notes
        if "Soul-Check Hints" in stripped_line:
            skip_mode = True
            continue
        if skip_mode:
            if stripped_line.startswith('#') or stripped_line.startswith('---'):
                skip_mode = False
            else:
                continue

        # Handle Headers
        if stripped_line.startswith('#'):
            level = 0
            text = stripped_line
            if stripped_line.startswith('# '):
                level = 1
                text = stripped_line[2:]
            elif stripped_line.startswith('## '):
                level = 2
                text = stripped_line[3:]
            elif stripped_line.startswith('### '):
                level = 3
                text = stripped_line[4:]
            elif stripped_line.startswith('#### '):
                level = 4
                text = stripped_line[5:]
            
            text = text.strip()
            
            if "Śikṣāsamuccaya" in text and level == 1: continue
            if "Scholarly Translation" in text and level == 2: continue

            if "Chapter" in text and level == 1:
                doc.add_page_break()

            doc.add_heading(text, level=level)

        # Handle Blockquotes
        elif stripped_line.startswith('>'):
            clean_text = stripped_line.lstrip('> ').strip()
            add_markdown_paragraph(doc, clean_text, style='Quote')
            
        # Handle Lists
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
             add_markdown_paragraph(doc, stripped_line[2:], style='List Bullet')
        
        # Normal Text
        else:
             add_markdown_paragraph(doc, stripped_line)

    doc.save(OUTPUT_FILE)
    print(f"Created Publication DOCX: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_scholarly_docx()

