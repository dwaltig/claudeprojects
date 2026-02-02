import re
import os
from docx import Document
from docx.shared import Pt

# Configuration
INPUT_FILE = "/Users/williamaltig/claudeprojects/Sutra_Projects/Shikshasamucchaya/01_TRANSLATIONS/Shikshasamucchaya_Blues_Edition.md"
OUTPUT_FILE = "/Users/williamaltig/claudeprojects/Sutra_Projects/Shikshasamucchaya/01_TRANSLATIONS/Shikshasamucchaya_The_Long_Game_Audio.docx"

def clean_for_tts(text):
    """
    Cleans markdown text for Text-to-Speech readability.
    """
    # 1. Remove footnote markers [^1], [1], etc. AND Unicode superscripts (¹²³⁴⁵⁶⁷⁸⁹⁰)
    text = re.sub(r'\[\^?\d+\]', '', text)
    text = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰]', '', text)
    
    # 2. Remove bold/italic markers (*, **) but keep content
    # Handle bold **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    # Handle italic *text*
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # 3. Remove links [text](url) -> text
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    
    # 4. Remove HTML tags
    text = re.sub(r'<.*?>', '', text)
    
    # 5. Remove horizontal rules
    if re.match(r'^[-*_]{3,}\s*$', text):
        return None
        
    # 6. Clean blockquote markers (Aggressive)
    # Remove leading > and optional space
    text = re.sub(r'^>\s?', '', text)
        
    return text.strip()

def create_audio_docx():
    doc = Document()
    
    # Set default style to be readable
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    skip_toc = False
    title_found = False
    first_chapter_found = False
    
    for line in lines:
        stripped_line = line.strip()
        
        # Skip empty lines
        if not stripped_line:
            continue
            
        # Skip redundant intro blurbs (User requested removal/endnote)
        if stripped_line.startswith("*A Blues vernacular translation"):
            continue
            
        # Detect Start of TOC
        if stripped_line == "## Table of Contents":
            skip_toc = True
            continue
            
        # Detect End of TOC (The first real content after TOC usually starts with a Chapter or an HR)
        if skip_toc:
            if stripped_line.startswith("# Chapter") or stripped_line.startswith("## Note on the Translation"):
                skip_toc = False
            else:
                continue

        # Skip links that look like TOC entries [Chapter X...](#chapter-x)
        if re.match(r'^- \[.*?\]\(#.*?\)$', stripped_line):
           continue

        # Headers -> Heading Styles
        if stripped_line.startswith('#'):
            # Remove Sanskrit/Parentheticals
            clean_text = re.sub(r'\s*\(.*?\)', '', stripped_line.lstrip('#').strip())
            clean_text = clean_for_tts(clean_text)
            
            if not clean_text:
                continue

            # Special Title Page Logic
            if not title_found:
                # First H1 is the Title
                if stripped_line.startswith('# '):
                    doc.add_heading(clean_text, level=0) # Level 0 = Title Style
                    title_found = True
                    continue
            
            # Subtitles on Title Page (H2s before any chapters)
            if title_found and not first_chapter_found:
                if stripped_line.startswith('## '):
                    p = doc.add_paragraph(clean_text)
                    p.style = 'Subtitle'
                    continue
                # If we hit the first Chapter, mark it
                if "Chapter" in stripped_line or "Note on" in stripped_line:
                    first_chapter_found = True
                    # Insert Page Break before first real content
                    doc.add_page_break()

            # Flatten Hierarchy for Professional Audio Look
            # H1 -> Heading 1 (Chapter)
            # H2 -> Heading 2 (Part)
            # H3/H4 -> Bold Paragraphs (Not Headings)
            
            level = 0
            if stripped_line.startswith('# '): 
                level = 1
            elif stripped_line.startswith('## '): 
                level = 2
            # Note: We purposely do NOT map H3/H4 to Headings to avoid "Title Stack"
            
            if level > 0:
                doc.add_heading(clean_text, level=level)
            elif stripped_line.startswith('### ') or stripped_line.startswith('#### '):
                # Add as Bold Paragraph
                p = doc.add_paragraph()
                runner = p.add_run(clean_text)
                runner.bold = True
                p.style = 'Normal' # Ensure it's not a heading
        
        # Lists
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            clean_text = clean_for_tts(stripped_line[2:])
            if clean_text:
                doc.add_paragraph(clean_text, style='List Bullet')
        elif re.match(r'^\d+\. ', stripped_line):
            # Ordered list
            clean_text = clean_for_tts(stripped_line)
            if clean_text:
                doc.add_paragraph(clean_text, style='List Number')
                
        # Tables (TTS hates tables, convert to text or skip)
        elif stripped_line.startswith('|'):
            # Simple approach: Replace pipes with commas or spaces for flow
            clean_text = stripped_line.replace('|', ' ').strip()
            # Remove dashed divider lines
            if '---' in clean_text:
                continue
            clean_text = clean_for_tts(clean_text)
            if clean_text:
                doc.add_paragraph(clean_text)
                
        # Normal Para
        else:
            clean_text = clean_for_tts(stripped_line)
            if clean_text:
                doc.add_paragraph(clean_text)

    doc.save(OUTPUT_FILE)
    print(f"Created audio-ready DOCX: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_audio_docx()
