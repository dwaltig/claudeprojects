#!/usr/bin/env python3
"""
Create ElevenReader-optimized Word documents for Śikṣāsamuccaya Blues translations.
Generates DOCXs for all completed chapters.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

PROJECT_DIR = "/Users/williamaltig/claudeprojects/Sutra_Projects/Shikshasamucchaya/01_TRANSLATIONS"

CHAPTERS = [
    {
        "source": "Chapter_01_Danapramita_Blues_Complete.md",
        "output": "Chapter_01_Danapramita_Blues_ELEVENREADER.docx",
        "title": "Chapter 1: The Road of Generosity"
    },
    {
        "source": "Chapter_02_Saddharmaparigraha_Blues.md",
        "output": "Chapter_02_Saddharmaparigraha_Blues_ELEVENREADER.docx",
        "title": "Chapter 2: Holding On to What's Real"
    },
    {
        "source": "Chapter_03_Dharmabhanakadiraksa_Blues.md",
        "output": "Chapter_03_Dharmabhanakadiraksa_Blues_ELEVENREADER.docx",
        "title": "Chapter 3: Watching Your Back"
    },
]

def clean_for_tts(text):
    """Remove all TTS-unfriendly elements."""
    text = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰]+', '', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^#+\s*$', '', text, flags=re.MULTILINE)
    text = text.replace('Master John Kim', 'William Altig')
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def create_chapter_docx(chapter_info):
    """Create a Word document optimized for ElevenReader."""
    source_path = os.path.join(PROJECT_DIR, chapter_info["source"])
    output_path = os.path.join(PROJECT_DIR, chapter_info["output"])
    
    if not os.path.exists(source_path):
        print(f"  Warning: {source_path} not found, skipping...")
        return False
    
    with open(source_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    content = clean_for_tts(content)
    lines = content.split('\n')
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Compendium of Training")
    run.bold = True
    run.font.size = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Śikṣāsamuccaya")
    run.italic = True
    run.font.size = Pt(18)
    
    author_line = doc.add_paragraph()
    author_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_line.add_run("by Śāntideva")
    run.font.size = Pt(14)
    
    chapter_line = doc.add_paragraph()
    chapter_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = chapter_line.add_run(f"\n\n{chapter_info['title']}")
    run.bold = True
    run.font.size = Pt(16)
    
    translator_line = doc.add_paragraph()
    translator_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = translator_line.add_run("\n\nBlues Vernacular Translation by William Altig")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    endnotes_buffer = []
    theological_buffer = []
    capture_endnotes = False
    capture_theological = False
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        if 'Blues Translation by' in stripped or 'Blues Vernacular Translation by' in stripped:
            continue
        
        # Part headings (# Part I: etc)
        if stripped.startswith('# Part') or (stripped.startswith('# ') and 'Part' in stripped):
            section = stripped.replace('# ', '').strip()
            doc.add_heading(section, level=1)
            continue
        
        # Chapter title (skip, already in title page)
        if stripped.startswith('# Chapter'):
            continue
        
        # Endnotes section
        if stripped == '## Endnotes' or stripped == 'Endnotes':
            capture_endnotes = True
            capture_theological = False
            continue
        
        # Theological Notes section
        if stripped == '## Theological Notes' or 'Theological Notes' in stripped:
            capture_endnotes = False
            capture_theological = True
            continue
        
        # End of chapter marker
        if stripped.startswith('Thus ends') or stripped.startswith('*Thus ends'):
            capture_endnotes = False
            capture_theological = False
            continue
        
        # Capture notes for appendix
        if capture_endnotes:
            endnotes_buffer.append(stripped)
            continue
        
        if capture_theological:
            theological_buffer.append(stripped)
            continue
        
        # Section headings (## Something)
        if stripped.startswith('## '):
            section = stripped.replace('## ', '')
            doc.add_heading(section, level=2)
            continue
        
        # Subsection headings (### Something)
        if stripped.startswith('### '):
            subsection = stripped.replace('### ', '')
            doc.add_heading(subsection, level=3)
            continue
        
        # Skip standalone markers
        if stripped == '#' or stripped == '---':
            continue
        
        # Regular paragraph
        doc.add_paragraph(stripped)
    
    # Appendix: Endnotes
    if endnotes_buffer:
        doc.add_page_break()
        doc.add_heading("Appendix: Endnotes", level=1)
        for line in endnotes_buffer:
            doc.add_paragraph(line)
    
    # Appendix: Theological Notes
    if theological_buffer:
        doc.add_page_break()
        doc.add_heading("Appendix: Theological Notes", level=1)
        for line in theological_buffer:
            if line.startswith('### '):
                doc.add_heading(line.replace('### ', ''), level=2)
            else:
                doc.add_paragraph(line)
    
    doc.save(output_path)
    return True

def main():
    print("Creating ElevenReader-optimized Śikṣāsamuccaya Blues DOCXs...")
    print("=" * 60)
    
    for chapter in CHAPTERS:
        print(f"\nProcessing: {chapter['title']}...")
        success = create_chapter_docx(chapter)
        if success:
            print(f"  ✓ Created: {chapter['output']}")
        else:
            print(f"  ✗ Failed: {chapter['output']}")
    
    print("\n" + "=" * 60)
    print("ElevenReader formatting applied:")
    print("  - Title pages with William Altig authorship")
    print("  - Part titles as Heading 1")
    print("  - Section titles as Heading 2")
    print("  - Endnotes moved to APPENDIX")
    print("  - Theological Notes moved to APPENDIX")
    print("  - Footnote markers removed")
    print("  - Clean paragraphs for TTS")

if __name__ == "__main__":
    main()
