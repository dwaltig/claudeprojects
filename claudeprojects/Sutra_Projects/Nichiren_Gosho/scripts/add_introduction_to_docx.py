import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

DOCX_PATH = Path('/Users/williamaltig/claudeprojects/Sutra_Projects/Nichiren_Gosho/03_FINAL_DOCUMENTS/KAIMOKU_SHO_AUDIO_SCRIPT.docx')

if not DOCX_PATH.exists():
    print(f'Error: {DOCX_PATH} not found')
    sys.exit(1)

doc = Document(DOCX_PATH)

# Create introduction content
intro_heading = 'Introduction'
intro_text = (
    "This audio version of the Kaimoku Sho provides a guided reading of the text, "
    "structured with clear chapter headings to facilitate listening. "
    "The following sections correspond to the original script prepared for ElevenReader, "
    "with consistent Heading 1 levels for each chapter."
)

# Insert at the beginning: a heading and a paragraph
# Insert new paragraph at start of document
first_paragraph = doc.paragraphs[0]
# Insert heading before first paragraph
heading_para = doc.add_paragraph()
heading_para.style = doc.styles['Heading 1']
heading_para.add_run(intro_heading)
# Insert introduction paragraph after heading
intro_para = doc.add_paragraph()
intro_para.add_run(intro_text)

# Move the new heading and intro to the top by reordering elements
# doc._body is the underlying XML; easiest is to rebuild document
new_doc = Document()
# Add heading and intro first
new_doc.add_paragraph(intro_heading, style='Heading 1')
new_doc.add_paragraph(intro_text)
# Then copy existing content
for element in doc.element.body:
    new_doc.element.body.append(element)

# Save back to same file (overwrite)
new_doc.save(DOCX_PATH)
print('Added introduction heading and paragraph to the DOCX')
