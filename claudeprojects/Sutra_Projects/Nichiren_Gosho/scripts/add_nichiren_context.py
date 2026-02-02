import sys
from pathlib import Path
from docx import Document

DOCX_PATH = Path('/Users/williamaltig/claudeprojects/Sutra_Projects/Nichiren_Gosho/03_FINAL_DOCUMENTS/KAIMOKU_SHO_AUDIO_SCRIPT.docx')

if not DOCX_PATH.exists():
    print(f'Error: {DOCX_PATH} not found')
    sys.exit(1)

doc = Document(DOCX_PATH)

# Content to add
heading_text = 'About Nichiren and This Work'
paragraph_text = (
    "Nichiren (1222–1282) was a Japanese Buddhist monk who founded the Nichiren tradition, "
    "centered on the Lotus Sutra as the ultimate teaching.  His treatise, the Kaimoku Sho, "
    "offers a concise exposition of the core principles of the Lotus Sutra, emphasizing the "
    "power of chanting the daimoku and the importance of faith in the universal truth of the "
    "Buddha’s teaching.  This audio reading presents the text in a clear, structured format, "
    "allowing listeners to grasp the historical context, the philosophical depth, and the "
    "practical guidance that Nichiren provides for personal transformation and societal harmony."
)

# Insert after the existing introduction heading and paragraph
# Find the index of the introduction heading (Heading 1 with text 'Introduction')
intro_idx = None
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1' and para.text.strip() == 'Introduction':
        intro_idx = i
        break

if intro_idx is None:
    print('Introduction heading not found; appending at end')
    # Append at end
    doc.add_paragraph(heading_text, style='Heading 2')
    doc.add_paragraph(paragraph_text)
else:
    # Insert heading and paragraph after the intro paragraph (which is intro_idx+1)
    # docx does not support direct insertion, so we rebuild the document
    new_doc = Document()
    for i, para in enumerate(doc.paragraphs):
        # copy existing paragraph
        new_doc.add_paragraph(para.text, style=para.style.name)
        # after intro paragraph, insert new content
        if i == intro_idx + 1:
            new_doc.add_paragraph(heading_text, style='Heading 2')
            new_doc.add_paragraph(paragraph_text)
    # Save
    doc = new_doc

# Overwrite the file
doc.save(DOCX_PATH)
print('Added Nichiren context section to the DOCX')
