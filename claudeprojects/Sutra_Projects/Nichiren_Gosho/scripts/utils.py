import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Convert a Markdown file to DOCX using pandoc."""
    subprocess.run([
        "pandoc",
        str(md_path),
        "-o",
        str(docx_path),
        "--metadata",
        "title=Kaimoku Sho Audio Script",
    ], check=True)


def refine_docx_structure(docx_path: Path):
    """
    Surgically fix the DOCX structure:
    1. Remove 'Part X' distracting headers.
    2. Normalize Chapter headers to Heading 1.
    3. Ensure Title and Copyright are at the top.
    4. Insert/Update Introduction and Nichiren Context after Copyright.
    5. Remove any duplicates of these sections.
    """
    doc = Document(docx_path)
    
    # 1. & 2. Filter/Normalize paragraphs
    # We will build a list of "good" paragraphs and then rebuild the document
    
    new_doc = Document()
    
    # Injected content definitions...
    intro_title = "Introduction"
    intro_content = (
        "This audio version of the Kaimoku Sho provides a guided reading of the text, "
        "structured with clear chapter headings to facilitate listening. "
        "The following sections correspond to the original script prepared for ElevenReader, "
        "with consistent Heading 1 levels for each chapter."
    )
    
    context_title = "About Nichiren and This Work"
    context_content = (
        "Nichiren (1222–1282) was a Japanese Buddhist monk who founded the Nichiren tradition, "
        "centered on the Lotus Sutra as the ultimate teaching. His treatise, the Kaimoku Sho, "
        "offers a concise exposition of the core principles of the Lotus Sutra, emphasizing the "
        "power of chanting the daimoku and the importance of faith in the universal truth of the "
        "Buddha’s teaching. This audio reading presents the text in a clear, structured format, "
        "allowing listeners to grasp the historical context, the philosophical depth, and the "
        "practical guidance that Nichiren provides for personal transformation and societal harmony."
    )

    # Keywords to strip out of any text
    distracting_phrases = [
        "Part One", "Part Two", "Part 1", "Part 2", 
        "End of Part One", "End of Part Two",
        "(Upper Volume)", "(Lower Volume)", "The Upper Volume", "The Lower Volume"
    ]
    
    processed_paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Keep occasional empty lines if they were intentional/spacing, 
            # but usually we want to skip if it's just garbage
            continue
            
        # Strip distracting phrases from the middle of sentences (e.g. Chapter titles)
        clean_text = text
        for phrase in distracting_phrases:
            # Try to catch variants like "End of Part One — "
            if phrase in clean_text:
                clean_text = clean_text.replace(phrase, "").replace(" — ", " ").replace("  ", " ").strip()
        
        # If the whole line was just a distraction, skip it
        if not clean_text or clean_text.upper() == "PART":
            continue
            
        # Ignore existing duplicates of our injected sections
        if clean_text == intro_title or clean_text == context_title:
            continue
        if clean_text == intro_content or clean_text == context_content:
            continue
            
        # Create a proxy for the paragraph with cleaned text
        para.text = clean_text
        
        # Normalize Chapters to Heading 1
        if clean_text.upper().startswith("CHAPTER"):
            para.style = doc.styles['Heading 1']
            
        processed_paragraphs.append(para)

    # Rebuild the document
    # 1. Add Title and Copyright first
    # 2. Add Introduction and Context
    # 3. Add the rest
    
    # We'll assume the first few paragraphs up to "Copyright" or the first date line are the header
    header_count = 0
    found_copyright = False
    for i, p in enumerate(processed_paragraphs):
        header_count += 1
        if "Copyright" in p.text or "All rights reserved" in p.text:
            found_copyright = True
            # Check next one if it's the specific copyright text
            continue
        if found_copyright and i < 5: # Keep a few lines of copyright info
            continue
        if i > 2 and not found_copyright: # Fallback if copyright not found
             header_count = 2
             break
        if found_copyright:
            break

    # Add header
    for i in range(header_count):
        p = processed_paragraphs[i]
        new_doc.add_paragraph(p.text, style=p.style.name)
        
    # Inject sections
    new_doc.add_paragraph(intro_title, style='Heading 1')
    new_doc.add_paragraph(intro_content)
    new_doc.add_paragraph(context_title, style='Heading 1') # Consistent level
    new_doc.add_paragraph(context_content)
    
    # Add the rest
    for i in range(header_count, len(processed_paragraphs)):
        p = processed_paragraphs[i]
        new_doc.add_paragraph(p.text, style=p.style.name)
        
    new_doc.save(docx_path)
    print(f"Refined structure and removed duplicates in {docx_path.name}")


def generate_pdf(docx_path: Path, pdf_path: Path):
    """Convert DOCX to PDF using pandoc."""
    try:
        subprocess.run([
            "pandoc",
            str(docx_path),
            "-o",
            str(pdf_path),
        ], check=True)
        print(f"Generated PDF at {pdf_path}")
    except Exception as e:
        print(f"PDF generation failed: {e}")
