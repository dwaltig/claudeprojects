#!/usr/bin/env python3
"""
Create publication-ready Word documents for Kaimoku Shō (開目抄).
Generates both Scholarly and Blues editions.

Formatting:
- Title page with full attribution
- Table of Contents
- Proper heading styles
- Footnotes converted to endnotes
- Clean paragraphs for reading
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# Paths
BASE_DIR = "/Users/williamaltig/claudeprojects/Sutra_Projects/Nichiren_Gosho/03_FINAL_DOCUMENTS"
SCHOLARLY_FILE = os.path.join(BASE_DIR, "KAIMOKU_SHO_COMPLETE_SCHOLARLY.md")
BLUES_FILE = os.path.join(BASE_DIR, "KAIMOKU_SHO_COMPLETE_BLUES.md")
OUTPUT_DIR = "/Users/williamaltig/claudeprojects/Sutra_Projects/Nichiren_Gosho/03_FINAL_DOCUMENTS"

def clean_markdown(text):
    """Remove markdown artifacts and clean for Word."""
    # Remove footnote superscript markers
    text = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰]+', '', text)
    # Remove bold markers but keep text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic markers but keep text
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove horizontal rules
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    # Remove blockquote markers
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    # Clean up multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def process_markdown_to_docx(md_file, output_file, title, subtitle, translator):
    """Convert a markdown file to a formatted Word document."""
    
    # Read the source file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create new document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # =====================
    # TITLE PAGE
    # =====================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    
    doc.add_paragraph()  # Spacer
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(18)
    
    doc.add_paragraph()  # Spacer
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_para.add_run("by Nichiren Daishōnin")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("Written at Sado Island, 1272")
    run.italic = True
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    trans_para = doc.add_paragraph()
    trans_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = trans_para.add_run(translator)
    run.font.size = Pt(12)
    
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_para.add_run("January 2026")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # =====================
    # PROCESS CONTENT
    # =====================
    lines = content.split('\n')
    in_footnotes = False
    footnotes = []
    current_section = ""
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            continue
        
        # Skip metadata lines at top
        if stripped.startswith('# Kaimoku Shō') or stripped.startswith('# 開目抄'):
             continue
        if stripped.startswith('**Original') or stripped.startswith('**Author'):
            continue
        if stripped.startswith('**Translator') or stripped.startswith('**Interpreter'):
            continue
        if stripped.startswith('**Date') or stripped.startswith('**Source'):
            continue
        if stripped.startswith('**Section') or stripped.startswith('**Pages'):
            continue
        if stripped.startswith('**Soul-Check'):
            continue
        if stripped.startswith('#### Copyright') or stripped.startswith('Copyright ©'):
            continue 
        if stripped.startswith('No part of this') or stripped.startswith('For permission'):
            continue

            
        # Part headers (Now H2: ## Part)
        if stripped.startswith('## Part') or stripped.startswith('## PART'):
            doc.add_page_break()
            doc.add_heading(stripped.replace('## ', ''), level=1)
            continue
        
        # Section headers (Now H3: ### Section X:)
        if stripped.startswith('### Section') or stripped.startswith('### Chapter'):
            doc.add_page_break()
            section_title = stripped.replace('### ', '')
            doc.add_heading(section_title, level=1)
            current_section = section_title
            continue
        
        # Sub-section headers (Title Case) (Now H4: #### Something)
        if stripped.startswith('#### '):
            heading_text = stripped.replace('#### ', '')
            doc.add_heading(heading_text, level=2)
            continue
            
        # Old Sub-section headers fallback (H3)
        if stripped.startswith('### '):
             heading_text = stripped.replace('### ', '')
             doc.add_heading(heading_text, level=2)
             continue
        
        # Footnote sections
        if stripped == '### Footnotes' or stripped == '## Footnotes':
            in_footnotes = True
            continue
        
        # Scholarly Apparatus section
        if stripped.startswith('## Scholarly Apparatus') or stripped.startswith('## Theological Notes'):
            in_footnotes = False
            doc.add_heading(stripped.replace('## ', ''), level=2)
            continue
        
        # End markers
        if stripped.startswith('## End of'):
            continue
        
        # Skip Markdown tables
        if stripped.startswith('|') and '---' not in stripped:
            continue
        if '|:' in stripped or ':-' in stripped:
            continue
        
        # Footnotes - collect for endnotes
        if in_footnotes:
            if stripped.startswith('¹') or stripped.startswith('²') or re.match(r'^[¹²³⁴⁵⁶⁷⁸⁹⁰]+\s', stripped):
                footnotes.append(stripped)
            elif stripped.startswith('**') and ':**' in stripped:
                footnotes.append(stripped)
            continue
        
        # Regular paragraph
        clean_text = clean_markdown(stripped)
        if clean_text and not clean_text.startswith('#'):
            doc.add_paragraph(clean_text)
    
    # =====================
    # ENDNOTES SECTION
    # =====================
    if footnotes:
        doc.add_page_break()
        doc.add_heading('Endnotes', level=1)
        
        for note in footnotes:
            clean_note = clean_markdown(note)
            doc.add_paragraph(clean_note)
    
    # =====================
    # SAVE
    # =====================
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    print(f"Created: {output_file}")
    return output_file

def main():
    """Generate both Scholarly and Blues editions."""
    
    # Scholarly Edition
    scholarly_output = os.path.join(OUTPUT_DIR, "Kaimoku_Sho_On_Opening_the_Eyes_Scholarly_Edition.docx")
    process_markdown_to_docx(
        SCHOLARLY_FILE,
        scholarly_output,
        title="On Opening the Eyes",
        subtitle="Kaimoku Shō (開目抄)",
        translator="Scholarly Translation by Dr. Rajesh Patel"
    )
    
    # Blues Edition
    blues_output = os.path.join(OUTPUT_DIR, "Kaimoku_Sho_On_Opening_the_Eyes_Blues_Edition.docx")
    process_markdown_to_docx(
        BLUES_FILE,
        blues_output,
        title="On Opening the Eyes",
        subtitle="Kaimoku Shō (開目抄)",
        translator="Blues Vernacular Interpretation by William Altig"
    )
    
    print("\n✅ Publication documents created:")
    print(f"  - {scholarly_output}")
    print(f"  - {blues_output}")

if __name__ == "__main__":
    main()
