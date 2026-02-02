from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_wisdom_questionnaire():
    doc = Document()
    
    # Title
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    heading = doc.add_paragraph('Wisdom Submission Questionnaire')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(16)
    
    subheading = doc.add_paragraph('Complete this questionnaire and email as a Word document to submissions@wisdompubs.org.')
    subheading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subheading.runs[0].bold = True
    
    doc.add_paragraph('_' * 80)

    # Content Data
    content = [
        ("ABOUT YOU", True),
        ("1. Your name, residence, phone number(s), and email(s)", False),
        ("Name: William Altig\nResidence: Houston, Texas, USA\nPhone: 832-212-0915\nEmail: dwaltig@gmail.com", "Answer"),
        
        ("2. Please list the names and residences of any co-contributors.", False),
        ("N/A", "Answer"),
        
        ("3. Write a 1-2-paragraph “About the Author”–style description of yourself.", False),
        ("William Altig is a scholar-practitioner and cultural translator who brings together deep knowledge of Buddhist philosophy with the spiritual language of African American blues and gospel tradition. For over twenty years, Altig has engaged in intensive study of the Lotus Sutra, culminating in this project to make its revolutionary non-dualism accessible to contemporary readers without sacrificing doctrinal depth.\n\nAltig lives in Houston, where he continues his Buddhist practice and serves a growing online community of over 2,500 listeners. He is an independent scholar whose work on “The Blues as Buddhist Epistemology” is currently under academic review. He is available for workshops, readings, and educational presentations on the intersection of spiritual wisdom traditions.", "Answer"),
        
        ("4. Distill this down to a 2-or-3-sentence identifying blurb that could appear on the back cover of your book. The last sentence must be, “S/he lives in _______.”", False),
        ("William Altig is an author and spiritual practitioner who translates Buddhist philosophy into the vernacular of the American Blues tradition. His “Blues Sutras” project has built a dedicated following of over 2,500 listeners, bridging ancient wisdom with contemporary accessibility. He lives in Houston, Texas.", "Answer"),
        
        ("5. Please list highlights of titles, dates, and publishers of other books you have written or edited.", False),
        ("This is Mr. Altig's first book-length publication.", "Answer"),
        
        ("6. Please list highlights of other important published writings—magazine articles, manuals, etc.", False),
        ("• “The Blues as Buddhist Epistemology” (Scholarly article currently under review at Philosophy East and West)\n• “The Dharma Sings the Blues” (Article under review at Contemporary Buddhism)\n• “Changing Poison into Medicine” (published on LessWrong / Alignment Forum, Dec 2025)", "Answer"),
        
        ("7. Please list highlights of any prizes, honors, awards received, titles, and memberships in professional associations.", False),
        ("N/A", "Answer"),
        
        ("8. Please list any important relationships with well-known teachers, centers, practices, etc.", False),
        ("• Affiliation: Independent Scholar-Practitioner.\n• Practice: Over 40 years of dedicated study and practice in the Mahayana tradition, focusing on the Lotus Sutra.", "Answer"),
        
        ("9. Do you have personal, professional, or institutional mailing lists to which it would be appropriate to send notifications and announcements about your book? If so, how many members does it have?", False),
        ("Yes. I maintain a growing digital platform with over 2,500 active listeners and subscribers across ElevenReader and other audio platforms who follow the “Dharma Reborn” series.", "Answer"),
        
        ("10. Please provide the URLs to your website, blog, Facebook page, Instagram profile, Twitter profile, and/or any other significant online profiles that you maintain. Please also list how many followers/subscribers you have on each platform.", False),
        ("• Spotify: https://open.spotify.com/artist/1Rg1HUIwMhqzyytGu2pt1p?si=6ibiDCk6SiickYs9HFCg6A\n• Apple Music: https://music.apple.com/us/artist/william-altig/1744374903\n• SoundCloud: https://soundcloud.com/d-william-altig\n• YouTube: https://www.youtube.com/@dwaltig\n• Instagram: https://www.instagram.com/the.william.altig.choir/\n• Facebook: https://www.facebook.com/profile.php?id=100091943075710\n• iHeartRadio: https://www.iheart.com/artist/william-altig-43055898/\n• TikTok: https://www.tiktok.com/@williamaltig\n• ElevenReader: https://elevenreader.io/authors/1VRKnaoxbsVuR8UV3yuF\n• Blog: https://williamaltig.blog\n\nSubscribers: ~2,500+ active listeners across audio platforms.", "Answer"),
        
        ("ABOUT YOUR BOOK", True),
        ("1. Working title and subtitle:", False),
        ("The Blues Sutras: A Vernacular Translation of the Lotus Sutra", "Answer"),
        
        ("2. A provisional table of contents with brief chapter summaries.", False),
        ("INTRODUCTORY MATERIAL\n• Translator’s Note: The Dharma and the Blues – Explains the “cultural translation” methodology.\n• Scholarly Introduction: Contextualizes the text within Mahayana history and Kumārajīva’s legacy.\n• Reader’s Introduction: A guide to reading the sutra as a spiritual practice.\n\nTHE 28 CHAPTERS\n1. The Opening (Introduction): The great assembly gathers; the Buddha enters samadhi, shaking the cosmos.\n2. The Loving Tricks (Expedient Means): The Buddha reveals the “One Vehicle”—all teachings are skillful means to lead beings to Buddhahood.\n3. The Parable (A Parable): The famous “Burning House” parable; the father saves his children with “toy carts” (skillful means).\n4. Faith and Understanding: The “Prodigal Son” parable, told from the son's perspective.\n5. The Parable of the Medicinal Herbs: The Dharma rain falls equally on all, but plants grow according to their capacity.\n6. The Naming (Bestowal of Prophecies): The Buddha predicts future Buddhahood for his major disciples.\n7. The Phantom City Blues: The parable of a guide creating a “phantom city” to let weary travelers rest before the final journey.\n8. When the Five Hundred Finally Heard Their Names Called: Prophecies extending to lesser-known disciples; the “Gem in the Jacket” parable.\n9. When the Faithful Get Their Due: Prophecies for Ananda and Rahula.\n10. Honor the Word-Carriers: The sanctity of the text itself; the “Digging for Water” parable.\n11. When the Jeweled Tower Rose Up: A massive stupa rises from the earth containing a past Buddha who verifies the teaching.\n12. When Your Enemy Was Your Teacher (Devadatta): The radical revelation that even the Buddha's enemy will become a Buddha, and the 8-year-old Dragon Princess achieves enlightenment instantly.\n13. When the Road Gets Rough: The assembly vows to spread the sutra in evil ages despite persecution.\n14. How to Keep Your Soul Clean: The “Peaceful Practices”—how a Bodhisattva should live and teach.\n15. When the Underground Army Rose Up: Countless Bodhisattvas emerge from the earth, pledging to keep the Dharma.\n16. How Long the Tathāgata’s Been Alive: The revelation of the Buddha's eternal life span; the “Physician and his Children” parable.\n17. Counting Up the Grace: The immense merit of hearing and believing this chapter.\n18. The Joy That Multiplies: The benefit of simply rejoicing in the sutra.\n19. When Your Senses Wake Up: The transformation of the six senses through practice.\n20. The Bodhisattva Who Never Looked Down: The story of Sadaparibhuta, who bowed to everyone saying, “I dare not despise you, for you will become a Buddha.”\n21. When the Lord Showed His Hand: The Buddha displays supernatural powers to entrust the sutra to the assembly.\n22. The Passing On (Entrustment): The final transfer of the Dharma responsibility.\n23. Medicine King Bodhisattva: The story of extreme devotion and the burning of the body as an offering.\n24. Wonderful Sound Bodhisattva: A visitor from another world demonstrates the vastness of the cosmos.\n25. The One Who Hears the World Crying (Avalokiteśvara): The “Universal Gate” chapter; the manifold saving forms of Guanyin/Avalokiteśvara.\n26. The Protection Songs (Dhāraṇī): Magic spells and guardians protecting the sutra's holders.\n27. The Beautiful King and His Good Friends: The story of King Subhavyuha transformed by his two sons.\n28. Universal Worthy Comes to Encourage You: Bodhisattva Samantabhadra vows to protect those who keep the sutra in the final age.\n\nSCHOLARLY APPARATUS\n• Sanskrit/Chinese Term Index: Cross-referencing Blues vernacular terms with original Sanskrit/Chinese.\n• Glossary of Key Buddhist Terms: Definitions of major concepts.", "Answer")
    ]
    
    for text, type_ in content:
        if type_ is True:
            # Section Header
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run("\n" + text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.all_caps = True
        elif type_ == "Answer":
            # Answer text
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
        else:
            # Question text
            p = doc.add_paragraph(text)
            p.runs[0].bold = True

    doc.save('Wisdom_Submission_Questionnaire_COMPLETED.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_wisdom_questionnaire()
