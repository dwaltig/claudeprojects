import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_formatted_paragraph(doc, text):
    """Parses simple Markdown (**bold**, *italic*) and adds to document."""
    p = doc.add_paragraph()
    
    # Split by bold first: **text**
    # This regex captures the delimiter and the content: (\*\*.*?\*\*)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold content
            content = part[2:-2]
            run = p.add_run(content)
            run.bold = True
        else:
            # Inside non-bold parts, look for italics: *text*
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    # Italic content
                    content = subpart[1:-1]
                    run = p.add_run(content)
                    run.italic = True
                else:
                    # Plain text
                    if subpart:
                        p.add_run(subpart)

def add_markdown_content(doc, md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('- '):
            # List item
            p = doc.add_paragraph(style='List Bullet')
            # remove '- ' and format
            text = line[2:]
            # We can't use our helper easily with style arg in add_paragraph 
            # so let's just do a simple run addition for lists or accept plain text
            # For simplicity, let's just clean it:
            clean_text = text.replace('**', '').replace('*', '')
            p.add_run(clean_text)
        elif line.startswith('|'):
             # Table row - skip or simplistic render
             # Tables are hard. Let's just print simple text
             p = doc.add_paragraph(line)
             p.paragraph_format.left_indent = Inches(0.5)
        else:
            add_formatted_paragraph(doc, line)

def add_html_chapter(doc, html_path):
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')
    
    # Extract Title (h1)
    h1 = soup.find('h1')
    if h1:
        doc.add_heading(h1.get_text(), level=1)
    
    h2 = soup.find('h2')
    if h2:
        doc.add_heading(h2.get_text(), level=2)
        
    body = soup.find('body')
    if not body:
        return

    for element in body.find_all(['p', 'h3', 'div']):
        if element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            if 'verse' in element.parent.get('class', []):
                p = doc.add_paragraph(element.get_text())
                p.paragraph_format.left_indent = Inches(0.5)
                p.italic = True
            else:
                # Basic text, strip any HTML tags just to be safe
                doc.add_paragraph(element.get_text())
        elif element.name == 'div' and 'verse' in element.get('class', []):
            pass 

def create_sample_chapters():
    doc = Document()
    
    # Title Page
    doc.add_heading('THE BLUES SUTRAS', 0)
    p = doc.add_paragraph('A Vernacular Translation of the Lotus Sutra')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph('William Altig')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n\n')
    
    p = doc.add_paragraph('SAMPLE CHAPTERS')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Introduction
    print("Processing Introduction...")
    add_markdown_content(doc, '/Users/williamaltig/claudeprojects/Lotus_Sutra/HYBRID_EDITION_INTRODUCTION.md')
    doc.add_page_break()
    
    # Chapter 3
    print("Processing Chapter 3...")
    add_html_chapter(doc, '/Users/williamaltig/claudeprojects/Lotus_Sutra/01_BLUES_INTERPRETATION/lotus_sutra_chapters_html/chapter_03.html')
    doc.add_page_break()
    
    # Chapter 12
    print("Processing Chapter 12...")
    add_html_chapter(doc, '/Users/williamaltig/claudeprojects/Lotus_Sutra/01_BLUES_INTERPRETATION/lotus_sutra_chapters_html/chapter_12.html')
    
    output_path = '/Users/williamaltig/claudeprojects/Lotus_Sutra/Sample_Chapters_Wisdom_Parallax.docx'
    doc.save(output_path)
    print(f"Created {output_path}")

if __name__ == "__main__":
    create_sample_chapters()
