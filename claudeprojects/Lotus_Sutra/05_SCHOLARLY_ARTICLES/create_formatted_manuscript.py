#!/usr/bin/env python3
"""
Create professionally formatted Word document from manuscript with markdown italics.
Converts *Title* to actual italics and applies academic formatting.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

def create_formatted_document(input_file, output_file):
    """Create a professionally formatted Word document from text file."""

    # Read the manuscript
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Process each line
    lines = content.split('\n')

    for line in lines:
        # Skip empty lines but preserve them
        if not line.strip():
            doc.add_paragraph()
            continue

        # Determine if this is a title/heading
        is_title = False
        is_section_header = False
        is_abstract = False
        is_keywords = False
        is_notes_header = False
        is_works_cited_header = False
        is_body_paragraph = False

        # Strip leading tab for processing (used to mark body paragraphs in plain text)
        clean_line = line.lstrip('\t')

        if clean_line.strip() == 'The Blues as Buddhist Epistemology':
            is_title = True
        elif clean_line.strip().startswith('How African-American Musical Tradition'):
            is_title = True
        elif clean_line.strip() == 'Abstract':
            is_abstract = True
        elif clean_line.strip().startswith('Keywords:'):
            is_keywords = True
        elif clean_line.strip() in ['Introduction', 'Defining "Shared Epistemology"',
                              'Section 1: Philosophical Parallels--The Shared Epistemology',
                              'Section 2: Historical Evidence and Precedent',
                              'Section 3: Translation as Epistemological Bridge',
                              'Section 4: Implications for Contemporary Practice',
                              'Conclusion']:
            is_section_header = True
        elif clean_line.strip() == 'Notes':
            is_notes_header = True
        elif clean_line.strip() == 'Works Cited':
            is_works_cited_header = True
        else:
            # Regular body paragraph (may have had leading tab)
            is_body_paragraph = True

        # Add paragraph
        p = doc.add_paragraph()

        # Set paragraph format
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # Center titles
        if is_title:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Apply first-line indent to body paragraphs (replaces the tab character)
        if is_body_paragraph and clean_line.strip():
            paragraph_format.first_line_indent = Inches(0.5)

        # Process the text with italics (use clean_line without leading tab)
        process_line_with_italics(p, clean_line, is_title, is_section_header,
                                  is_abstract, is_keywords, is_notes_header,
                                  is_works_cited_header)

    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

def process_line_with_italics(paragraph, text, is_title=False, is_section_header=False,
                              is_abstract=False, is_keywords=False, is_notes_header=False,
                              is_works_cited_header=False):
    """Process a line of text and add runs with appropriate formatting."""

    # Pattern to find markdown italics (supports *text* and _text_)
    pattern = r'[_*]([^*_]+)[_*]'

    # Find all matches
    matches = list(re.finditer(pattern, text))

    if not matches:
        # No italics, just add the text
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return

    # Process text with italics
    last_end = 0
    for match in matches:
        # Add text before the italicized part
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        # Add italicized text
        run = paragraph.add_run(match.group(1))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True

        last_end = match.end()

    # Add any remaining text after last match
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def fix_en_dashes(text):
    """Replace hyphens in page ranges with en-dashes."""
    # Pattern for page ranges like 32-33 or 349-370
    return re.sub(r'(\d+)-(\d+)', r'\1–\2', text)

if __name__ == '__main__':
    input_file = '/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/ARTICLE_4_Blues_Buddhist_Epistemology_BLIND_FOR_SUBMISSION.txt'
    output_file = '/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/The_Blues_as_Buddhist_Epistemology.docx'

    create_formatted_document(input_file, output_file)
    print("Formatting complete!")
    print("- Times New Roman, 12pt font")
    print("- Double spacing throughout")
    print("- 1-inch margins")
    print("- Markdown italics converted to actual italics")
    print("- Ready for academic journal submission")
