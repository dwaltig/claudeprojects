
import re
from docx import Document
from docx.shared import Pt, Inches

def generate_tts_docx():
    input_file = "/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/TIANTAI_TRANSFORMER_COMPARATIVE_ANALYSIS.md"
    output_docx = "/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/TIANTAI_TRANSFORMER_COMPARATIVE_ANALYSIS_TTS.docx"

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # --- CLEANING FOR TTS ---

    # 1. Remove Tables
    # Identify table blocks and remove them or replace with note
    # Simple regex to remove lines starting with |
    content = re.sub(r'\|.*\|', '', content)

    # 2. Remove Citations like [1], [1, 2]
    content = re.sub(r'\[\d+(?:, \d+)*\]', '', content)

    # 3. Remove References Section
    if "## References" in content:
        content = content.split("## References")[0]

    # 4. Expand Abbreviations
    content = content.replace("e.g.", "for example")
    content = content.replace("i.e.", "that is")
    content = content.replace("vs.", "versus")

    # 5. Remove Math Syntax slightly for readability
    # $W_Q$ -> W_Q
    content = content.replace('$', '')
    # \text{Attention} -> Attention
    content = re.sub(r'\\text\{(\w+)\}', r'\1', content)
    # Remove other latex commands like \frac, \left, \right
    content = re.sub(r'\\[a-zA-Z]+', '', content)
    content = content.replace('{', '').replace('}', '')

    # --- CREATE DOCX ---
    doc = Document()
    
    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Handle Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        else:
            # Clean Markdown Bold/Italic for TTS completely
            # Use regex to strip the markers but keep the text
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Bold
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line) # Italic
            
            # Final safety sweep for any remaining asterisks (e.g. unmatched)
            clean_line = clean_line.replace('*', '')
            
            p = doc.add_paragraph(clean_line)
            p.paragraph_format.line_spacing = 1.5

    doc.save(output_docx)
    print(f"Created TTS Optimized DOCX: {output_docx}")

if __name__ == "__main__":
    generate_tts_docx()
