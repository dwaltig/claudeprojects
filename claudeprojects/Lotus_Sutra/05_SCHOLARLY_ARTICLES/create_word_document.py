#!/usr/bin/env python3
"""
Convert The Blues Has Been Teaching Dharma All Along from Markdown to Word (.docx)
WITH WORKING PAGE NUMBERS - Using different method
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re

def add_page_number_to_footer(section):
    """Add page numbers to footer using XML approach"""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Clear existing content
    paragraph.clear()
    
    # Add page number using XML
    run = paragraph.add_run()
    
    # Create the page number field using XML
    fldChar_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> PAGE </w:instrText>')
    fldChar_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_end)
    
    # Format the run
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

def create_formatted_docx():
    # Create document
    doc = Document()
    
    # Set up page margins (1 inch all around - standard)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add page numbers to footer
        add_page_number_to_footer(section)
    
    # Read the markdown file
    with open('The_Blues_Has_Been_Teaching_Dharma_All_Along.txt', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines and horizontal rules
        if not line or line == '---':
            i += 1
            continue
        
        # Main title (# The Blues...)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.name = 'Times New Roman'
            i += 1
            continue
        
        # Subtitle (## Recognizing...)
        if line.startswith('## ') and i < 5:  # Only first few are subtitles
            subtitle = line[3:].strip()
            p = doc.add_heading(subtitle, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
            i += 1
            continue
        
        # Section headers (## Opening, ## The Same Truth...)
        if line.startswith('## '):
            header = line[3:].strip()
            p = doc.add_heading(header, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
                run.bold = True
            i += 1
            continue
        
        # Author byline
        if line.startswith('By '):
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
                run.italic = True
            i += 1
            continue
        
        # Block quotes (lines starting with >)
        if line.startswith('> '):
            quote_text = line[2:].strip()
            p = doc.add_paragraph(quote_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
                run.italic = True
            i += 1
            continue
        
        # Regular paragraphs
        if line:
            # Process inline formatting (italics)
            p = doc.add_paragraph()
            
            # Split by asterisks to handle italics
            parts = re.split(r'(\*[^*]+\*)', line)
            
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    # Italic text
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    # Regular text
                    run = p.add_run(part)
                
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
            
            # Set paragraph spacing
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
        
        i += 1
    
    # Save the document
    output_file = 'The_Blues_Has_Been_Teaching_Dharma_All_Along.docx'
    doc.save(output_file)
    print(f"✅ Created: {output_file}")
    print(f"   - Proper headers and formatting")
    print(f"   - Italics preserved")
    print(f"   - Times New Roman, 12pt")
    print(f"   - PAGE NUMBERS in footer (centered)")
    print(f"   - Ready for submission")
    print(f"\n⚠️  NOTE: You may need to update fields in Word:")
    print(f"   - Open document in Word")
    print(f"   - Press Ctrl+A (select all) or Cmd+A (Mac)")
    print(f"   - Press F9 to update all fields")
    print(f"   - Or right-click footer and select 'Update Field'")

if __name__ == '__main__':
    create_formatted_docx()
