#!/usr/bin/env python3
"""
Convert Diamond Sutra Tricycle article from Markdown to Word format.
Applies academic formatting: 12pt Times New Roman, double-spaced, 1-inch margins.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re

def create_tricycle_article():
    """Create professionally formatted Word document from markdown."""
    
    input_file = '/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/ARTICLE_5_DIAMOND_SUTRA_DRAFT.md'
    output_file = '/Users/williamaltig/claudeprojects/Lotus_Sutra/05_SCHOLARLY_ARTICLES/ARTICLE_5_Diamond_Sutra_Blues_TRICYCLE.docx'
    
    # Read the manuscript
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Process each line
    lines = content.split('\n')
    in_blockquote = False
    blockquote_text = []
    
    for line in lines:
        stripped = line.strip()
        
        # Skip markdown horizontal rules
        if stripped == '---':
            continue
        
        # Handle blockquotes
        if stripped.startswith('>'):
            in_blockquote = True
            quote_text = stripped.lstrip('> ').strip()
            if quote_text:
                blockquote_text.append(quote_text)
            continue
        elif in_blockquote and stripped == '':
            # End of blockquote - write it out
            if blockquote_text:
                full_quote = ' '.join(blockquote_text)
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
                add_formatted_text(p, full_quote)
                blockquote_text = []
            in_blockquote = False
            doc.add_paragraph()
            continue
        elif in_blockquote:
            blockquote_text.append(stripped)
            continue
        
        # Skip empty lines but preserve them
        if not stripped:
            doc.add_paragraph()
            continue
        
        # Handle headers
        if stripped.startswith('# '):
            # Main title
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            run = p.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
            continue
        
        if stripped.startswith('## '):
            # Section header
            header_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(header_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            continue
        
        # Handle byline (italic)
        if stripped.startswith('*By ') or stripped.startswith('*William'):
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            # Remove asterisks
            byline = stripped.strip('*')
            run = p.add_run(byline)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
            continue
        
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        
        add_formatted_text(p, stripped)
    
    # Flush any remaining blockquote
    if blockquote_text:
        full_quote = ' '.join(blockquote_text)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.right_indent = Inches(0.5)
        add_formatted_text(p, full_quote)
    
    # Save document
    doc.save(output_file)
    print(f"Document saved to: {output_file}")
    return output_file

def add_formatted_text(paragraph, text):
    """Add text with markdown italics converted to Word italics."""
    
    # Pattern for markdown italics *text*
    pattern = r'\*([^*]+)\*'
    
    matches = list(re.finditer(pattern, text))
    
    if not matches:
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return
    
    last_end = 0
    for match in matches:
        # Text before italics
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        
        # Italicized text
        run = paragraph.add_run(match.group(1))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        
        last_end = match.end()
    
    # Remaining text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

if __name__ == '__main__':
    output = create_tricycle_article()
    print("Formatting complete!")
    print("- Times New Roman, 12pt font (14pt for title)")
    print("- Double spacing throughout")
    print("- 1-inch margins")
    print("- Blockquotes indented")
    print("- Markdown italics converted to Word italics")
    print("- Ready for Tricycle submission")
