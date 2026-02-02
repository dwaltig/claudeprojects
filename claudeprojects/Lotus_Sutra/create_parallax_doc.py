from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_parallax_proposal():
    doc = Document()
    
    # Title
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    heading = doc.add_paragraph('PROPOSAL: THE BLUES SUTRAS')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)
    
    sub = doc.add_paragraph('William Altig')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph('Houston, Texas | 832-212-0915 | dwaltig@gmail.com')
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('_' * 80)

    # Content
    sections = [
        ("PROJECT OVERVIEW", True),
        ("Title: The Blues Sutras: A Vernacular Translation of the Lotus Sutra\nGenre: Buddhist Translation / Spiritual Nonfiction\nEstimated Word Count: 180,000 words\nStatus: Complete manuscript available", False),
        
        ("1. SYNOPSIS", True),
        ("“No mud, no lotus.” This profound teaching from Thich Nhat Hanh is the heartbeat of the Blues.", False),
        ("The Blues Sutras is a revolutionary vernacular translation of the Lotus Sutra, Mahayana Buddhism’s most influential text. While scholarly translations often lock the Dharma behind academic terminology, this project renders the text into the spiritual vernacular of the African-American Blues tradition—America’s indigenous wisdom language. Just as the Buddha used the vernacular Magadhi to reach the masses, this translation uses the Blues idiom to make the profound non-dualism of the Lotus Sutra visceral, accessible, and deeply moving.", False),
        ("The project is not a summary or a commentary, but a rigorous, full-length translation of all 28 chapters. It preserves the text’s theological depth and structural integrity while amplifying its affective power. By translating the ancient “Skillful Means” (upaya) of the Buddha into the “Skillful Means” of the Blues, the text demonstrates that the transformation of suffering (dukkha) into liberation is a universal human technology, shared by ancient Indian sages and Mississippi Delta musicians alike.", False),
        ("At its core, this is a work of Engaged Buddhism. It refuses to leave the Dharma in the ivory tower. It brings the Sutra down to the street corner, the prison, and the hospital room, speaking directly to those who know suffering intimately and offering them the revolutionary promise that they, too, are Buddhas.", False),

        ("2. A NOTE ON METHODOLOGY & RESPECT", True),
        ("In the current publishing climate, a project involving vernacular translation requires clear ethical grounding. This work approaches the Blues not as a “style” to be mimicked or parodied, but as a sacred wisdom tradition equivalent in dignity to the Buddhist canon.", False),
        ("The translation is an act of homage that places African-American spiritual technology—specifically its method of transforming suffering into liberation—into dialogue with Mahayana philosophy. It avoids caricature by adhering strictly to the doctrinal content of the Sutra, simply rendering it through the grammar of American resilience. It is a work of “cultural translation” that honors the prophetic voice of the oppressed as a vehicle for the Dharma.", False),

        ("3. MARKET ANALYSIS & AUDIENCE", True),
        ("Primary Audience:", False),
        ("• Engaged Buddhists: readers of Thich Nhat Hanh, Pema Chödrön, and angel Kyodo williams who seek a Dharma that engages directly with the grit of real life.\n• Secular/Cultural Seekers: The “Spiritual But Not Religious” demographic who find traditional liturgy alienating but resonate with the raw honesty of the Blues/Gospel tradition.\n• African-American Practitioners: Providing a Dharma gate that honors Black cultural forms as valid vessels for enlightenment.", False),
        ("Comparable Titles:", False),
        ("• The Gospel According to Jesus, Queen of Heaven (Jo Clifford) – A radical re-visioning of scripture that honors the marginalized.\n• After Buddhism (Stephen Batchelor) – Reimagining Dharma for the secular/modern age.\n• Radical Dharma (Rev. angel Kyodo williams) – Exploring the intersection of blackness, liberation, and Buddhist practice.", False),
        ("Why Parallax?", False),
        ("Parallax Press is the natural home for this work because the Blues is the practice of “transforming garbage into flowers.” No other publisher has so consistently championed the idea that suffering is the compost for enlightenment. This translation is a musical elaboration of that central Parallax truth.", False),

        ("4. ABOUT THE AUTHOR", True),
        ("William Altig is a scholar-practitioner and cultural translator with over 40 years of experience in Mahayana Buddhism. He is an Independent Scholar whose work focuses on the intersection of American vernacular music and Buddhist epistemology.", False),
        ("His current projects include “The Blues as Buddhist Epistemology” (scholarly article under review at Philosophy East and West) and “The Dharma Sings the Blues” (under review at Contemporary Buddhism). He maintains a growing digital platform with over 2,500 active listeners across audio platforms like ElevenReader, where his “Dharma Reborn” series has built a dedicated following.", False),
        ("He lives in Houston, Texas.", False),
        ("Online Platform:\n• ElevenReader: https://elevenreader.io/authors/1VRKnaoxbsVuR8UV3yuF\n• Blog: https://williamaltig.blog", False)
    ]
    
    for text, is_header in sections:
        if is_header:
            p = doc.add_paragraph()
            run = p.add_run("\n" + text)
            run.bold = True
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(12)

    doc.save('Proposal_Parallax_The_Blues_Sutras_Altig.docx')
    print("Proposal document created successfully.")

if __name__ == "__main__":
    create_parallax_proposal()
