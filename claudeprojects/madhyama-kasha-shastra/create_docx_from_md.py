import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    doc = Document()
    # Simple conversion: treat headings, bullet points, tables, code blocks
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    in_code_block = False
    for line in lines:
        stripped = line.rstrip('\n')
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            # add as preformatted text using default style
            doc.add_paragraph(stripped)
            continue
        if stripped.startswith('#'):
            level = stripped.count('#')
            text = stripped.lstrip('#').strip()
            doc.add_heading(text, level=level)
        elif stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('|'):
            # simplistic table handling: ignore for now
            continue
        elif stripped.strip() == '':
            doc.add_paragraph('')
        else:
            doc.add_paragraph(stripped)
    doc.save(docx_path)

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print('Usage: python create_docx_from_md.py <input_md> <output_docx>')
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    md_to_docx(md_path, docx_path)
