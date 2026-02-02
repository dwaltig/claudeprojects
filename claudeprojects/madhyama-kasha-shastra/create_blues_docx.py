#!/usr/bin/env python3
"""
Create a properly formatted DOCX from the Blues Lyrics markdown.
Preserves verse structure with proper line breaks and italics.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_rule(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('───────────────────────────')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

def process_text_with_formatting(paragraph, text):
    """Process text with *italic* and **bold** formatting."""
    # Pattern for bold and italic
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
        elif part.strip():
            run = paragraph.add_run(part)
            run.font.name = 'Georgia'
            run.font.size = Pt(11)

def create_blues_docx(md_path: Path, docx_path: Path):
    doc = Document()
    
    # Set up document margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse YAML frontmatter
    title = "Mūlamadhyamakakārikā Blues"
    subtitle = "The Lyrics"
    author = "Master John Kim"
    
    frontmatter_match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
    if frontmatter_match:
        fm = frontmatter_match.group(1)
        for line in fm.split('\n'):
            if line.startswith('title:'):
                title = line.split(':', 1)[1].strip()
            elif line.startswith('subtitle:'):
                subtitle = line.split(':', 1)[1].strip()
            elif line.startswith('author:'):
                author = line.split(':', 1)[1].strip()
        content = content[frontmatter_match.end():]
    
    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.name = 'Georgia'
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.name = 'Georgia'
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run(f"Interpreted by {author}")
    author_run.font.size = Pt(14)
    author_run.font.name = 'Georgia'
    
    # Add page break after title
    doc.add_page_break()
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    verse_buffer = []  # Collect verse lines
    
    def flush_verse_buffer():
        """Output collected verse lines as a single paragraph with line breaks."""
        nonlocal verse_buffer
        if verse_buffer:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            for idx, verse_line in enumerate(verse_buffer):
                # Check if entire verse block is italic
                if verse_line.startswith('*') and verse_line.endswith('*'):
                    run = p.add_run(verse_line[1:-1])
                    run.italic = True
                else:
                    process_text_with_formatting(p, verse_line)
                run = p.runs[-1] if p.runs else p.add_run('')
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                # Add line break if not last line
                if idx < len(verse_buffer) - 1:
                    p.add_run('\n')
            verse_buffer = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines (but flush verse buffer first)
        if not line.strip():
            flush_verse_buffer()
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() in ['* * *', '---', '***']:
            flush_verse_buffer()
            add_horizontal_rule(doc)
            i += 1
            continue
        
        # Handle chapter headings (# Chapter...)
        if line.startswith('# ') and not line.startswith('## ') and not line.startswith('### '):
            flush_verse_buffer()
            heading_text = line[2:].strip()
            doc.add_page_break()  # New page for each chapter
            h = doc.add_heading(heading_text, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(20)
            i += 1
            continue
        
        # Handle malformed headings like "### ## *When..."
        if re.match(r'^#+\s*#+', line):
            flush_verse_buffer()
            heading_text = re.sub(r'^#+\s*#+\s*', '', line)
            heading_text = heading_text.strip('*').strip()
            h = doc.add_heading(heading_text, level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(14)
                run.italic = True
            i += 1
            continue
        
        # Handle subheadings (## or ### ...)
        if line.startswith('## ') or line.startswith('### '):
            flush_verse_buffer()
            level = 2 if line.startswith('## ') else 3
            heading_text = line.lstrip('#').strip()
            heading_text = heading_text.strip('*').strip()
            h = doc.add_heading(heading_text, level=level)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Georgia'
                run.font.size = Pt(14) if level == 2 else Pt(12)
                run.italic = True
            i += 1
            continue
        
        # Handle verse lines (starting with |)
        if line.startswith('|'):
            verse_text = line[1:].strip()
            if verse_text:  # Skip empty pipe lines
                verse_buffer.append(verse_text)
            i += 1
            continue
        
        # Handle bold section markers (**The Teaching:** etc.)
        if line.startswith('**') and ':**' in line:
            flush_verse_buffer()
            marker_text = line.replace('**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(marker_text)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Handle verse markers like **(Verse 1)**
        if line.startswith('**(') and line.endswith(')**'):
            flush_verse_buffer()
            verse_marker = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(verse_marker)
            run.bold = True
            run.font.name = 'Georgia'
            run.font.size = Pt(10)
            i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.', line):
            flush_verse_buffer()
            list_text = re.sub(r'^\d+\.\s*', '', line)
            list_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', list_text)
            list_text = re.sub(r'\*([^*]+)\*', r'\1', list_text)
            p = doc.add_paragraph(style='List Number')
            process_text_with_formatting(p, list_text)
            i += 1
            continue
        
        # Regular paragraph
        flush_verse_buffer()
        p = doc.add_paragraph()
        process_text_with_formatting(p, line)
        
        i += 1
    
    # Flush any remaining verses
    flush_verse_buffer()
    
    # Save the document
    doc.save(docx_path)
    print(f"Created: {docx_path}")

if __name__ == '__main__':
    base = Path('/Users/williamaltig/claudeprojects/madhyama-kasha-shastra/01_TRANSLATIONS')
    md_path = base / 'Mūlamadhyamakakārikā_Blues_Lyrics_Only.md'
    docx_path = base / 'Mūlamadhyamakakārikā_Blues_Lyrics_Only.docx'
    
    create_blues_docx(md_path, docx_path)
