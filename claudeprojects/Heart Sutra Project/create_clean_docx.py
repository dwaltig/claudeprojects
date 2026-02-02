#!/usr/bin/env python3
"""Create clean DOCX from VERSION_D markdown."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def process_markdown_text(text):
    """Convert markdown formatting to plain text, preserving content."""
    # Remove bold/italic markers
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # bold+italic
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)      # bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)          # italic
    text = re.sub(r'_([^_]+)_', r'\1', text)          # underline/italic
    
    # Remove code markers
    text = re.sub(r'`([^`]+)`', r'\1', text)
    
    # Clean up footnote references
    text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
    
    return text

def create_clean_docx():
    """Create properly formatted DOCX from VERSION_D."""
    
    with open('/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    lines = content.split('\n')
    skip_next = False
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        if skip_next:
            skip_next = False
            continue
            
        # Skip empty lines
        if not line:
            continue
        
        # Skip horizontal rules
        if line.startswith('---'):
            continue
        
        # Title
        if line.startswith('# '):
            title_text = process_markdown_text(line[2:])
            p = doc.add_paragraph(title_text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            doc.add_paragraph()  # blank line
        
        # Author
        elif '**William Altig**' in line:
            p = doc.add_paragraph('William Altig')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            doc.add_paragraph()  # blank line
        
        # Section headings (##)
        elif line.startswith('## '):
            heading = process_markdown_text(line[3:])
            # Remove section numbers like "1. ", "2. " etc
            heading = re.sub(r'^\d+\.\s+', '', heading)
            p = doc.add_heading(heading, level=1)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.bold = True
        
        # Subsection headings (###)
        elif line.startswith('### '):
            subheading = process_markdown_text(line[4:])
            p = doc.add_heading(subheading, level=2)
            run = p.runs[0]
            run.font.size = Pt(12)
        
        # Footnotes
        elif line.startswith('[^'):
            match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
            if match:
                num = match.group(1)
                text = process_markdown_text(match.group(2))
                p = doc.add_paragraph(f'[{num}] {text}')
                run = p.runs[0]
                run.font.size = Pt(10)
        
        # Regular paragraphs
        else:
            clean_text = process_markdown_text(line)
            if clean_text:  # Only add non-empty paragraphs
                doc.add_paragraph(clean_text)
    
    output = '/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D.docx'
    doc.save(output)
    print(f"✅ Created clean DOCX: {output}")

if __name__ == '__main__':
    create_clean_docx()
