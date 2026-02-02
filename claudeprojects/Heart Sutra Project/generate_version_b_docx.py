#!/usr/bin/env python3
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def create_submission_docx(input_file, output_file, anonymize=False):
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found.")
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Set default style to Times New Roman 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process metadata and anonymization
    lines = content.split('\n')
    processed_lines = []
    
    in_abstract = False
    
    for line in lines:
        # Anonymization logic: remove author name if blind
        if anonymize:
            line = re.sub(r'William Altig', '[Author Name Redacted for Peer Review]', line, flags=re.IGNORECASE)
        
        # Handle simple markdown formatting
        # Bold
        line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
        # Italics (rough handling)
        # Note: python-docx handles bold/italic via 'runs', so we'll do basic replacement for now 
        # or just strip markdown for the generator
        
        processed_lines.append(line)

    for line in processed_lines:
        line_stripped = line.strip()
        
        if not line_stripped:
            doc.add_paragraph()
            continue
            
        if line_stripped == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("* * *")
            continue
            
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
        elif line.startswith("[^"):
            # Put footnotes at the end in a real doc, but for now just add as text
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10)
        else:
            # Basic italic handling for *text*
            p = doc.add_paragraph()
            parts = re.split(r'(\*[^*]+\*)', line)
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    p.add_run(part)

    doc.save(output_file)
    print(f"Created: {output_file}")

if __name__ == "__main__":
    # Generate Author Version B
    create_submission_docx(
        'ARTICLE_Hellhounds_VERSION_B.md', 
        'Beyond_Appropriation_VERSION_B_AUTHOR.docx', 
        anonymize=False
    )
    # Generate Blind Version B
    create_submission_docx(
        'ARTICLE_Hellhounds_VERSION_B.md', 
        'Beyond_Appropriation_VERSION_B_BLIND.docx', 
        anonymize=True
    )
    # Generate Response Letter
    create_submission_docx(
        'JAR_RESPONSE_LETTER_VERSION_B.md', 
        'JAR_RESPONSE_LETTER_VERSION_B.docx', 
        anonymize=False
    )
