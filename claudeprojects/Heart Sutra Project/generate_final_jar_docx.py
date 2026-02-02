
#!/usr/bin/env python3
"""
Generate final JAR submission DOCX files from markdown.
Creates both a standard and a BLIND version.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import re
import os

def create_formatted_document(input_file, output_file, blind=False):
    """Create a professionally formatted Word document from text file."""

    # Read the manuscript
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document properties
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Process each line
    lines = content.split('\n')

    for line in lines:
        # Skip empty lines but consider them for spacing if needed (though usually we handle spacing via paragraph format)
        # For this specific manuscript, we want to preserve paragraph breaks
        if not line.strip():
            # If it's just an empty line in MD, we might want a blank line in DOCX or just separate paras.
            # Standard academic format uses indent for new para, no extra space.
            # But let's follow the previous script's logic which adds a paragraph
            if line == '': # Strictly empty, not just whitespace
                doc.add_paragraph() 
            continue

        clean_line = line.strip()
        
        # BLINDING LOGIC
        if blind:
            if "William Altig" in clean_line:
                continue
            if "Independent Scholar" in clean_line: # Common affiliation line
                continue

        # Skip markdown formatting artifacts if they are standalone
        if clean_line.startswith('---') or clean_line.startswith('==='):
            continue

        # Determine structural role
        is_title = False
        is_centered = False
        is_abstract = False
        is_keywords = False
        
        # Heuristics for Titles/Headers
        if clean_line.startswith('# '):
            clean_line = clean_line.replace('# ', '').strip()
            is_title = True
        elif clean_line.startswith('## '):
            clean_line = clean_line.replace('## ', '').strip()
            is_centered = True # Section headers often left or centered. Previous script favored identifying specific headers.
        
        # Specific Content Checks
        if "Hellhounds and the Void" in clean_line and len(clean_line) < 100:
            is_title = True
        
        if clean_line == "Abstract":
            is_centered = True
            is_abstract = True
            
        if clean_line.startswith("Keywords:"):
            is_keywords = True

        # Create Paragraph
        p = doc.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        if is_title:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif is_centered:
             # Check if it's a section header to align left or center? 
             # Previous script centered headers like 'Introduction'.
             # Let's align left for standard subheads unless it's Abstract/Title
             if clean_line in ["Abstract", "Introduction", "Conclusion", "Notes", "Works Cited"]:
                 paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
             else:
                 paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            # Body text - Indent
            # But don't indent if it's the start of Abstract or Keywords?
            # Standard: First line indent 0.5 inch
            paragraph_format.first_line_indent = Inches(0.5)

        # Process Italics (Markdown *text* or _text_)
        process_line_with_italics(p, clean_line)

    doc.save(output_file)
    print(f"Saved: {output_file}")


def process_line_with_italics(paragraph, text):
    """Process a line of text and add runs with appropriate formatting."""
    # Simple regex for *italic* or _italic_
    # Note: This handles single * or _. It doesn't handle **bold**. 
    # The user's synthesis used *Void* and *śūnyatā* (italics).
    # The file has **William Altig** (bold). I should handle bold too if possible, but minimal scope is key.
    
    # Let's use a tokenizer approach for * text *
    # Current script only needs to handle the *words* I added and standard italics.
    
    # Regex matches: (**bold** OR *italic* OR _italic_)
    # Order matters: check double chars first
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)', text)
    
    for part in parts:
        # Handle empty strings from split
        if not part:
            continue
            
        # Check BOLD (**text**)
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
        # Check ITALIC (*text*)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
        # Check ITALIC (_text_)
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            
        # Plain text
        else:
             run = paragraph.add_run(part)
             run.font.name = 'Times New Roman'
             run.font.size = Pt(12)

if __name__ == "__main__":
    input_md = "/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_JAR_SUBMISSION_FINAL.md"
    
    # 1. Non-Blind
    output_std = "/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_JAR_SUBMISSION_FINAL.docx"
    create_formatted_document(input_md, output_std, blind=False)
    
    # 2. Blind
    output_blind = "/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_JAR_SUBMISSION_FINAL_BLIND.docx"
    create_formatted_document(input_md, output_blind, blind=True)
