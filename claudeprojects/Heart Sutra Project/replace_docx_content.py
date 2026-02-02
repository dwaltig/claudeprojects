#!/usr/bin/env python3
"""Replace content in existing DOCX template."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def process_markdown_text(text):
    """Clean markdown formatting."""
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
    return text

# Open existing template
doc = Document('/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D_READABLE.docx')

# Clear all existing content
for element in doc.element.body:
    doc.element.body.remove(element)

# Read VERSION_D markdown
with open('/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')

for line in lines:
    line = line.strip()
    
    if not line or line.startswith('---'):
        continue
    
    # Title
    if line.startswith('# '):
        text = process_markdown_text(line[2:])
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)
    
    # Author
    elif '**William Altig**' in line:
        p = doc.add_paragraph('William Altig')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        doc.add_paragraph()
    
    # Headings
    elif line.startswith('## '):
        text = re.sub(r'^\d+\.\s+', '', process_markdown_text(line[3:]))
        p = doc.add_heading(text, level=1)
        p.runs[0].font.size = Pt(12)
    
    elif line.startswith('### '):
        p = doc.add_heading(process_markdown_text(line[4:]), level=2)
        p.runs[0].font.size = Pt(12)
    
    # Footnotes
    elif line.startswith('[^'):
        match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
        if match:
            num, text = match.group(1), process_markdown_text(match.group(2))
            p = doc.add_paragraph(f'[{num}] {text}')
            p.runs[0].font.size = Pt(10)
    
    # Regular text
    else:
        clean = process_markdown_text(line)
        if clean:
            doc.add_paragraph(clean)

doc.save('/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D_READABLE.docx')
print("✅ Created readable DOCX using your existing file as template")
