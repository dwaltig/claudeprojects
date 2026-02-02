import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def markdown_to_docx(input_md, output_docx, blind=False):
    doc = Document()
    
    # Set default font to Times New Roman 12pt, double-spaced
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    
    if not os.path.exists(input_md):
        print(f"Error: {input_md} not found.")
        return

    with open(input_md, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Anonymization for Blind version
    if blind:
        content = re.sub(r'William Altig', '[REDACTED FOR PEER REVIEW]', content)
        content = re.sub(r'I approach this as a white practitioner and independent scholar\.', 
                        'The author approaches this as a white practitioner and independent scholar.', content)

    # Split into main content and endnotes section
    notes_match = re.search(r'^## Notes', content, re.MULTILINE)
    if notes_match:
        main_content = content[:notes_match.start()]
        notes_section = content[notes_match.start():]
    else:
        main_content = content
        notes_section = ""
    
    # Parse endnote definitions
    endnotes = {}
    word_count_line = ""
    if notes_section:
        wc_match = re.search(r'^\*\*Word Count:\*\*.*', notes_section, re.MULTILINE)
        if wc_match:
            word_count_line = wc_match.group(0)

        endnote_pattern = re.compile(r'^\[\^(\d+)\]:\s*(.*?)(?=\n\[\^|\Z)', re.MULTILINE | re.DOTALL)
        for match in endnote_pattern.finditer(notes_section):
            num = match.group(1)
            text = match.group(2).strip()
            text = re.sub(r'\n\s{2,}', ' ', text)
            endnotes[num] = text
    
    lines = main_content.split('\n')
    
    # Group lines into logical units
    units = []
    current_type = None
    current_buffer = []

    def flush():
        nonlocal current_type, current_buffer
        if current_buffer:
            if current_type == 'para':
                units.append(('para', ' '.join(current_buffer)))
            elif current_type == 'quote':
                # Join quote lines and strip the '> ' from the start of each
                text = ' '.join(b[2:] if b.startswith('> ') else b for b in current_buffer)
                units.append(('quote', text))
            elif current_type == 'list':
                for item in current_buffer:
                    units.append(('list', item[2:]))
            elif current_type == 'heading':
                units.append(('heading', current_buffer[0]))
            elif current_type == 'hr':
                units.append(('hr', '---'))
            current_buffer = []
        current_type = None

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush()
            continue

        if line.startswith('#'):
            flush()
            current_type = 'heading'
            current_buffer.append(line)
            flush()
        elif line.startswith('> '):
            if current_type != 'quote':
                flush()
                current_type = 'quote'
            current_buffer.append(line)
        elif line.startswith('* ') or line.startswith('- '):
            if current_type != 'list':
                flush()
                current_type = 'list'
            current_buffer.append(line)
        elif stripped == '---' or stripped == '------------------------------------------------------------------------':
            flush()
            units.append(('hr', '---'))
        else:
            if current_type != 'para':
                flush()
                current_type = 'para'
            current_buffer.append(line)
    
    flush()
    
    for u_type, text in units:
        if u_type == 'heading':
            level = len(re.match(r'^#+', text).group())
            clean_text = text.lstrip('#').strip()
            p = doc.add_heading('', level=level)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _process_inline_formatting_with_endnotes(p, clean_text)
        elif u_type == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            _process_inline_formatting_with_endnotes(p, text)
        elif u_type == 'list':
            p = doc.add_paragraph(style='List Bullet')
            _process_inline_formatting_with_endnotes(p, text)
        elif u_type == 'hr':
            doc.add_page_break()
        elif u_type == 'para':
            p = doc.add_paragraph()
            _process_inline_formatting_with_endnotes(p, text)

    # Add Endnotes section
    if endnotes:
        doc.add_page_break()
        doc.add_heading('Notes', level=1)
        for num in sorted(endnotes.keys(), key=int):
            p = doc.add_paragraph()
            run = p.add_run(num + ". ")
            run.font.superscript = True
            _process_inline_formatting(p, endnotes[num])
            
    if word_count_line:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _process_inline_formatting(p, word_count_line)

    doc.save(output_docx)
    print(f"Generated: {output_docx}")

def _process_inline_formatting_with_endnotes(paragraph, text):
    parts = re.split(r'(\[\^\d+\])', text)
    for part in parts:
        fn_match = re.match(r'\[\^(\d+)\]', part)
        if fn_match:
            run = paragraph.add_run(fn_match.group(1))
            run.font.superscript = True
        else:
            _process_inline_formatting(paragraph, part)

def _process_inline_formatting(paragraph, text):
    # Split by bold-italic, bold, and italic markers
    pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    v_prefix = "VERSION_D"
    input_file = f"ARTICLE_Hellhounds_{v_prefix}.md"
    response_file = f"JAR_RESPONSE_LETTER_{v_prefix}.md"
    
    markdown_to_docx(input_file, f"ARTICLE_Hellhounds_{v_prefix}_AUTHOR.docx", blind=False)
    markdown_to_docx(input_file, f"ARTICLE_Hellhounds_{v_prefix}_BLIND.docx", blind=True)
    
    if os.path.exists(response_file):
         markdown_to_docx(response_file, f"JAR_RESPONSE_LETTER_{v_prefix}.docx", blind=False)
    else:
        # Check for non-prefixed response letter if needed
        alt_response = f"JAR_RESPONSE_LETTER_{v_prefix}.md"
        if os.path.exists(alt_response):
            markdown_to_docx(alt_response, f"JAR_RESPONSE_LETTER_{v_prefix}.docx", blind=False)
