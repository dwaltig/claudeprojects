#!/usr/bin/env python3
"""Convert Hellhounds v3.2 markdown to DOCX for JoAR submission."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def create_v3_2_docx():
    """Create DOCX from v3.2 markdown."""
    
    input_path = '/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_v3.2.md'
    output_path = '/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_v3.2_SUBMISSION.docx'

    # Read the markdown file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (first line starting with #)
        if line.startswith('# '):
            title = line[2:]
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        
        # Author line
        elif line.startswith('**William Altig**'):
            p = doc.add_paragraph('William Altig')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            
        # Abstract section handling
        elif line.startswith('## Abstract'):
            p = doc.add_paragraph('Abstract')
            p.runs[0].bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Keywords
        elif line.startswith('## **Keywords:**'):
             # Extract keywords text
            keywords_text = line.replace('## **Keywords:**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run('Keywords: ')
            run.bold = True
            p.add_run(keywords_text)

        # Horizontal rules - skip
        elif line.startswith('---'):
            pass
            
        # Section headings (##)
        elif line.startswith('## '):
            # Clean up heading text if it has numbering like "## 1. Intro"
            # JoAR follows Chicago style, usually just Title Case or centered.
            # We'll keep the numbers if present in MD, but ensure formatting is heading-like
            heading_text = line[3:].strip()
            # Optional: Remove leading "1. " if specific style requires, but keeping for clarity
            p = doc.add_heading(heading_text, level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = None # Default black
        
        # Subsection headings (###)
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_heading(heading_text, level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
                run.font.color.rgb = None
        
        # Footnotes entries (at the end)
        elif line.startswith('[^'):
            # Extract footnote number and text
            match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
            if match:
                num = match.group(1)
                text = match.group(2)
                
                # Format styling for footnotes
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) # Simple unbold for docx plain
                text = re.sub(r'\*(.*?)\*', r'\1', text)     # Simple unitalic
                
                p = doc.add_paragraph(f'{num}. {text}')
                p.style = doc.styles['No Spacing']
                p.runs[0].font.size = Pt(10)
        
        # Regular paragraphs
        else:
            # Handle standard paragraphs
            p = doc.add_paragraph()
            
            # Very basic markdown parsing for bold/italic within text
            # This is a simple parser; robust parsing would require a library or more complex logic
            # For this script we will strip MD markers or rely on a regex replacer for common cases.
            # Ideally we want to KEEP italics/bold.
            
            # Split by markdown markers and add runs? That's complex.
            # Simplified approach: Use regex to replace MD with nothing for clean text, 
            # OR ideally, simple formatting.
            
            # Let's try to preserve italics at least, as book titles need them.
            # Regex to find *text*
            
            parts = re.split(r'(\*[^\*]+\*|\*\*[^\*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    if part: # If not empty
                        # Clean up footnote refs [^1] -> superscript numbers or just [1]
                        # For docx submission, usually simple text [1] is fine, or true footnotes.
                        # This script mimics the old one which seemingly just kept [1].
                        clean_part = re.sub(r'\[\^(\d+)\]', r'[\1]', part)
                        p.add_run(clean_part)

        i += 1
    
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

if __name__ == '__main__':
    create_v3_2_docx()
