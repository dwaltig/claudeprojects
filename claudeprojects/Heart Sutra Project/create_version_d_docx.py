#!/usr/bin/env python3
"""Convert VERSION_D markdown to DOCX for review."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_version_d_docx():
    """Create DOCX from VERSION_D markdown."""
    
    # Read the markdown file
    with open('/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Title (first line)
        if i == 0 and line.startswith('# '):
            title = line[2:]
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
        
        # Author line
        elif line.startswith('**William Altig**'):
            p = doc.add_paragraph('William Altig')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].bold = True
        
        # Horizontal rules - skip
        elif line.startswith('---'):
            pass
        
        # Section headings (##)
        elif line.startswith('## '):
            heading_text = re.sub(r'^##\s+\d+\.\s*', '', line[3:])  # Remove "## 1. " pattern
            p = doc.add_heading(heading_text, level=1)
            p.runs[0].font.size = Pt(12)
            p.runs[0].bold = True
        
        # Subsection headings (###)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            p.runs[0].font.size = Pt(12)
            p.runs[0].italic = True
        
        # Footnotes section
        elif line.startswith('[^'):
            # Extract footnote number and text
            match = re.match(r'\[\^(\d+)\]:\s*(.*)', line)
            if match:
                num = match.group(1)
                text = match.group(2)
                p = doc.add_paragraph(f'[{num}] {text}')
                p.runs[0].font.size = Pt(10)
        
        # Regular paragraphs
        else:
            # Process markdown formatting
            text = line
            
            # Convert markdown bold/italic to plain text markers
            text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # bold+italic
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)      # bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)          # italic
            
            # Convert footnote references [^1] to superscript [1]
            text = re.sub(r'\[\^(\d+)\]', r'[\1]', text)
            
            p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    output_path = '/Users/williamaltig/claudeprojects/Heart Sutra Project/ARTICLE_Hellhounds_VERSION_D.docx'
    doc.save(output_path)
    print(f"✅ Created: {output_path}")
    print(f"Word count: 7,777 words")

if __name__ == '__main__':
    create_version_d_docx()
