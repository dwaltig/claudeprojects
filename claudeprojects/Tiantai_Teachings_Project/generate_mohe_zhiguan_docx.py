import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_text(paragraph):
    """
    Helper function to handle **bold** and *italics* text within a paragraph.
    Uses regex to tokenize the text and applies styles to runs.
    """
    text = paragraph.text
    paragraph.clear() # Clear the initial text to rebuild with runs

    # Pattern: (**.*?**)|(\*.*?)
    pattern = re.compile(r'(\*\*.*?\*\*)|(\*.*?\*)')
    
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
            
        run = paragraph.add_run()
        
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part

def add_markdown_to_document(document, md_file_path):
    """
    Reads a Markdown file and adds its content to the DOCX Document object.
    """
    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: File not found - {md_file_path}")
        p = document.add_paragraph(f"[MISSING FILE: {os.path.basename(md_file_path)}]")
        p.runs[0].bold = True
        return

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            p = document.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        
        # Lists (Bullet points)
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(line[2:], style='List Bullet')
            format_text(p)
        elif re.match(r'^\d+\.', line): # Numbered lists
            text = re.sub(r'^\d+\.\s*', '', line)
            p = document.add_paragraph(text, style='List Number')
            format_text(p)
            
        # Blockquotes
        elif line.startswith('> '):
            p = document.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.5)
            p.italic = True
            
        # Normal Paragraphs
        else:
            p = document.add_paragraph(line)
            format_text(p)

def generate_full_document(file_list, output_path):
    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for index, file_path in enumerate(file_list):
        print(f"Processing: {file_path}")
        add_markdown_to_document(document, file_path)
        
        # Add page break after each file (except the last one)
        if index < len(file_list) - 1:
            document.add_page_break()

    document.save(output_path)
    print(f"Successfully created: {output_path}")

def main():
    base_dir = "Tiantai_Teachings_Project/01_TRANSLATIONS/Great_Cessation"
    scholarly_dir = f"{base_dir}/Scholarly"
    practitioner_dir = f"{base_dir}/Practitioner"
    root_dir = "Tiantai_Teachings_Project/01_TRANSLATIONS/Great_Cessation"

    scholarly_files = [
        f"{scholarly_dir}/Title_Page_SCHOLARLY.md",
        f"{root_dir}/Copyright_Page.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Fascicle_1_Part_1_SCHOLARLY_MASTER.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Fascicle_1_Part_2_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Fascicle_1_Part_3_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Fascicle_1_Part_4_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_1_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_2_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_3_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_02_Explaining_the_Name_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_03_Characteristics_of_Substance_Part_1_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_03_Characteristics_of_Substance_Part_2_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_04_Inclusion_of_Dharmas_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_05_Biased_and_Perfect_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_06_Preparatory_Expedients_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_1_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_The_Inconceivable_Object_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_2_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_3_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_4_SCHOLARLY.md",
        f"{scholarly_dir}/Mohe_Zhiguan_Chapter_07_Secondary_Objects_Summary_SCHOLARLY.md",
    ]

    practitioner_files = [
        f"{practitioner_dir}/Title_Page_PRACTITIONER.md",
        f"{root_dir}/Copyright_Page.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Fascicle_1_Part_1_PRACTITIONER_MASTER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Fascicle_1_Part_2_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Fascicle_1_Part_3_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Fascicle_1_Part_4_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_1_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_2_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_01_Grand_Purpose_Part_3_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_02_Explaining_the_Name_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_03_Characteristics_of_Substance_Part_1_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_03_Characteristics_of_Substance_Part_2_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_04_Inclusion_of_Dharmas_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_05_Biased_and_Perfect_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_06_Preparatory_Expedients_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_1_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_The_Inconceivable_Object_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_2_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_3_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_Correct_Contemplation_Part_4_PRACTITIONER.md",
        f"{practitioner_dir}/Mohe_Zhiguan_Chapter_07_Secondary_Objects_Summary_PRACTITIONER.md",
    ]

    print("Generating Scholarly Document...")
    generate_full_document(
        scholarly_files, 
        f"{scholarly_dir}/Mohe_Zhiguan_FULL_SCHOLARLY.docx"
    )

    print("\nGenerating Practitioner Document...")
    generate_full_document(
        practitioner_files, 
        f"{practitioner_dir}/Mohe_Zhiguan_FULL_PRACTITIONER.docx"
    )

if __name__ == "__main__":
    main()