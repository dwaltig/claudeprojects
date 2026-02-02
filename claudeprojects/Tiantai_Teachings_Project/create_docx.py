import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_file_path, docx_file_path):
    """
    Converts a simple Markdown file to a DOCX file.
    Handles headers (#, ##, etc.), bold (**text**), and italics (*text*).
    """
    document = Document()
    
    # Set default style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    except FileNotFoundError:
        print(f"Error: File not found - {md_file_path}")
        return

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            p = document.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=4)
        
        # Lists (Bullet points)
        elif line.startswith('* ') or line.startswith('- '):
            p = document.add_paragraph(line[2:], style='List Bullet')
            format_text(p)
        elif re.match(r'^\d+\.', line): # Numbered lists
            # Remove the number and dot (e.g., "1. ")
            text = re.sub(r'^\d+\.\s*', '', line)
            p = document.add_paragraph(text, style='List Number')
            format_text(p)
            
        # Blockquotes
        elif line.startswith('> '):
            p = document.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.5)
            p.italic = True
            
        # Normal Paragraphs
        else:
            p = document.add_paragraph(line)
            format_text(p)

    document.save(docx_file_path)
    print(f"Successfully created: {docx_file_path}")

def format_text(paragraph):
    """
    Helper function to handle **bold** and *italics* text within a paragraph.
    Uses regex to tokenize the text and applies styles to runs.
    """
    text = paragraph.text
    paragraph.clear() # Clear the initial text to rebuild with runs

    # Regex to capture **bold** and *italics*
    # Group 1: **bold**
    # Group 2: *italics*
    # The non-capturing groups (?:...) are used to split but we want to keep delimiters for logic or just matching groups.
    # Actually, a simpler way is to split and keep delimiters, or iterate.
    
    # Pattern: (**.*?**)|(\*.*?\*)
    pattern = re.compile(r'(\*\*.*?\*\*)|(\*.*?\*)')
    
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
            
        run = paragraph.add_run()
        
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]
            run.italic = True
        else:
            run.text = part

def main():
    base_dir = "Tiantai_Teachings_Project/01_TRANSLATIONS/Great_Cessation"
    
    files_to_convert = [
        {
            "src": f"{base_dir}/Scholarly/Mohe_Zhiguan_ESSENTIALS_SCHOLARLY.md",
            "dest": f"{base_dir}/Scholarly/Mohe_Zhiguan_ESSENTIALS_SCHOLARLY.docx"
        },
        {
            "src": f"{base_dir}/Practitioner/Mohe_Zhiguan_ESSENTIALS_PRACTITIONER.md",
            "dest": f"{base_dir}/Practitioner/Mohe_Zhiguan_ESSENTIALS_PRACTITIONER.docx"
        }
    ]

    for item in files_to_convert:
        markdown_to_docx(item["src"], item["dest"])

if __name__ == "__main__":
    main()