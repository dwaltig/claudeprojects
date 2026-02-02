#!/usr/bin/env python3
"""
Generate DOCX files for Vimalakirti Sutra Volume 3 (Chapters 10-14)
Creates separate Scholarly and Blues edition documents.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

# Configuration
TRANSLATIONS_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project/01_TRANSLATIONS")
OUTPUT_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project/06_FINAL_DOCUMENTS")

VOLUME_3_CHAPTERS = [
    ("Chapter_10_Accumulated_Fragrance", "Chapter 10: Accumulated Fragrance (香積佛品)"),
    ("Chapter_11_Bodhisattva_Practices", "Chapter 11: Bodhisattva Practices (菩薩行品)"),
    ("Chapter_12_Seeing_Akshobhya", "Chapter 12: Seeing Akṣobhya Buddha (見阿閦佛品)"),
    ("Chapter_13_Dharma_Offering", "Chapter 13: Dharma Offering (法供養品)"),
    ("Chapter_14_Entrusting", "Chapter 14: Entrusting (囑累品)"),
]

def create_title_page(doc, edition):
    """Add a title page to the document."""
    # Add some spacing
    for _ in range(5):
        doc.add_paragraph()
    
    # Title
    title = doc.add_heading("The Vimalakīrti Nirdeśa Sūtra", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"\n{edition} Edition")
    subtitle_run.font.size = Pt(18)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Volume
    volume = doc.add_paragraph()
    volume_run = volume.add_run("\nVolume 3: Chapters 10-14")
    volume_run.font.size = Pt(16)
    volume.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Source
    source = doc.add_paragraph()
    source_run = source.add_run("\n\nTranslated from the Classical Chinese of Kumārajīva (T0475)")
    source_run.font.size = Pt(12)
    source_run.italic = True
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Translator
    translator = doc.add_paragraph()
    if edition == "Scholarly":
        translator_text = "Scholarly Translation by Dr. Rajesh Patel"
    else:
        translator_text = "Blues Translation by Master John Kim"
    trans_run = translator.add_run(f"\n\n{translator_text}")
    trans_run.font.size = Pt(14)
    translator.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Page break
    doc.add_page_break()

def add_toc(doc, chapters):
    """Add a table of contents."""
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    for _, title in chapters:
        p = doc.add_paragraph()
        p.add_run(title)
    
    doc.add_page_break()

def add_content_to_doc(doc, content, chapter_title):
    """Add markdown content to document with basic formatting."""
    # Add chapter heading
    heading = doc.add_heading(chapter_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')
    in_verse = False
    
    for line in lines:
        line = line.strip()
        
        # Skip the title line (already added as heading)
        if line.startswith('# Chapter') or line.startswith('# '):
            continue
        
        # Skip metadata lines
        if line.startswith('**Scholarly Translation') or line.startswith('**Blues Translation'):
            p = doc.add_paragraph()
            run = p.add_run(line.replace('**', ''))
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Handle horizontal rules
        if line == '---':
            doc.add_paragraph('─' * 50)
            continue
            
        # Handle headings
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
            
        # Handle verse (lines starting with *)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.italic = True
            p.paragraph_format.left_indent = Inches(0.5)
            continue
            
        # Handle bold text with **
        if line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            continue
            
        # Handle numbered lists
        if line and line[0].isdigit() and '. ' in line[:4]:
            doc.add_paragraph(line, style='List Number')
            continue
            
        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            continue
            
        # Regular paragraphs
        if line:
            # Convert markdown formatting
            p = doc.add_paragraph()
            # Simple handling - just add the text
            # Remove markdown bold/italic markers for cleaner output
            clean_line = line.replace('**', '').replace('*', '')
            p.add_run(clean_line)
        else:
            # Empty line = paragraph break
            doc.add_paragraph()

def create_volume_docx(edition):
    """Create a DOCX for the specified edition."""
    doc = Document()
    
    # Set up document margins
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Add title page
    create_title_page(doc, edition)
    
    # Add TOC
    add_toc(doc, VOLUME_3_CHAPTERS)
    
    # Add each chapter
    for file_base, chapter_title in VOLUME_3_CHAPTERS:
        suffix = "Scholarly" if edition == "Scholarly" else "Blues"
        file_path = TRANSLATIONS_DIR / f"{file_base}_{suffix}.md"
        
        if file_path.exists():
            print(f"Adding: {file_path.name}")
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            add_content_to_doc(doc, content, chapter_title)
            doc.add_page_break()
        else:
            print(f"Warning: File not found: {file_path}")
    
    # Save document
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT_DIR / f"Vimalakirti_Sutra_Volume_3_{edition.upper()}.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path

def main():
    print("=" * 60)
    print("Vimalakirti Sutra Volume 3 DOCX Generator")
    print("=" * 60)
    
    # Create Scholarly edition
    print("\n--- Creating Scholarly Edition ---")
    scholarly_path = create_volume_docx("Scholarly")
    
    # Create Blues edition
    print("\n--- Creating Blues Edition ---")
    blues_path = create_volume_docx("Blues")
    
    print("\n" + "=" * 60)
    print("Generation Complete!")
    print(f"Scholarly: {scholarly_path}")
    print(f"Blues: {blues_path}")
    print("=" * 60)

if __name__ == "__main__":
    main()
