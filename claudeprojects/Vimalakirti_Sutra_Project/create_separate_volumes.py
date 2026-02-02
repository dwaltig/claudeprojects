#!/usr/bin/env python3
"""
Generate two separate Word documents for Vimalakirti Sutra Volume 1:
1. Scholarly-only version
2. Blues-only version
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

TRANSLATIONS_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project/01_TRANSLATIONS")
OUTPUT_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project")

SCHOLARLY_CHAPTERS = [
    "Chapter_1_Buddha_Land_Scholarly.md",
    "Chapter_2_Skillful_Means_Scholarly.md",
    "Chapter_3_Disciples_Scholarly.md",
    "Chapter_4_Bodhisattvas_Scholarly.md",
]

BLUES_CHAPTERS = [
    "Chapter_1_Buddha_Land_Blues.md",
    "Chapter_2_Skillful_Means_Blues.md",
    "Chapter_3_Disciples_Blues.md",
    "Chapter_4_Bodhisattvas_Blues.md",
]

def markdown_to_docx_paragraph(doc, line):
    """Convert markdown line to Word paragraph."""
    if not line.strip():
        doc.add_paragraph()
        return
    
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith("#### "):
        doc.add_heading(line[5:].strip(), level=4)
    elif line.strip() in ["---", "***", "___"]:
        p = doc.add_paragraph()
        p.add_run("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif line.startswith("> "):
        doc.add_paragraph(line[2:].strip(), style="Quote")
    else:
        p = doc.add_paragraph()
        text = line.strip()
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = p.add_run(part[1:-1])
                run.italic = True
            else:
                p.add_run(part)


def process_file(doc, filepath):
    """Process markdown file into document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    for line in content.split('\n'):
        markdown_to_docx_paragraph(doc, line)


def create_volume(chapters, output_name, subtitle, translator):
    """Create a volume document."""
    doc = Document()
    
    # Style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Title
    title = doc.add_heading("The Vimalakirti Sutra", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    trans = doc.add_paragraph()
    trans.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trans.add_run(translator).italic = True
    
    doc.add_paragraph()
    
    # TOC
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("Chapter 1: Buddha Land (佛國品)")
    doc.add_paragraph("Chapter 2: Skillful Means (方便品)")
    doc.add_paragraph("Chapter 3: Disciples (弟子品)")
    doc.add_paragraph("Chapter 4: Bodhisattvas (菩薩品)")
    
    doc.add_page_break()
    
    # Chapters
    for i, chapter in enumerate(chapters, 1):
        print(f"  Processing {chapter}...")
        filepath = TRANSLATIONS_DIR / chapter
        if filepath.exists():
            process_file(doc, filepath)
            doc.add_page_break()
    
    # Save
    output_path = OUTPUT_DIR / output_name
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    print("Creating Scholarly Volume...")
    create_volume(
        SCHOLARLY_CHAPTERS,
        "Vimalakirti_Sutra_Volume_1_SCHOLARLY.docx",
        "Volume 1: Chapters 1-4 (Scholarly Edition)",
        "Translation by Dr. Rajesh Patel"
    )
    
    print("\nCreating Blues Volume...")
    create_volume(
        BLUES_CHAPTERS,
        "Vimalakirti_Sutra_Volume_1_BLUES.docx",
        "Volume 1: Chapters 1-4 (Blues Edition)",
        "Translation by Master John Kim"
    )
    
    print("\nDone!")
