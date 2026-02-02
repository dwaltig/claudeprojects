#!/usr/bin/env python3
"""
Generate Word document for Vimalakirti Sutra Volume 1.
Combines all four chapters (both Scholarly and Blues versions).
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

TRANSLATIONS_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project/01_TRANSLATIONS")
OUTPUT_FILE = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project/Vimalakirti_Sutra_Volume_1.docx")

CHAPTERS = [
    ("Chapter_1_Buddha_Land_Scholarly.md", "Chapter_1_Buddha_Land_Blues.md"),
    ("Chapter_2_Skillful_Means_Scholarly.md", "Chapter_2_Skillful_Means_Blues.md"),
    ("Chapter_3_Disciples_Scholarly.md", "Chapter_3_Disciples_Blues.md"),
    ("Chapter_4_Bodhisattvas_Scholarly.md", "Chapter_4_Bodhisattvas_Blues.md"),
]

def markdown_to_docx_paragraph(doc, line, current_style="Normal"):
    """Convert a markdown line to a Word paragraph with basic formatting."""
    # Skip empty lines
    if not line.strip():
        doc.add_paragraph()
        return
    
    # Handle headers
    if line.startswith("# "):
        p = doc.add_heading(line[2:].strip(), level=1)
        return
    elif line.startswith("## "):
        p = doc.add_heading(line[3:].strip(), level=2)
        return
    elif line.startswith("### "):
        p = doc.add_heading(line[4:].strip(), level=3)
        return
    elif line.startswith("#### "):
        p = doc.add_heading(line[5:].strip(), level=4)
        return
    
    # Handle horizontal rules
    if line.strip() in ["---", "***", "___"]:
        p = doc.add_paragraph()
        p.add_run("─" * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    
    # Handle block quotes
    if line.startswith("> "):
        p = doc.add_paragraph(line[2:].strip(), style="Quote")
        return
    
    # Regular paragraph - handle inline formatting
    p = doc.add_paragraph()
    
    # Process inline formatting (bold, italic)
    text = line.strip()
    
    # Simple approach: split by formatting markers and process
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def process_markdown_file(doc, filepath, is_scholarly=True):
    """Process a markdown file and add to document."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    for line in lines:
        markdown_to_docx_paragraph(doc, line)


def create_volume_1_docx():
    """Create the Volume 1 Word document."""
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_heading("The Vimalakirti Sutra", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Volume 1: Chapters 1-4")
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    translators = doc.add_paragraph()
    translators.alignment = WD_ALIGN_PARAGRAPH.CENTER
    translators.add_run("Scholarly Translation by Dr. Rajesh Patel\n").italic = True
    translators.add_run("Blues Translation by Master John Kim").italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Table of Contents placeholder
    toc = doc.add_heading("Contents", level=1)
    doc.add_paragraph("Chapter 1: Buddha Land (佛國品)")
    doc.add_paragraph("Chapter 2: Skillful Means (方便品)")
    doc.add_paragraph("Chapter 3: Disciples (弟子品)")
    doc.add_paragraph("Chapter 4: Bodhisattvas (菩薩品)")
    
    doc.add_page_break()
    
    # Process each chapter
    for i, (scholarly, blues) in enumerate(CHAPTERS, 1):
        print(f"Processing Chapter {i}...")
        
        # Scholarly version
        scholarly_path = TRANSLATIONS_DIR / scholarly
        if scholarly_path.exists():
            process_markdown_file(doc, scholarly_path, is_scholarly=True)
            doc.add_page_break()
        
        # Blues version
        blues_path = TRANSLATIONS_DIR / blues
        if blues_path.exists():
            process_markdown_file(doc, blues_path, is_scholarly=False)
            doc.add_page_break()
    
    # Save
    doc.save(OUTPUT_FILE)
    print(f"\nSaved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    create_volume_1_docx()
