#!/usr/bin/env python3
"""
Generate separate Scholarly and Blues DOCX files for Vimalakirti Sutra Volume 2
"""

import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    exit(1)

BASE_DIR = Path("/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project")
TRANS_DIR = BASE_DIR / "01_TRANSLATIONS"

VOLUME_2_CHAPTERS = [
    ("Chapter_5_Manjushri_Inquires", "Chapter 5: Manjushri Inquires About the Illness"),
    ("Chapter_6_Inconceivable", "Chapter 6: The Inconceivable"),
    ("Chapter_7_Observing_Beings", "Chapter 7: Observing Sentient Beings"),
    ("Chapter_8_Buddha_Path", "Chapter 8: The Buddha Path"),
    ("Chapter_9_Non_Duality", "Chapter 9: Entering the Gate of Non-Duality"),
]

def read_markdown(filepath):
    """Read markdown file and return content."""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def add_content_to_doc(doc, content, chapter_title):
    """Add markdown content to document with basic formatting."""
    # Add chapter heading
    heading = doc.add_heading(chapter_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = content.split('\n')
    in_verse = False
    
    for line in lines:
        line = line.strip()
        
        # Skip the title line (already added as heading)
        if line.startswith('# Chapter') or line.startswith('# '):
            continue
        
        # Skip metadata lines
        if line.startswith('**Scholarly Translation') or line.startswith('**Blues Translation'):
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        # Handle horizontal rules
        if line == '---':
            doc.add_paragraph('─' * 50)
            continue
            
        # Handle headings
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            continue
            
        # Handle verse (lines starting with *)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.italic = True
            p.paragraph_format.left_indent = Inches(0.5)
            continue
            
        # Handle regular paragraphs
        if line:
            # Handle footnote markers
            line = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)  # Remove bold markers
            p = doc.add_paragraph(line)

def create_volume(chapters, suffix, title):
    """Create a DOCX file for the given chapters."""
    doc = Document()
    
    # Add title page
    title_para = doc.add_heading(f'The Vimalakirti Sutra\nVolume 2', level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Add table of contents header
    doc.add_heading('Table of Contents', level=1)
    for _, chapter_title in chapters:
        doc.add_paragraph(chapter_title)
    
    doc.add_page_break()
    
    # Process each chapter
    for file_prefix, chapter_title in chapters:
        filename = f"{file_prefix}_{suffix}.md"
        filepath = TRANS_DIR / filename
        
        if filepath.exists():
            print(f"  Processing {filename}...")
            content = read_markdown(filepath)
            add_content_to_doc(doc, content, chapter_title)
            doc.add_page_break()
        else:
            print(f"  WARNING: {filename} not found!")
    
    return doc

def main():
    print("Creating Scholarly Volume 2...")
    scholarly_doc = create_volume(
        VOLUME_2_CHAPTERS, 
        "Scholarly",
        "Scholarly Edition\nTranslation by Dr. Rajesh Patel"
    )
    scholarly_path = BASE_DIR / "Vimalakirti_Sutra_Volume_2_SCHOLARLY.docx"
    scholarly_doc.save(scholarly_path)
    print(f"Saved: {scholarly_path}")
    
    print("\nCreating Blues Volume 2...")
    blues_doc = create_volume(
        VOLUME_2_CHAPTERS,
        "Blues", 
        "Blues Edition\nTranslation by Master John Kim"
    )
    blues_path = BASE_DIR / "Vimalakirti_Sutra_Volume_2_BLUES.docx"
    blues_doc.save(blues_path)
    print(f"Saved: {blues_path}")
    
    print("\nDone!")

if __name__ == "__main__":
    main()
