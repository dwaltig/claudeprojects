#!/usr/bin/env python3
"""
Create ElevenReader-optimized Word documents for Vimalakirti Sutra Blues translations.
ElevenReader Formatting Expert Rules Applied:
- No Table of Contents
- Chapter titles as Heading 1
- Section titles as Heading 2
- Endnotes/Notes moved to APPENDIX at end
- Clean paragraphs for natural TTS reading
- Footnote markers removed
- Markdown artifacts cleaned
- Author: William Altig (not AI agents)
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# Paths
PROJECT_DIR = "/Users/williamaltig/claudeprojects/Vimalakirti_Sutra_Project"

VOLUMES = [
    {
        "source": "Vimalakirti_Sutra_Volume_1_BLUES.md",
        "output": "Vimalakirti_Sutra_Volume_1_ELEVENREADER.docx",
        "title": "The Vimalakirti Sutra: Volume 1",
        "subtitle": "Chapters 1-5"
    },
    {
        "source": "Vimalakirti_Sutra_Volume_2_BLUES.md",
        "output": "Vimalakirti_Sutra_Volume_2_ELEVENREADER.docx",
        "title": "The Vimalakirti Sutra: Volume 2",
        "subtitle": "Chapters 6-9"
    },
    {
        "source": "Vimalakirti_Sutra_Volume_3_BLUES.md",
        "output": "Vimalakirti_Sutra_Volume_3_ELEVENREADER.docx",
        "title": "The Vimalakirti Sutra: Volume 3",
        "subtitle": "Chapters 10-14"
    },
    {
        "source": "05_COMPILED/Vimalakirti_Sutra_Blues_Complete.md",
        "output": "Vimalakirti_Sutra_Complete_ELEVENREADER.docx",
        "title": "The Vimalakirti Sutra: Complete Blues Edition",
        "subtitle": "Complete Translation"
    },
]

def clean_for_tts(text):
    """Remove all TTS-unfriendly elements aggressively."""
    # 0. Remove Table of Contents Block
    # Look for "## Table of Contents" and remove everything until the next horizontal rule
    text = re.sub(r'## Table of Contents.*?---', '', text, flags=re.DOTALL)

    # 1. Remove HTML comments
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    
    # 2. Remove footnote superscript markers (¹²³⁴⁵⁶⁷⁸⁹⁰)
    text = re.sub(r'[¹²³⁴⁵⁶⁷⁸⁹⁰]+', '', text)
    
    # 3. Remove Images
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    
    # 4. Remove Links but keep text: [text](url) -> text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # 5. Remove Bold/Italic formatting
    # Note: We do this iteratively to handle some nesting, though meaningful text is preserved
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)     # Italic
    text = re.sub(r'__(.*?)__', r'\1', text)     # Bold (underscore)
    text = re.sub(r'_(.*?)_', r'\1', text)       # Italic (underscore)
    
    # 6. Remove List markers (turn lists into plain lines)
    # Unordered list bullets (*, -, +) at start of line
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
    # Ordered list numbers (1., 2.) at start of line
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
    
    # 7. Remove Blockquotes (>)
    text = re.sub(r'^>\s*', '', text, flags=re.MULTILINE)
    
    # 8. Remove Horizontal Rules
    text = re.sub(r'^\s*[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    
    # 9. Remove standalone hash symbols on their own lines (but keep headings)
    text = re.sub(r'^#+\s*$', '', text, flags=re.MULTILINE)
    
    # 10. Replace AI agent references with author name
    text = text.replace('Master John Kim', 'William Altig')
    text = text.replace('by Master John Kim', 'by William Altig')
    text = text.replace('Blues Translation by Master John Kim', 'Blues Vernacular Translation by William Altig')
    
    # 11. Clean up multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

def create_volume_docx(volume_info):
    """Create a Word document optimized for ElevenReader."""
    source_path = os.path.join(PROJECT_DIR, volume_info["source"])
    output_path = os.path.join(PROJECT_DIR, volume_info["output"])
    
    if not os.path.exists(source_path):
        print(f"  Warning: {source_path} not found, skipping...")
        return False
    
    with open(source_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Clean for TTS
    content = clean_for_tts(content)
    lines = content.split('\n')
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(volume_info["title"])
    run.bold = True
    run.font.size = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(volume_info["subtitle"])
    run.italic = True
    run.font.size = Pt(16)
    
    author_line = doc.add_paragraph()
    author_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author_line.add_run("\n\nBlues Vernacular Translation by William Altig")
    run.font.size = Pt(12)

    series_line = doc.add_paragraph()
    series_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = series_line.add_run("\nSeries: Dharma Reborn, Volume 7")
    run.italic = True
    run.font.size = Pt(10)
    
    doc.add_page_break()
    
    endnotes_buffer = []
    theological_buffer = []
    capture_endnotes = False
    capture_theological = False
    current_chapter = ""
    chapter_count = 0
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # Skip author attribution lines
        if 'Blues Translation by' in stripped or 'Blues Vernacular Translation by' in stripped:
            continue

        # Skip Markdown Table of Contents links (redundant with clean_for_tts but safe)
        if stripped.startswith('- [') and '](' in stripped:
            continue
        
        # Generic Header Detection
        header_match = re.match(r'^(#+)\s+(.*)', stripped)
        if header_match:
            hash_str = header_match.group(1)
            text = header_match.group(2).strip()
            level = len(hash_str)
            
            # Special Handling for Endnotes/Theological headers
            if "Endnotes" in text and level <= 2:
                capture_endnotes = True
                capture_theological = False
                endnotes_buffer.append(f"\n### {current_chapter}\n")
                continue
            
            if "Theological Notes" in text and level <= 2:
                capture_endnotes = False
                capture_theological = True
                theological_buffer.append(f"\n### {current_chapter}\n")
                continue
            
            # If we hit any other header, stop capturing notes
            capture_endnotes = False
            capture_theological = False
            
            # Special Logic for Chapter 1 references or "Chapter X"
            if level == 1:
                # Page break logic for major specific sections if seemingly a new chapter
                # Arbitrary rule: If it says "Chapter", treat as new unit
                if "Chapter" in text or "Volume" in text:
                     # Save previous chapter notes catch-all (logic already implicit by buffers)
                     pass
                
                # Update current chapter context
                current_chapter = text
                chapter_count += 1
                
                if chapter_count > 1 and ("Chapter" in text or chapter_count == 2):
                    # Add page break before new major sections except the very first one
                    doc.add_page_break()
            
            # Add the heading to the doc
            # Map levels: # -> 1, ## -> 2, ### -> 3
            # Limit to 3 for simplicity, or 4 if needed
            doc_level = min(level, 3) 
            doc.add_heading(text, level=doc_level)
            continue

        # Detect end of chapter marker (legacy check)
        if stripped.startswith('**End of Chapter'):
            capture_endnotes = False
            capture_theological = False
            continue
        
        # Capture notes for appendix
        if capture_endnotes:
            endnotes_buffer.append(stripped)
            continue
        
        if capture_theological:
            theological_buffer.append(stripped)
            continue
        
        # Skip standalone # or ---
        if stripped == '#' or stripped == '---':
            continue
        
        # Regular paragraph
        doc.add_paragraph(stripped)
    
    # ==========================================
    # APPENDIX: ENDNOTES
    # ==========================================
    if endnotes_buffer:
        doc.add_page_break()
        doc.add_heading("Appendix: Endnotes", level=1)
        doc.add_paragraph("These scholarly annotations are grouped by chapter for reference.")
        
        for line in endnotes_buffer:
            if line.startswith('\n### '):
                chapter_title = line.strip().replace('### ', '')
                doc.add_heading(chapter_title, level=2)
            else:
                doc.add_paragraph(line)
    
    # ==========================================
    # APPENDIX: THEOLOGICAL NOTES
    # ==========================================
    if theological_buffer:
        doc.add_page_break()
        doc.add_heading("Appendix: Theological Notes", level=1)
        doc.add_paragraph("These deeper reflections explore the Blues-Buddhist resonances.")
        
        for line in theological_buffer:
            if line.startswith('\n### '):
                chapter_title = line.strip().replace('### ', '')
                doc.add_heading(chapter_title, level=2)
            else:
                doc.add_paragraph(line)
    
    # Save
    doc.save(output_path)
    return True

def main():
    print("Creating ElevenReader-optimized Vimalakirti Sutra volumes...")
    print("=" * 60)
    
    for volume in VOLUMES:
        print(f"\nProcessing: {volume['title']}...")
        success = create_volume_docx(volume)
        if success:
            print(f"  ✓ Created: {volume['output']}")
        else:
            print(f"  ✗ Failed to create {volume['output']}")
    
    print("\n" + "=" * 60)
    print("ElevenReader formatting applied to all volumes:")
    print("  - Title pages with author William Altig")
    print("  - Chapter titles as Heading 1")
    print("  - Section titles as Heading 2")
    print("  - Endnotes moved to APPENDIX")
    print("  - Theological Notes moved to APPENDIX")
    print("  - Footnote markers removed")
    print("  - 'Master John Kim' replaced with 'William Altig'")
    print("  - Clean paragraphs for natural TTS reading")

if __name__ == "__main__":
    main()
