import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration
BASE_DIR = "/Users/williamaltig/claudeprojects/Miao-lo/Guketsu_Project/01_TRANSLATIONS"
OUTPUT_FILE = "/Users/williamaltig/claudeprojects/Guketsu_Scholarly_Publication_Ready.docx"

def get_fascicle_files(fascicle_num):
    """Finds and sorts files for a specific fascicle."""
    files = [f for f in os.listdir(BASE_DIR) if f.startswith(f"Fascicle_{fascicle_num:02d}") and "Scholarly" in f]
    
    sections = []
    preface = None
    opening = None
    closing = None
    
    for f in files:
        if "Preface" in f:
            preface = f
        elif "Opening" in f:
            opening = f
        elif "Closing" in f:
            closing = f
        elif "Section" in f:
            match = re.search(r"Section_?0?(\d+)", f)
            if match:
                 sec_num = int(match.group(1))
                 sections.append((sec_num, f))
        elif "Index" in f:
             pass 

    sections.sort(key=lambda x: x[0])
    
    ordered_files = []
    if preface: ordered_files.append(preface)
    if opening: ordered_files.append(opening)
    for _, f in sections: ordered_files.append(f)
    if closing: ordered_files.append(closing)
    
    return ordered_files

def read_markdown_file(filename):
    """Reads markdown file and returns lines."""
    path = os.path.join(BASE_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()

def clean_text_fragment(text):
    """Removes stray markdown symbols from text segments."""
    return text.replace('*', '').replace('#', '').replace('|', '')

def parse_markdown_to_runs(paragraph, text):
    """
    Parses a string with markdown (**bold**, *italic*) and adds runs to the paragraph.
    Aggressively cleans syntax markers.
    """
    # 1. Bold: kw**bold**kw
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            # BOLD segment
            content = part[2:-2]
            # Check for inner italic
            subparts = re.split(r'(\*.*?\*)', content)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) >= 2:
                    # Inner italic inside bold
                    run_text = subpart[1:-1]
                    run = paragraph.add_run(clean_text_fragment(run_text))
                    run.bold = True
                    run.italic = True
                else:
                    if subpart:
                        run = paragraph.add_run(clean_text_fragment(subpart))
                        run.bold = True
        else:
            # NON-BOLD segment. Check for Italic.
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) >= 2:
                    run_text = subpart[1:-1]
                    run = paragraph.add_run(clean_text_fragment(run_text))
                    run.italic = True
                else:
                    if subpart:
                        paragraph.add_run(clean_text_fragment(subpart))

def create_table_from_lines(doc, table_lines):
    """Creates a Word table from markdown table lines."""
    if not table_lines:
        return

    headers = [c.strip() for c in table_lines[0].strip('| \n').split('|')]
    rows = []
    
    for line in table_lines[2:]:
        if not line.strip(): continue
        row_cells = [c.strip() for c in line.strip('| \n').split('|')]
        rows.append(row_cells)
    
    if not headers and not rows: return
    
    # Validation: rows might have diff length than headers if pipe missing
    cols_count = len(headers)
    table = doc.add_table(rows=1, cols=cols_count)
    table.style = 'Table Grid'
    
    # Fill headers
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i]._element.clear_content()
            p = hdr_cells[i].add_paragraph()
            parse_markdown_to_runs(p, h)
            for run in p.runs:
                run.bold = True

    # Fill rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                 row_cells[i]._element.clear_content()
                 p = row_cells[i].add_paragraph()
                 parse_markdown_to_runs(p, cell_data)

def clean_header_text(text):
    """Removes 'Part N' and 'Fascicle N' artifacts from headers."""
    # Removes ", Part \d+" or "Part \d+: " 
    text = re.sub(r', Part \d+', '', text)
    text = re.sub(r'Part \d+: ', '', text)
    # Remove "Fascicle \d+" prefix since that is now the parent header
    text = re.sub(r'^Fascicle \d+[,:]? ', '', text)
    return text.strip()

def process_markdown_to_docx():
    doc = Document()
    
    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    doc.add_paragraph('The Guketsu Commentary (Selected Fascicles)\nBy Zhanran (Miao-lo)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Add Table of Contents ---
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()

    # Buffer for sequential content
    content_buffer = [] 
    
    # --- Add General Introduction ---
    print("Adding General Introduction...")
    content_buffer.append(('heading_1', 'General Introduction'))
    intro_path = "/Users/williamaltig/claudeprojects/GUKETSU_INTRODUCTION.md"
    if os.path.exists(intro_path):
        with open(intro_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            for line in lines:
                line = line.strip()
                if not line or line.startswith('# '): continue 
                if line.startswith('* ') or line.startswith('- '):
                     content_buffer.append(('list_bullet', line[2:]))
                else:
                     content_buffer.append(('paragraph', line))
    content_buffer.append(('page_break', '')) 
    
    # Table parsing state
    in_table = False
    table_buffer = []

    # Process files
    for i in range(1, 11):
        files = get_fascicle_files(i)
        if not files: continue
        
        print(f"Processing Fascicle {i}...")
        
        if i > 1:
            content_buffer.append(('page_break', ''))
        
        # Add Semantic H1 for the Fascicle
        content_buffer.append(('heading_1', f"Fascicle {i}"))

        for filename in files:
            lines = read_markdown_file(filename)

            for line in lines:
                line = line.strip()
                if not line:
                    if in_table:
                        content_buffer.append(('table', table_buffer))
                        table_buffer = []
                        in_table = False
                    continue

                # --- Metadata Stripping ---
                if line.startswith('<!--') or line.startswith('Next:') or line.startswith('**Next') or line.startswith('**Translated by'):
                    continue
                
                # Check for various metadata key formats (bolded or plain)
                clean_line = line.replace('*', '').strip()
                if clean_line.startswith('Source:') or clean_line.startswith('Author:') or clean_line.startswith('Translator:') or clean_line.startswith('Date:') or clean_line.startswith('CBETA:') or clean_line.startswith('Section:'):
                    continue
                
                # Robust Stripping
                if "The Architect" in clean_line or "MAS Agent" in clean_line: continue
                if "Colophon" in clean_line: continue
                if "🎉" in clean_line: continue
                if "Translation continues" in clean_line or "complete. Translation continues" in clean_line or "translation complete" in clean_line or "fully translated" in clean_line: continue

                if "END OF FASCICLE" in line.upper() or "END OF SECTION" in line.upper(): continue
                if line.startswith('**END'): continue
                if line.startswith('---'): continue
                if clean_line == "Footnotes" or clean_line == "## Footnotes": continue

                # --- Header Processing ---
                if line.startswith('# '):
                    text = clean_header_text(line[2:])
                    # Demote file-level title (#) to H2, as H1 is now the Fascicle container
                    content_buffer.append(('heading_2', text))
                    continue

                if line.startswith('## '):
                    header_text = line[3:]
                    # Skip distinct "Translation" or "Scholarly Translation" headers
                    if "Scholarly Translation" in header_text or "Translation" == header_text.strip() or "Colophon" in header_text or "Footnotes" in header_text:
                        continue
                    # Demote ## to H3
                    content_buffer.append(('heading_3', clean_header_text(header_text)))
                    continue
                
                if line.startswith('### '):
                     # Demote ### to H4
                    content_buffer.append(('heading_4', clean_header_text(line[4:])))
                    continue

                if line.startswith('#### '):
                     # Demote #### to H5 (or bold paragraph if preferred, sticking to H5 for structure)
                     # Word standard headings usually go deep enough.
                    content_buffer.append(('heading_5', clean_header_text(line[5:])))
                    continue

                # --- Table Parsing ---
                if line.startswith('|'):
                    in_table = True
                    table_buffer.append(line)
                    continue
                elif in_table:
                    content_buffer.append(('table', table_buffer))
                    table_buffer = []
                    in_table = False
                
                # --- List Bullet Handling ---
                if line.startswith('* ') or line.startswith('- '):
                    content = line[2:].strip()
                    content_buffer.append(('list_bullet', content))
                    continue

                # --- Blockquote ---
                if line.startswith('> '):
                    content = line[2:].strip()
                    if content.startswith('*') and content.endswith('*'): content = content[1:-1]
                    content_buffer.append(('blockquote', content))
                    continue

                # --- Standard Paragraph ---
                content_buffer.append(('paragraph', line))
            
            # Catch trailing table
            if in_table:
                content_buffer.append(('table', table_buffer))
                table_buffer = []
                in_table = False

    # --- Append Methodology Appendix ---
    print("Appending Methodology Appendix...")
    content_buffer.append(('page_break', ''))
    
    methodology_path = "/Users/williamaltig/claudeprojects/TRANSLATION_METHODOLOGY.md"
    if os.path.exists(methodology_path):
        with open(methodology_path, 'r', encoding='utf-8') as f:
            meth_lines = f.readlines()
            for line in meth_lines:
                line = line.strip()
                if not line: continue
                if line.startswith('# '):
                    # Title of Appendix
                    content_buffer.append(('heading_1', line[2:]))
                elif line.startswith('## '):
                    content_buffer.append(('heading_2', line[3:]))
                elif line.startswith('### '):
                    content_buffer.append(('heading_3', line[4:]))
                elif line.startswith('* ') or line.startswith('- '):
                    content_buffer.append(('list_bullet', line[2:]))
                else:
                    content_buffer.append(('paragraph', line))

    # --- Rendering ---
    print("Rendering Content...")
    for item_type, content in content_buffer:
        if item_type == 'table':
            create_table_from_lines(doc, content)
        elif item_type == 'page_break':
            doc.add_page_break()
        elif item_type.startswith('heading_'):
            level = int(item_type.split('_')[1])
            p = doc.add_heading(level=level)
            parse_markdown_to_runs(p, content)
        elif item_type == 'list_bullet':
            p = doc.add_paragraph(style='List Bullet')
            parse_markdown_to_runs(p, content)
        elif item_type == 'blockquote':
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            parse_markdown_to_runs(p, content)
            for run in p.runs: run.italic = True
        else:
            p = doc.add_paragraph(style='Normal')
            parse_markdown_to_runs(p, content)

    doc.save(OUTPUT_FILE)
    print(f"Successfully created: {OUTPUT_FILE}")

if __name__ == "__main__":
    process_markdown_to_docx()
