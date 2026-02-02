from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_blues_manuscript():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading("THE LUTHIER'S HANDBOOK", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Wiring the Rig of Enlightenment\n\nVolume 1: The Setup & The Mechanics")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # Introduction
    doc.add_heading('INTRODUCTION: THE OVERTURE', 1)
    
    doc.add_paragraph("This project ain't about poetry. It's about mechanics. Miao-lo is telling you that if you think God is only in your heart and not in your shoes, you're walking barefoot on broken glass. Everything sings. Everything.")
    
    doc.add_heading('Track 1: The Silence and the Noise', 2)
    doc.add_paragraph("""Miao-lo is looking at the scene in the 8th Century and seeing two types of bad roadies destroying the music scene.

First, you got the 'Empty Rig' guys. They say, 'The music is just silence, man, so why bother plugging in the amp?' They think 'Void' means you don't need the Gear (The Teachings). Miao-lo says they're just lazy. If you don't have a path, you’re just standing in the dark.

Second, you got the 'No Manual' guys. They say, 'I don't need no sheet music, I just play what I feel!' Miao-lo says they’re violating the 'Physics of the Groove.' Even the Legends—the ancient Buddhas—they left Blueprints (The Sutras). They knew that if you don't write down the wiring diagram, the next generation is gonna blow every fuse in the building.

He’s looking at the scene and saying, 'It’s a mess out here.' Ten guys walk into the shop claiming they can fix a guitar, but maybe one of 'em actually knows how to use a soldering iron. We need the Manual back on the bench.""")

    doc.add_heading('Track 2: The Lineage (Checking the Warranty)', 2)
    doc.add_paragraph("""Zhiyi is The Architect. He designed the first High-Fidelity sound system for the soul. He wasn't just guessing; he mapped out the 'One Great Groove' so people wouldn't drown in the noise.

Then he handed the keys to Guanding, who was the Recording Engineer—he made sure every word was taped. It went through a few more hands until it got to Me (Miao-lo).

And I’m looking at these old tapes and realizing the quality is starting to drop. People are misreading the labels. The sound is getting fuzzy. It’s time for a Remaster.""")

    doc.add_heading('Track 3: The Ten Intentions', 2)
    doc.add_paragraph("""Miao-lo steps up to the mic and says, 'Look, I didn't want to write this manual. I had to. The gear is breaking down, and nobody knows how to fix it anymore.'

1. I Ain't Making This Up: This isn't my solo album. This is the original track. I'm just cleaning up the hiss.
2. Don't Forget Where You Came From: You guys learned the chords, but then you went off chasing some new fad and forgot the Roots.
3. The Telephone Game: If I don't write this down now, fifty years from now, people are gonna think 'Middle Way' means 'Mediocre.' The signal is getting weak.
4. The Bedroom Players: For the cats practicing alone in their rooms who don't have a teacher to show them the fingering. This manual is for you.
5. Theory AND Practice: You can't just read the sheet music (Doctrine) and you can't just bang on the strings (Contemplation). You need both. This book connects the Theory to the Gig.
6. The Key Changes: I'm gonna show you where the song shifts gears—the 'Joints' (Guanjie). If you miss these changes, you'll be playing in the wrong key.
7. Preserving the Legacy: I'm archiving the Master Tapes so they don't rot in the basement.
8. My Own Cheat Sheet: Honestly? I need this too. Even the Head Engineer needs a schematic sometimes so he doesn't wire the power to the ground.
9. Fact-Check Me: I'm putting my notes on the table. If I'm wrong, tell me. I ain't too proud to learn.
10. The Big Love: At the end of the day, it's all for the audience. We do this so the people can hear the music.""")

    doc.add_heading('Track 4: The Lost Tapes', 2)
    doc.add_paragraph("""Master Zhiyi did three big sessions at the Yuquan Temple. 
Take 1: Was a massive 20-track monster. Way too long, lots of feedback. 
Take 2: They cut it down to 10 tracks. Cleaned it up a bit. 
Take 3: This is the one we’re using.

Now, some of the roadies in the back are whispering that Take 3 is just the 'Radio Edit'—you know, the short version for people with no attention span. They think Take 2 is the 'Deep Cut.'

Miao-lo says: 'Nah, man. Check the levels.' He actually sat down and counted the words. He realized Take 3 isn't shorter; it's just tighter. It’s the Final Master. Guanding went back in and edited the EQ, compressed the signal, and removed the noise. 

So don't go looking for those rough demos. Take 3 is the Remastered Definitive Edition. It’s got all the same notes, but the clarity is 100% higher.""")

    doc.add_heading('Track 5: The Wiring of the Title (MAHA)', 2)
    doc.add_paragraph("""Why do we call it 'Great' (Maha)? It ain't just because it's a long book. Miao-lo is breaking down the word 'MAHA'. He says it's a Power Chord. It's got three notes played at the exact same time:

1. 'Great' (Emptiness): This is the Bass Note. It's huge. It swallows everything up. Emptiness means there are no boundaries, no fences. It's the Open Road. That's 'Great.'
2. 'Many' (Provisional): This is the Rhythm. The thousands of different sounds, the intricate fingerpicking, the details of life. You got sorrow, joy, rent, whiskey, love—that's the 'Many.'
3. 'Superior' (The Middle Way): This is the Harmonic. It's the sweet spot where the Bass (Empty) and the Rhythm (Many) lock in perfectly. It's better than just silence, and it's better than just noise. It's Music. That's 'Superior.'

So when you say 'Maha,' you aren't saying 'Big.' You're saying 'The Whole Damn Thing.'""")

    doc.add_page_break()
    doc.add_heading('PART II: THE RIG SETUPS (Fascicle 2)', 1)

    doc.add_heading('Track 24: Constant Sitting (The Studio Session)', 2)
    doc.add_paragraph("""We call it 'One Practice' (Yi Xing) because you only do One Thing: you Sit. No walking, no dancing, no chores. The Three Inputs: You gotta wire your Body, Mouth, and Mind to the same channel. If your Body is sitting but your Mouth is whispering and your Mind is wandering, you got 'Phase Cancellation.' The signal disappears. The Venue: You can play with the full band in the meditation hall, but if you really want to capture the 'Master Take,' you gotta go to a Private Studio (a secluded spot). No distractions. No outside noise. The Engineering Goal: You're trying to reach a state where there is zero friction between your life and the One Great Groove.""")

    doc.add_heading('Track 28: Constant Walking (The Marathon Tour)', 2)
    doc.add_paragraph("""This is the High-Stakes Session. In the shop, we call this the 'Stress Test.' The Buddha-Standing Rig: The name 'Buddha-Standing' means the Buddhas themselves come out from the wings and stand right there on the stage with you. But they only show up if you keep the rhythm going. The No-Stop Clause: You walk. 90 days. He says 'not a moment's rest.' This means you don't slow down the BPM. You don't take a 'Five.' If you stop walking, the voltage drops and the Buddhas disappear. The Fraud Policy: Miao-lo is being a tough roadie here. He says: 'Don't go around calling yourself a "Spiritual Person" if you’re actually just lazy.' The Engineering Goal: You walk until your ego is so exhausted it can't hold up the 'False Self' anymore. When the 'Small I' collapses, the 'Great Groove' (The Buddha) is all that’s left standing.""")

    doc.add_heading('Track 32: The Street Groove (Following Intent)', 2)
    doc.add_paragraph("""This is the Field Recording. You aren't in the studio anymore. You’re out in the street, and everything is a sound source. Miao-lo says: 'If you really want to be a Pro, you gotta be able to find the groove anywhere.' He calls this 'Following your Intent'.

It works through the Four Phases of the Signal:
1. The Silent Circuit (Unarisen): The thought hasn't hit yet. The fader is down. But the potential is there.
2. The Signal Pulse (About to Arise): You feel the impulse. 'I think I want a drink.' 'I think I'm mad.' That’s the electricity starting to move.
3. The Sound (Active Thinking): Now you’re in it. You’re playing the note. 'I am mad.'
4. The Decay (Concluded): The note is over, but the reverb is still in the room.

Miao-lo’s secret? 'Observe the Nine Realms.' If a 'Hell' thought (Anger) pops up, you don't run away. You look at it. You find the 'Three Truths' in the middle of your anger. The Engineering Goal: You turn the 'Nine Realms' (the messy, distorted parts of life) into the 'Buddha Realm' (High Fidelity). You realize that Enlightenment is a Verb—it’s the act of watching your own mind as it moves through the day.""")

    doc.add_page_break()
    doc.add_heading('PART III: THE TEN PEDALS (The 10 Vehicles)', 1)
    
    doc.add_paragraph("Miao-lo calls these the 'Ten Vehicles,' but in the shop, we call 'em the Pedal Board. These are the effects you stomp on to change the reality of the signal.")

    doc.add_heading('Pedal 1: The Hologram (Ichinen Sanzen)', 2)
    doc.add_paragraph("""Most people think their mind is like a small box (a Noun) where they keep their thoughts. Miao-lo says: 'Nah, man. Your mind is a Vibration (a Verb) that contains the Entire Catalog of the universe.'

The 'Jie-er' Moment (The Nano-Pulse): He uses this word Jie-er. It means a tiny, weak, flickering pulse of energy. It’s that split-second thought you have while you’re reaching for your coffee.
The Full Stack (3000 Realms): Miao-lo says that in that one nano-pulse, the Whole Shop is wired in. The 10 Worlds (from Hell to Buddha). The 10 Suchnesses (The Specs). The 3 Realms (The Environment). 10 x 10 x 3 x 10 = 3000.

The Inconceivable Wiring: You can't 'think' your way into this. You can't pick it apart with a screwdriver. It’s 'Inconceivable' because the Subject (You) and the Object (The World) are the Same Signal.

The Engineering Logic: Think of a single screw on the floor of the shop. If you look at it with a regular eye, it’s just a screw. But if you put it under the Tiantai Laser, you see that the chrome on that screw is reflecting the whole workshop. You see the workbench, the guitars, the engineer, and the street outside. Miao-lo is saying: 'Your mind is that chrome screw.'""")

    doc.add_heading('Pedal 10: The Governor Delete (Freedom from Dharma-Love)', 2)
    doc.add_paragraph("""Miao-lo is looking at the Speed Governor.

You’re cruising at the 'Likeness Stage' (Purple). You’re fast, you’re smooth, and you’re safe. The engine is purring perfectly at the limit. But there’s a problem. The 'Governor' (Dharma Love) is kicking in. It’s saying: 'This is fast enough. Don't push it any harder or you might break something.' It’s keeping you safe at the 'Summit' of the ordinary world, but it won't let you break the sound barrier into the True World.

The Trap: Stagnation. You’re stuck at the redline, but you aren't shifting gears.
The Fix: You gotta Delete the Governor. You gotta Bypass the Safety.

You realize that 'Safety' (Attachment to Purity) is just another form of blockage. You cut the wire that limits the engine. You push past the 'Purple' zone into the 'True Red.' It feels dangerous because you're leaving the 'Likeness' behind, but that’s the only way to get to the First Abode (The Destination).

The Lesson: If you love the ride too much, you'll never get out of the car. You gotta be willing to redline the engine to break through the wall.""")

    doc.add_page_break()
    doc.add_heading('PART IV: THE RIG DOCTOR (Fascicle 8)', 1)
    
    doc.add_paragraph("Miao-lo doesn't just talk about theory; he talks about what happens when the gear breaks down in the real world. This is the 'Illness' chapter.")

    doc.add_heading('Track 54: Diagnostics', 2)
    doc.add_paragraph("""Miao-lo is pulling out the Oscilloscope. Before you try to 'fix' the noise, you gotta identify the source. He lists the 6 Causes of Signal Failure:

1. The Four Elements: Your Earth (Structure), Water (Cooling), Fire (Power), and Wind (Signal Flow) are out of balance.
2. Improper Food/Drink: You're feeding the rig bad power (Diet).
3. Improper Meditation: You're using the wrong settings for your system. (Yes, bad meditation can make you sick).
4. External Demons: Someone else is messing with your gear.
5. Mara (The Devil): The 'Head of the Studio' is trying to shut you down.
6. Karma: This is a factory defect from a previous model.

The Pulse Check: He talks about reading the pulse like checking the Waveform.
Spring (Wiry): The signal is tight and tense.
Summer (Surging): The signal is too hot (Overloading).
Autumn (Floating): The signal is weak and drifting.
Winter (Deep): The signal is buried in the noise floor.""")

    doc.add_heading('Track 57: The Six Breaths (The Frequency Generator)', 2)
    doc.add_paragraph("""Miao-lo says: 'Every illness has a resonant frequency.' If you hit the counter-frequency, you cancel out the sickness.

1. Chui: Like blowing out a candle. It fixes the 'Cold.'
2. Hu: Like blowing on hot soup. It fixes the 'Heat.'
3. He: A deep guttural sound. It fixes 'Agitation' (Heart trouble).
4. Si: A hissing sound. It fixes 'Exhaustion' (Lungs).
5. Xu: Cures Phlegm.
6. Xi: Cures Pain.

The Lesson: You are a musical instrument. If you're out of tune, you don't just take a pill. You use Sound (Breath) to vibrate the chassis back into alignment. You literally Hum yourself to health.""")

    doc.save('Miao-lo/Guketsu_Project/01_TRANSLATIONS/Guketsu_Blues.docx')
    print("Expanded Blues Edition created.")

if __name__ == "__main__":
    create_blues_manuscript()
