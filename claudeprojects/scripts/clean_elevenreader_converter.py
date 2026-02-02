#!/usr/bin/env python3
"""
Clean DOCX Converter for ElevenReader
Fixes: excessive whitespace, inconsistent layout, page numbers

Key differences from previous scripts:
- NO page breaks (ElevenReader handles pagination)
- NO page numbers
- Consistent single blank lines between paragraphs
- Clean, minimal formatting
- Maximum 2 consecutive blank lines
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys
import os

def create_clean_document():
    """Create document with clean, consistent styles - NO page numbers"""
    doc = Document()
    
    # Remove any default headers/footers (which might have page numbers)
    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True
        # Clear any existing footer content
        for paragraph in section.footer.paragraphs:
            paragraph.clear()
    
    # Set up clean styles
    styles = doc.styles
    
    # Normal text - clean, readable
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Georgia'
    normal_font.size = Pt(12)
    
    # Heading 1 - Chapter titles
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Georgia'
    h1_font.size = Pt(16)
    h1_font.bold = True
    
    # Heading 2 - Section headers
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Georgia'
    h2_font.size = Pt(14)
    h2_font.bold = True
    
    # Heading 3 - Subsection headers
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Georgia'
    h3_font.size = Pt(12)
    h3_font.bold = True
    
    return doc


def clean_excessive_whitespace(lines):
    """Ensure no more than 1 consecutive blank line"""
    cleaned = []
    blank_count = 0
    
    for line in lines:
        if line.strip() == '':
            blank_count += 1
            if blank_count <= 1:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)
    
    return cleaned


def convert_markdown_to_docx(markdown_file, output_file):
    """Convert markdown to clean DOCX for ElevenReader"""
    
    print(f"Converting: {markdown_file}")
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    lines = clean_excessive_whitespace(lines)
    
    doc = create_clean_document()
    
    # Remove page numbers from all sections
    for section in doc.sections:
        section.footer.is_linked_to_previous = False
        for para in section.footer.paragraphs:
            para.clear()
    
    i = 0
    last_was_blank = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip horizontal rules (--- or ===)
        if stripped in ['---', '==='] or (stripped.startswith('---') and len(stripped) < 10) or (stripped.startswith('===') and stripped.count('=') > 10):
            i += 1
            continue
        
        # Main title (# )
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = doc.add_paragraph(title_text, style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_was_blank = False
        
        # Heading 2 (## )
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_paragraph(heading_text, style='Heading 1')
            last_was_blank = False
        
        # Heading 3 (### )
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_paragraph(heading_text, style='Heading 2')
            last_was_blank = False
        
        # Heading 4 (#### )
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            doc.add_paragraph(heading_text, style='Heading 3')
            last_was_blank = False
        
        # Code blocks (```) - just add content without the markers
        elif stripped.startswith('```'):
            i += 1
            continue
        
        # Bold-only lines (like **Publication Information**)
        elif stripped.startswith('**') and stripped.endswith('**') and stripped.count('**') == 2:
            text = stripped[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            last_was_blank = False
        
        # Blank lines - SKIP entirely (no extra paragraph marks)
        elif stripped == '':
            i += 1
            continue
        
        # Regular content
        elif stripped:
            # Handle inline bold/italic
            text = stripped
            
            # Simple approach: just add the text, processing bold markers
            p = doc.add_paragraph()
            
            # Split on bold markers
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    # Handle italic within non-bold parts
                    italic_parts = re.split(r'(\*[^*]+\*)', part)
                    for ip in italic_parts:
                        if ip.startswith('*') and ip.endswith('*') and len(ip) > 2:
                            run = p.add_run(ip[1:-1])
                            run.italic = True
                        elif ip:
                            p.add_run(ip)
            
            last_was_blank = False
        
        i += 1
    
    # Set consistent margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Ensure no page numbers
        for para in section.footer.paragraphs:
            para.clear()
    
    # Save
    doc.save(output_file)
    print(f"✓ Created: {output_file}")
    print("  - No page numbers")
    print("  - No excessive whitespace")
    print("  - Consistent layout")
    print("  - Ready for ElevenReader")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 clean_elevenreader_converter.py <markdown_file> [output_file]")
        print("\nThis will convert a markdown file to a clean DOCX suitable for ElevenReader.")
        print("If output_file is not specified, it will use the input filename with _CLEAN.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        base = os.path.splitext(input_file)[0]
        output_file = f"{base}_CLEAN.docx"
    
    convert_markdown_to_docx(input_file, output_file)


if __name__ == '__main__':
    main()
