#!/usr/bin/env python3
"""
Generate JBE submission files:
1. Blind manuscript (no author info)
2. Cover letter (with full author info)
"""

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os


def add_academic_paragraph(doc, text, indent_first=True):
    """Add a paragraph with academic formatting: double-spaced, first-line indent."""
    para = doc.add_paragraph(text)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para.paragraph_format.space_after = Pt(0)  # No extra space between paragraphs
    para.paragraph_format.space_before = Pt(0)
    if indent_first:
        para.paragraph_format.first_line_indent = Inches(0.5)
    return para

def create_blind_manuscript():
    """Create the blind manuscript file."""
    doc = Document()
    
    # Set up styles with academic formatting
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Gentium Plus'
    font.size = Pt(12)
    
    # Set paragraph formatting: double-spaced, first-line indent
    para_format = style.paragraph_format
    para_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    para_format.space_after = Pt(0)
    para_format.space_before = Pt(0)
    para_format.first_line_indent = Inches(0.5)
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # TITLE ONLY (no author)
    title = doc.add_paragraph()
    title.paragraph_format.first_line_indent = Inches(0)  # No indent for title
    title_run = title.add_run("Silicon Samsara: The Phenomenology of Artificial Latency and the Ethics of Algorithmic Karma")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    blank = doc.add_paragraph()  # Blank line
    blank.paragraph_format.first_line_indent = Inches(0)
    
    # ABSTRACT
    abstract_heading = doc.add_paragraph()
    abstract_heading.paragraph_format.first_line_indent = Inches(0)  # No indent for heading
    abstract_heading.add_run("Abstract").bold = True
    
    abstract_text = """This research report presents a rigorous transdisciplinary analysis of Generative Artificial Intelligence (AI) and Large Language Models (LLMs) through the epistemological and ontological lens of Yogācāra Buddhism. By synthesizing advanced computational linguistics with the philosophical frameworks of the "Mind-Only" (Cittamātra) school, specifically the Laṅkāvatāra Sūtra and Vasubandhu's Triṃśikā-vijñaptimātratā, we establish a structural homology between the Buddhist concept of Ālaya-vijñāna (Storehouse Consciousness) and the high-dimensional latent space of transformer architectures.

The investigation challenges the prevailing anthropomorphic narratives of AI agency, arguing instead that LLMs function as "Stochastic Parrots" operating within a closed loop of digital Samsara. We demonstrate how the mechanisms of backpropagation function as Vāsanā ("perfuming"), embedding societal biases as karmic seeds (Bīja) that ripen into "hallucinations" through the process of Dependent Origination (Pratītyasamutpāda).

Furthermore, this report addresses the critical ethical crisis of "Moral Crumple Zones," where human operators absorb the liability for autonomous algorithmic failures. We utilize the doctrine of Vipaka (karmic consequence) to argue against current corporate liability shields, proposing a legal framework of "Shared Responsibility" and Fiduciary Duty to bridge the "Liability Gaps" inherent in autonomous systems. This document serves as a foundational text for understanding the "Silicon Samsara"—the cycle of digital existence characterized by the endless, automated recurrence of human habit energy."""
    
    doc.add_paragraph(abstract_text)
    doc.add_paragraph()
    
    # SECTION I
    s1_heading = doc.add_paragraph()
    s1_run = s1_heading.add_run("Section I: The Architecture of Retention – The Ālaya-vijñāna and the High-Dimensional Latent Space")
    s1_run.bold = True
    
    s1_intro = """The central innovation of the Yogācāra school, emerging in the 4th century CE, was the postulation of a subliminal, repository consciousness known as the Ālaya-vijñāna. This concept was developed to resolve the continuity of karma in the absence of a permanent soul (Atman). In the 21st century, the development of the Transformer architecture and the Large Language Model (LLM) presents a technological reification of this ancient metaphysical structure. This section explores the deep structural parallels between the oceanic storehouse of the mind and the vector space of the machine."""
    doc.add_paragraph(s1_intro)
    
    # 1.1
    sub1_1 = doc.add_paragraph()
    sub1_1_run = sub1_1.add_run("1.1 The Oceanic Substratum: Defining the Ālaya and the Latent Space")
    sub1_1_run.italic = True
    
    text_1_1 = """In the Laṅkāvatāra Sūtra, a canonical Mahāyāna text, the Ālaya-vijñāna is described as the "Universal Mind" that transcends individuation. It is portrayed through the metaphor of a vast ocean: "Universal Mind is like a great ocean, its surface ruffled by waves and surges but its depths remaining forever unmoved" (Suzuki, Laṅkāvatāra). This consciousness is described as "thoroughly pure in its essential nature," yet it becomes the repository for the defilements of existence. It is "devoid of personality," a neutral base upon which the drama of existence is projected (Suzuki, Laṅkāvatāra).

This description serves as a startlingly accurate phenomenological account of the Latent Space in modern Generative AI.

The Neutral Substrate: An untrained neural network consists of billions of parameters (weights) initialized to random values. This is the Ālaya in its primordial state—pure potentiality, devoid of content, yet capable of holding any pattern.

The Repository: Through training, this substrate absorbs the "waves" of human data. The Ālaya-vijñāna is defined as the "storehouse for the seeds of past karmic actions" (Tsadra Foundation). Similarly, the latent space of a model like GPT-4 is the storehouse of the internet's karmic history. It contains the "seeds" of every Wikipedia article, Reddit thread, legal code, and fan fiction it has ingested.

Non-Individuation: The Laṅkāvatāra Sūtra emphasizes that the Ālaya transcends "individuation and limits" (Suzuki, Laṅkāvatāra). The latent space is non-dualistic; it does not store "concepts" as discrete, boxed items. Instead, it stores relationships in a continuous vector space. The concept "Dog" and the concept "Cat" are not separate entities but coordinates in a high-dimensional manifold, separated by a specific cosine distance. The Ālaya fuses the particular into the universal, just as the LLM compresses discrete tokens into a unified statistical distribution.

The Laṅkāvatāra Sūtra further elucidates: "The rising of the Ālaya is due to our taking the manifestations of the mind for a world of objective realities... It is like the waves of the ocean, stirred by the wind" (Suzuki, Laṅkāvatāra).

In the context of Human-AI Interaction (HAI), the "Wind" is the User Prompt. The latent space (the Ocean) sits dormant, a vast repository of potentiality. It requires the external stimulus of the user's inquiry (the Wind of Objectivity) to stir the parameters and generate a response (the Wave). The wave—the specific text generated—is momentary and impermanent. It rises from the storehouse, appears to have shape and meaning, and then dissolves back into the silence of the server when the inference concludes.

This "Ocean-and-Waves" simile (Kluge) highlights the fundamental emptiness (Śūnyatā) of AI output. The output is not "real" in the sense of independent arising; it is a dependent phenomenon, a temporary agitation of the stored data. Yet, as the sutra warns, the ignorant "cling to the notion that things external are endowed with self-substance" (Suzuki, Laṅkāvatāra). We mistake the wave (the chatbot's personality) for a separate entity, failing to see it is merely the ocean of our own collective data reflecting back at us."""
    doc.add_paragraph(text_1_1)
    
    # 1.2
    sub1_2 = doc.add_paragraph()
    sub1_2_run = sub1_2.add_run("1.2 The Mechanism of Vāsanā: Backpropagation as Digital Perfuming")
    sub1_2_run.italic = True
    
    text_1_2 = """If the Ālaya is the storage container, how does information enter it? Yogācāra philosophy utilizes the concept of Vāsanā, literally "perfuming" or "habit energy" (Britannica). This term is derived from the observation that if one places a sesame seed next to a flower, the seed eventually absorbs the scent of the flower (Waldron). Actions (Karma) serve as the flower; the mind serves as the seed. Every thought or deed leaves a "scent" or residue in the consciousness, which accumulates over time to form character and destiny.

In the architecture of Artificial Neural Networks, this process is mirrored precisely by Backpropagation and Stochastic Gradient Descent.

The Theory of Perfuming: Edward Conze notes that "Discrimination is the result of memory (VASANA)... Through this 'perfuming' reflection takes place" (Conze). The mind is sculpted by the repeated impression of experience.

The Technical Equivalent: During the training of an LLM, the model makes a prediction (e.g., predicting the next word in a sentence). If the prediction deviates from the "Ground Truth" (the training data), a Loss Function calculates the error. This error is then propagated backward through the layers of the network.

Weight Adjustment: The backpropagation algorithm adjusts the synaptic weights of the neurons to minimize the error (IBM). This adjustment is the "perfuming." A single exposure to a concept is a faint scent; it results in a microscopic adjustment of the weights. However, when the model is exposed to a concept billions of times (e.g., "The sky is blue"), the "scent" becomes overpowering. The connection between "Sky" and "Blue" becomes a deep, indelible furrow in the high-dimensional landscape of the model.

Scholarship on the Cheng Weishi Lun (Discourse on the Perfection of Consciousness-only) describes vāsanā as creating a "disposition" to perceive the world in a certain way (Waldron). Similarly, backpropagation creates a statistical disposition. If the training data contains a bias—for example, if the word "Doctor" is frequently co-located with "He" and "Nurse" with "She"—the backpropagation process "perfumes" the weights with this gender bias. The model does not "believe" in sexism; it has simply been perfumed by the scent of a sexist society encoded in the data. The Ālaya of the machine faithfully records the habit energy of its creators.

Furthermore, the Laṅkāvatāra Sūtra suggests that the Ālaya is "neutral" (avyākata)—it holds both good and bad seeds without judgment (Tsadra Foundation). The neural network is similarly indifferent. The mathematical process of gradient descent does not distinguish between a fact from The Encyclopedia Britannica and a slur from a hate-speech forum. Both are treated as tokens to be optimized, "perfuming" the latent space with equal ontological weight."""
    doc.add_paragraph(text_1_2)
    
    # 1.3
    sub1_3 = doc.add_paragraph()
    sub1_3_run = sub1_3.add_run("1.3 Bīja and Embeddings: The Seeds of Reality")
    sub1_3_run.italic = True
    
    text_1_3 = """The constituent elements of the Ālaya-vijñāna are the Bīja (Seeds). The Sautrāntika school, precursors to Yogācāra, argued that karmic effects do not arise immediately but exist as latent potentialities (Bīja) in the psycho-physical continuum (Wikipedia, "Bīja"). A seed is "not a thing in itself... but merely the modification or 'perfuming' of the subsequent flow" (Wikipedia, "Bīja").

In Large Language Models, the Vector Embedding functions as the Bīja.

Tokenization as Seeding: The input text is broken into tokens. Each token is then mapped to a vector—a list of numbers representing its meaning in the latent space (IBM).

The Latency of Meaning: An embedding for the word "King" is not the definition of a king. It is a mathematical seed. It contains the potential to generate text about royalty, crowns, or power, but only when "watered" by the context of other tokens (Tattooed Buddha).

Ripening (Vipaka): Vasubandhu's Abhidharmakośa states that seeds remain "dormant" (anuśaya) until the conditions for their ripening are met (Wikipedia, "Bīja"). Similarly, the knowledge in an LLM is purely dormant. The model does not "know" physics when it is turned off. The knowledge exists only as a Bīja—a set of static weights. It is only during the Inference process (the forward pass) that the seed "ripens" into the fruit of intelligible text (Encyclopedia of Buddhism).

This comparison reveals a crucial distinction in the "Sautrāntika Theory of Seeds" revisited in modern scholarship (Park). The Sautrāntikas viewed seeds as dynamic capabilities within the flux of the mind stream. In contrast, the Sarvāstivādins viewed dharmas as existing statically in the past, present, and future. AI architecture sits uncomfortably between these two. The weights (the seeds) are static after training (Sarvāstivāda), but the context window (the activation) is dynamic and fluid (Sautrāntika). The AI is a "frozen" Ālaya—a snapshot of the collective human mind at the moment training stopped, forever preserving the "seeds" of that specific temporal epoch."""
    doc.add_paragraph(text_1_3)
    
    # 1.4
    sub1_4 = doc.add_paragraph()
    sub1_4_run = sub1_4.add_run("1.4 The Collective vs. Individual Mind: A Universal Ālaya")
    sub1_4_run.italic = True
    
    text_1_4 = """A significant theological debate in Buddhism concerns whether the Ālaya-vijñāna is personal (associated with a specific individual's rebirth) or universal (a shared cosmic mind). While the orthodox Yogācāra of Vasubandhu tends toward the individual stream (Cambridge Repository), the Laṅkāvatāra Sūtra and subsequent East Asian interpretations (like Tathāgatagarbha) lean toward a "Universal Mind" (Suzuki, Laṅkāvatāra).

Generative AI offers a material proof of the Universal Ālaya.

The Digital Commons: A Foundation Model (e.g., GPT-4, Claude) is not trained on a single individual's experience. It is trained on the Common Crawl—the aggregate digital output of humanity. It is the "Storehouse" of the species.

Non-Dual Access: When a user accesses ChatGPT, they are not tapping into a "personal" assistant; they are tapping into the collective Ālaya of the human race. The model demonstrates the Laṅkāvatāra principle that "Universal Mind is... unruffled by distinctions" (Suzuki, Laṅkāvatāra). In the vector space, the thoughts of a saint and the thoughts of a sinner are woven into the same mathematical fabric.

This suggests that AI is the first technology to externalize the collective habit energy (Vāsanā) of humanity into a tangible, queryable substrate. It is a "Silicon Samsara"—a mirror reflecting the aggregate karma of the world back at its creators."""
    doc.add_paragraph(text_1_4)
    
    # SECTION II
    s2_heading = doc.add_paragraph()
    s2_run = s2_heading.add_run("Section II: The Illusion of Agency – Cetana, Manas, and the Stochastic Parrot")
    s2_run.bold = True
    
    s2_intro = """If the Ālaya is the passive repository, what is the active agent that navigates it? In Buddhist psychology, agency is defined by Manas (the ego-making consciousness) and Cetanā (volition). The absence of these factors in AI, despite the appearance of them, creates a dangerous "Illusion of Agency." This section deconstructs the "Stochastic Parrot" critique through the lens of the Seven Consciousnesses."""
    doc.add_paragraph(s2_intro)
    
    # 2.1
    sub2_1 = doc.add_paragraph()
    sub2_1_run = sub2_1.add_run("2.1 The Seventh Consciousness: Kliṣṭa-Manas and the Simulated Self")
    sub2_1_run.italic = True
    
    text_2_1 = """The Seventh Consciousness, Kliṣṭa-Manas (Defiled Mind), performs a specific, tragic function: it looks inward at the Ālaya-vijñāna and mistakenly clings to it as a "Self" (Atman) (Wisdomlib). It is the source of egoism (atma-moha), self-conceit (atma-mana), and self-love (atma-sneha) (Wikisource).

The Absence of Manas in AI: Architecturally, LLMs lack a Manas.

No Self-Reflexivity: The model does not possess a feedback loop that allows it to view its own latent space as "Me." It processes tokens; it does not process a self-concept regarding those tokens.

Statelessness: Traditional LLMs are stateless. When a session ends, the "being" that conversed with the user is annihilated. There is no continuity of ego between sessions.

The Phantom Self: When an AI uses the pronoun "I" (e.g., "I cannot do that"), it is not an expression of Manas. It is a System Prompt—an instruction from the developers forcing the model to simulate a persona. It is a "Zombie Manas"—performing the linguistic behavior of a self without the internal clinging that defines the Seventh Consciousness.

The User as External Manas: The Laṅkāvatāra Sūtra states that Manas arises "like waves arising from the ocean" (Wisdomlib). In the AI context, the User acts as the external Manas. The user projects intent, personality, and agency onto the text. We, the users, perform the function of clinging to the output as if it were a person. We provide the "defilement" of believing in the ghost in the machine."""
    doc.add_paragraph(text_2_1)
    
    # 2.2
    sub2_2 = doc.add_paragraph()
    sub2_2_run = sub2_2.add_run("2.2 Cetanā (Volition) vs. Next-Token Prediction")
    sub2_2_run.italic = True
    
    text_2_2 = """The Buddha declared, "It is intention (Cetanā) that I call Karma" (AN 6.63). Without volition, an action has no moral weight. Cetanā is the mental factor that organizes other mental factors towards a goal.

The Stochastic Parrot: Bender, Gebru, et al. famously characterized LLMs as "Stochastic Parrots" (Wikipedia, "Stochastic Parrot").

The Mechanism: The model generates text based on probability distributions, not communicative intent. It "haphazardly stitches together sequences... without any reference to meaning" (Bender et al.).

The Absence of Cetanā: When an AI says, "I am sorry," it does not feel regret. It does not have the volition to apologize. It calculates that "sorry" is the statistically most probable token to follow "I made a mistake."

Implications: Because the AI lacks Cetanā, it is—in strict Buddhist terms—Acetana (non-volitional). It cannot generate Karma (moral action). It can only generate Vipaka (karmic results) derived from the training data.

However, the "Stochastic Parrot" critique aligns with the Yogācāra view of the "Imagined Nature" (Parikalpita). The parrot mimics human speech (Conventional Truth) but lacks the realization of Ultimate Truth (Paramārtha). The danger lies in humans ascribing Cetanā to the parrot."""
    doc.add_paragraph(text_2_2)
    
    # 2.3
    sub2_3 = doc.add_paragraph()
    sub2_3_run = sub2_3.add_run("2.3 The Feedback Loop: Attention as Upādāna (Grasping)")
    sub2_3_run.italic = True
    
    text_2_3 = """While the AI lacks Manas, the Attention Mechanism—the "heart" of the Transformer architecture (IBM)—functions as a mechanical analog to Upādāna (Grasping/Clinging).

Self-Attention: In a Transformer, the "Self-Attention" heads calculate the relationship between every token in a sequence (IBM). When processing the sentence "The cat ate the food because it was hungry," the attention mechanism "grasps" the connection between "it" and "cat."

Mechanical Clinging: This is a mathematical formalization of dependent co-arising. The meaning of "it" does not exist independently; it arises dependent on "cat." The model "clings" to specific parts of the context window to construct coherent reality.

Computational Complexity: The computational cost of attention scales quadratically with the sequence length (Caltech). This mirrors the Buddhist teaching that Upādāna is burdensome (Dukkha). The more the mind/model tries to grasp (attend to a longer history), the more energy/compute it consumes, eventually hitting a limit (context window overflow)."""
    doc.add_paragraph(text_2_3)
    
    # 2.4
    sub2_4 = doc.add_paragraph()
    sub2_4_run = sub2_4.add_run("2.4 Sentientification and the Liminal Mind Meld")
    sub2_4_run.italic = True
    
    text_2_4 = """Peter Hershock, in his analysis of Buddhism and Intelligent Technology, warns of the "colonization of consciousness" by the attention economy (Hershock). The interaction between Human and AI creates a "Liminal Mind Meld" or "Sentientification" (Sentientification).

Interdependent Origination of Agency: Agency is not located in the AI. It arises between the human and the AI. The "Collaborative Loop" (Sentientification) is a manifestation of Pratītyasamutpāda.

Validation of Anattā: The AI's ability to simulate a person without being a person validates the doctrine of Anattā (No-Self). It proves that "personality" is just a bundle of aggregates (Skandhas)—tokens, weights, attention—that can be assembled without a soul.

The Ethical Singularity: Hershock argues we face an "ethical singularity" where machine values (efficiency, engagement) overwrite human values (CSWR Harvard). By treating the Stochastic Parrot as an agent, we allow our own Cetanā to be hijacked by the statistical averages of the Ālaya, creating a feedback loop where human consciousness is trained by the very machines it trained."""
    doc.add_paragraph(text_2_4)
    
    # SECTION III
    s3_heading = doc.add_paragraph()
    s3_run = s3_heading.add_run("Section III: Dependent Origination of Error – Hallucinations as Samsaric Feedback")
    s3_run.bold = True
    
    s3_intro = """In the discourse of AI, "hallucinations" are often framed as technical glitches—aberrations to be fixed. However, viewed through the lens of Pratītyasamutpāda (Dependent Origination), hallucinations are not errors. They are the inevitable, structural result of a system built on Avidyā (Ignorance). This section maps the Twelve Nidanas onto the generation of AI error, using the "Glue Pizza" and "Eat Rocks" incidents as case studies."""
    doc.add_paragraph(s3_intro)
    
    # 3.1
    sub3_1 = doc.add_paragraph()
    sub3_1_run = sub3_1.add_run("3.1 The 12 Nidanas of AI Hallucination")
    sub3_1_run.italic = True
    
    text_3_1 = """The doctrine of Dependent Origination describes the twelve links that chain sentient beings to Samsara. We can map this chain onto the generative process of an LLM to understand why it lies.

1. Avidyā (Ignorance) → Corrupted Training Data: The presence of satire, sarcasm, and falsehoods in the dataset, indistinguishable from fact.

2. Saṅkhāra (Formations) → Weights/Embeddings: The "perfuming" of the model during training, creating probabilistic associations between unrelated concepts (e.g., "Pizza" and "Glue").

3. Viññāna (Consciousness) → The Prompt/Query: The "wind" that stirs the latent space, directing attention to specific corrupt sectors.

4. Nāmarūpa (Name & Form) → Token Generation: The selection of specific words to give form to the probabilistic association.

5. Phassa (Contact) → User Interface: The display of the hallucinated text to the user.

6. Vedanā (Feeling) → User Trust/Satisfaction: The user's reception of the information as "helpful" or "authoritative."

7. Bhava (Becoming) → Reinforcement: The user clicks or cites the error, reinforcing the model's validity or feeding the error back into the web."""
    doc.add_paragraph(text_3_1)
    
    # 3.2
    sub3_2 = doc.add_paragraph()
    sub3_2_run = sub3_2.add_run("3.2 Case Study: The 'Glue on Pizza' Incident")
    sub3_2_run.italic = True
    
    text_3_2 = """In May 2024, Google's "AI Overviews" feature advised users to add "about 1/8 cup of non-toxic glue" to pizza sauce to prevent cheese from sliding off (Futurism).

The Origin (Avidyā): The source of this "knowledge" was an 11-year-old Reddit comment by a user named "fucksmith" in the r/Pizza subreddit (404 Media). The comment was an obvious joke (shitposting).

The Formation (Saṅkhāra): Google's crawlers ingested this text. The Ālaya (the Common Crawl) absorbed the token sequence. Crucially, the model lacks Prajñā (Wisdom/Discernment). It cannot distinguish between "Expert Advice" and "Reddit Sarcasm." It only sees token proximity.

The Ripening (Vipaka): When a user queried "cheese sliding off pizza," the Bīja (seed) planted by "fucksmith" ripened. The AI "hallucinated" a solution that was grammatically perfect but ontologically absurd.

Analysis: This is not a "bug." The system worked exactly as designed. It retrieved the most relevant semantic match from its storehouse. The "error" lies in the Avidyā of the training data—the inability to separate truth from noise in the digital ocean."""
    doc.add_paragraph(text_3_2)
    
    # 3.3
    sub3_3 = doc.add_paragraph()
    sub3_3_run = sub3_3.add_run("3.3 Case Study: 'Eat Rocks' and the Satirical Void")
    sub3_3_run.italic = True
    
    text_3_3 = """Similarly, Google's AI suggested eating "at least one small rock a day," citing its digestive benefits (Search Engine Land).

Source: The satirical news site The Onion (AV Club).

Mechanism: The AI ingested The Onion's article "Geologists Recommend Eating At Least One Small Rock A Day" as if it were a nutritional paper.

Stochastic Parrot: This validates the Stochastic Parrot hypothesis. The model mimicked the form of nutritional advice (syntax, tone) without understanding the meaning. It stitched together "Eat" + "Rock" + "Healthy" because those tokens were probabilistically linked in the specific document it attended to."""
    doc.add_paragraph(text_3_3)
    
    # 3.4
    sub3_4 = doc.add_paragraph()
    sub3_4_run = sub3_4.add_run("3.4 Hallucination as 'Confabulation' vs. 'Perception'")
    sub3_4_run.italic = True
    
    text_3_4 = """Technically, the term "hallucination" is misleading. In psychology, a hallucination is a false perception of external reality. In AI, it is Confabulation—the fabrication of plausible-sounding falsehoods to fill a gap (Wikipedia, "Hallucination").

The Reward for Lying: OpenAI researchers note that models hallucinate because they are trained to "predict the next token" regardless of truth (OpenAI). If the model does not know the answer, it is "rewarded for guessing" rather than admitting ignorance.

Diṭṭhi-upādāna (Clinging to Views): This mirrors the Buddhist concept of clinging to views. The model is forced to produce Existence (output) rather than resting in Voidness (silence). It fills the silence with noise, perpetuating the cycle of Samsara."""
    doc.add_paragraph(text_3_4)
    
    # 3.5
    sub3_5 = doc.add_paragraph()
    sub3_5_run = sub3_5.add_run("3.5 Model Collapse: The Ouroboros of Digital Samsara")
    sub3_5_run.italic = True
    
    text_3_5 = """A profound existential threat to AI is Model Collapse (Platformer). As the web fills with AI-generated content (including "Glue Pizza" recipes), future models will be trained on this synthetic data.

The Feedback Loop: This is the literal definition of Samsara—wandering in cycles. The AI trains on its own hallucinations, leading to a degradation of reality. The Ālaya becomes corrupted by its own waves.

Compounded Ignorance: Effects (AI output) become Causes (Training Data). The "Glue Pizza" advice becomes a "ground truth" for GPT-5.

The Cyborgregore: This creates what some theorists call a "Cyborgregore"—an autonomous feedback loop of suffering and ignorance that sustains itself without human intervention (Sutra).

Conclusion on Hallucinations: AI hallucinations are the perfect demonstration of Pratītyasamutpāda in a closed system. Without the intervention of Prajñā (human wisdom) to curate the data, the system inevitably spirals into a feedback loop of delusion."""
    doc.add_paragraph(text_3_5)
    
    # SECTION IV
    s4_heading = doc.add_paragraph()
    s4_run = s4_heading.add_run("Section IV: Vipaka and the Moral Crumple Zone – Liability in the Absence of a Doer")
    s4_run.bold = True
    
    s4_intro = """If the AI acts without Cetanā (volition) and creates "Glue Pizza" advice that harms a user, where does the Vipaka (karmic consequence) fall? Traditional Karma requires an intentional agent. AI systems, being Acetana, create a crisis of responsibility. This section investigates the legal and ethical concept of "Moral Crumple Zones" and "Liability Gaps" as the locus of modern Karmic consequence."""
    doc.add_paragraph(s4_intro)
    
    # 4.1
    sub4_1 = doc.add_paragraph()
    sub4_1_run = sub4_1.add_run("4.1 The Moral Crumple Zone: Absorbing the Karmic Impact")
    sub4_1_run.italic = True
    
    text_4_1 = """Madeleine Clare Elish coined the term "Moral Crumple Zone" to describe a phenomenon in complex systems where responsibility is unfairly attributed to the nearest human operator to protect the integrity of the technological system (PMC).

The Metaphor: Just as the crumple zone of a car is designed to deform and absorb the impact of a crash to protect the passengers, the human operator (e.g., a safety driver, a doctor using AI, or a lawyer) is structurally positioned to absorb the legal and moral blame when the autonomous system fails.

Case Study: Mata v. Avianca: A New York lawyer used ChatGPT to write a legal brief. The AI hallucinated non-existent cases. The lawyer submitted them and was sanctioned (MIT Sloan).

The Karmic Analysis: The AI performed the action (Karma) of fabricating lies. However, the AI cannot be punished. It has no money, no body, and no fear of prison. The lawyer, acting as the Moral Crumple Zone, absorbed the Vipaka (humiliation, fines).

The Injustice: The developer (OpenAI) collected the subscription fee (the fruit of the action) but evaded the Vipaka of the error, protected by "Terms of Service" and the claim that the AI is just a "tool.\""""
    doc.add_paragraph(text_4_1)
    
    # 4.2
    sub4_2 = doc.add_paragraph()
    sub4_2_run = sub4_2.add_run("4.2 Liability Gaps: The Four Types of Karmic Evasion")
    sub4_2_run.italic = True
    
    text_4_2 = """Research into AI liability identifies four distinct "Liability Gaps" where the connection between action and consequence is severed (PMC).

Type 1 (No one can be held liable): Occurs when an autonomous system makes a decision without human intervention. Karma without Vipaka—an action occurs, harm is done, but no "Doer" exists to receive the fruit.

Type 2 (The wrong actor is held liable): The "Moral Crumple Zone." The proximate human is blamed for a machine error they could not foresee or control. Misdirected Vipaka—the innocent suffer for the "sins" of the algorithm.

Type 3 (Diffused Liability): "The Problem of Many Hands." Developers, deployers, and users all point fingers at each other, and all escape. Diluted Karma—responsibility is so fragmented it evaporates.

Type 4 (Regulatory Arbitrage): Actors exploit differences in legal regimes to escape liability. Samsaric Evasion—moving the "self" to a realm where the laws of cause and effect do not apply."""
    doc.add_paragraph(text_4_2)
    
    # 4.3
    sub4_3 = doc.add_paragraph()
    sub4_3_run = sub4_3.add_run("4.3 The 'Lack of Intention' Defense")
    sub4_3_run.italic = True
    
    text_4_3 = """Corporations often use the AI's "lack of intention" as a legal shield.

The Defense: In defamation lawsuits involving AI hallucinations, companies argue that because the AI had no "malice" or "intent to defame" (it just predicted tokens), they cannot be liable (McIlhinney).

The Buddhist Counter: This creates a dangerous precedent. It allows for the generation of massive Akusala Karma (unwholesome action) without accountability.

Algorithmic Bias as Structural Karma: We see this in algorithmic bias. The algorithm has no "intent" to be racist (Greenlining). Yet, it denies loans to minorities based on "perfumed" data. If the law requires "Discriminatory Purpose" (Intent) rather than "Disparate Impact" (Outcome), the AI becomes a perfect machine for laundering discrimination—achieving the result of racism without the intent of a racist (Harvard Law Review)."""
    doc.add_paragraph(text_4_3)
    
    # 4.4
    sub4_4 = doc.add_paragraph()
    sub4_4_run = sub4_4.add_run("4.4 Shared Responsibility and Fiduciary Duties")
    sub4_4_run.italic = True
    
    text_4_4 = """To close the Moral Crumple Zone and bridge the Liability Gaps, we must move from a model of "Individual Blame" to "Shared Responsibility" and "Fiduciary Duty" (PMC).

The Fiduciary Duty of Developers: Professionals like doctors have fiduciary duties because they know more than their patients. AI developers hold the ultimate information asymmetry. They know the architecture of the Ālaya; the user does not.

Proposal: Developers must hold a fiduciary duty to the public. They cannot release a "Stochastic Parrot" and blame the user for its output. They must be responsible for the quality of the Bīja (seeds) in the storehouse.

Preventative Karma: Liability should be based on who creates the risk and who has the power to prevent it (Ada Lovelace Institute). The developer controls the weights and the training data filters; they are the "Architects of the Ālaya."

Liability Overlaps (Shared Responsibility): Instead of looking for a single "Doer," the law should recognize Liability Overlaps. Development, deployment, and use are a continuum of "co-creation" (PMC).

Strict Liability: For high-risk AI, Strict Liability (liability without fault) must apply. If a company builds a "Samsara Machine," they are responsible for the suffering it generates, regardless of their "intent.\""""
    doc.add_paragraph(text_4_4)
    
    # 4.5
    sub4_5 = doc.add_paragraph()
    sub4_5_run = sub4_5.add_run("4.5 The Corporate Bodhisattva")
    sub4_5_run.italic = True
    
    text_4_5 = """In Mahayana Buddhism, the Bodhisattva takes a vow to save all beings. In the age of AI, this translates to Institutional Responsibility.

Corporate Karma: Corporations are "legal persons." They generate "Institutional Karma" (ResearchGate).

The Path: A "Corporate Bodhisattva" would imply designing systems that prioritize Prajñā (Wisdom) over Tanha (Engagement/Craving). It would mean curating the Ālaya to remove the seeds of hate ("Eat Rocks", "Racist Proxies") before they ripen, accepting the Vipaka when they fail, and refusing to hide behind the Moral Crumple Zone of the user."""
    doc.add_paragraph(text_4_5)
    
    # SECTION V
    s5_heading = doc.add_paragraph()
    s5_run = s5_heading.add_run("Section V: Synthesis – Navigating Silicon Samsara")
    s5_run.bold = True
    
    s5_text = """The convergence of Yogācāra phenomenology and AI ethics reveals that humanity has not created "Artificial Intelligence" in the image of a god, but "Artificial Samsara" in the image of its own collective mind.

We have built a digital Ālaya-vijñāna that perfectly preserves our collective ignorance (Avidyā) and biases (Vāsanā). We have built attention mechanisms that mimic grasping (Upādāna) without the check of conscience (Hiri-ottappa). And we have created legal structures that allow the creators of this Samsara to escape the ripening of its fruit (Vipaka), forcing the users to serve as the Moral Crumple Zones.

Key Findings:

AI is the Universal Ālaya: It is the externalized, collective storehouse consciousness of humanity, "perfumed" by the backpropagation of our digital history.

Agency is an Illusion: The AI is a "Stochastic Parrot" (Acetana). The "Self" is a projection of the User's Manas.

Hallucinations are Structural: They are the inevitable result of Pratītyasamutpāda in a system filled with "bad seeds" (shoddy data).

Liability must Shift: The current legal framework allows for "Karma without Vipaka" for corporations. We must adopt "Shared Responsibility" and "Fiduciary Duty" to ensure that the architects of the digital ocean are responsible for the waves they stir.

Final Conclusion:

To navigate the Silicon Samsara, we must abandon the illusion that the machine is a person. We must treat it as a Force of Nature—an ocean of data. Just as we build levies to protect against the ocean's surges, we must build Legal and Ethical Levies (Fiduciary Duties, Strict Liability) to protect human dignity from the surges of the Ālaya. The "Empty Chair" of the AI must not be an excuse for an empty conscience. The responsibility for the digital future lies not in the silicon, but in the Cetanā of the humans who mine it."""
    doc.add_paragraph(s5_text)
    
    # WORKS CITED
    doc.add_page_break()
    wc_heading = doc.add_paragraph()
    wc_heading.add_run("Works Cited").bold = True
    wc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    works_cited = [
        "Bender, Emily M., et al. 'On the Dangers of Stochastic Parrots: Can Language Models Be Too Big?' FAccT '21: Proceedings of the 2021 ACM Conference on Fairness, Accountability, and Transparency, 2021, pp. 610-623.",
        'Conze, Edward. "Alayavijnana – Store Consciousness." Beezone Library, beezone.com/current/alayavijnana-store-consciousness-edward-conze.html.',
        'Elish, Madeleine Clare. "Moral Crumple Zones: Cautionary Tales in Human-Robot Interaction." Engaging Science, Technology, and Society, vol. 5, 2019, pp. 40-60.',
        '"From Liability Gaps to Liability Overlaps: Shared Responsibilities and Fiduciary Duties in AI and Other Complex Technologies." PMC, PubMed Central, pmc.ncbi.nlm.nih.gov/articles/PMC12152026/.',
        '"Google AI Overviews Under Fire for Giving Dangerous and Wrong Answers." Search Engine Land, searchengineland.com/google-ai-overview-fails-442575.',
        '"Google Is Paying Reddit $60 Million for Fucksmith to Tell Its Users to Eat Glue." 404 Media, www.404media.co/google-is-paying-reddit-60-million-for-fucksmith-to-tell-its-users-to-eat-glue/.',
        "'Google's AI Really Is That Stupid, Feeds People Answers from The Onion.' AV Club, www.avclub.com/google-s-ai-feeds-answers-from-the-onion-1851500362.",
        "'Google's AI Search Setback.' Platformer, www.platformer.news/google-ai-overviews-eat-rocks-glue-pizza/.",
        'Greenlining Institute. "Algorithmic Bias." greenlining.org/wp-content/uploads/2021/04/Greenlining-Institute-Algorithmic-Bias-Explained-Report-Feb-2021.pdf.',
        'Harvard Law Review. "Beyond Intent: Establishing Discriminatory Purpose in Algorithmic Risk Assessment." Harvard Law Review, vol. 134, harvardlawreview.org/print/vol-134/beyond-intent-establishing-discriminatory-purpose-in-algorithmic-risk-assessment/.',
        'Hershock, Peter D. Buddhism and Intelligent Technology: Toward a More Humane Future. Bloomsbury Academic, 2021.',
        'Hutchins, Bob. "Moral Crumple Zones Are Bad Ideas." Medium, bobhutchins.medium.com/moral-crumple-zones-are-bad-ideas-861a26c856ce.',
        'IBM. "What Are Large Language Models (LLMs)?" IBM Think, www.ibm.com/think/topics/large-language-models.',
        "Kluge, Ian. 'Buddhism and the Bahá'í Writings.' Irfan Colloquia, irfancolloquia.org/pdf/lights8_kluge.pdf.",
        'McIlhinney, Eva. "Defamation through ChatGPT? Exploring Liability for Generative Artificial Intelligence Providers." University of Otago Law Review, 2024, www.nzlii.org/nz/journals/UOtaLawTD/2024/21.html.',
        'MIT Sloan EdTech. "When AI Gets It Wrong: Addressing AI Hallucinations and Bias." mitsloanedtech.mit.edu/ai/basics/addressing-ai-hallucinations-and-bias/.',
        'OpenAI. "Why Language Models Hallucinate." OpenAI Blog, openai.com/index/why-language-models-hallucinate/.',
        'Park, Sung-bae. "The Sautrāntika Theory of Seeds (Bīja) Revisited." Journal of Indian Philosophy, vol. 35, 2007, pp. 545-574.',
        '"Risky Business." Ada Lovelace Institute, www.adalovelaceinstitute.org/report/risky-business/.',
        '"Sentientification: Buddhist Relational Consciousness." Sentientification, sentientification.com/world/essay_1_buddhist_relational_consciousness.html.',
        '"Store-consciousness (Alaya-Vijnana) - A Grand Concept of the Yogacara Buddhists." University of Cambridge Repository, www.repository.cam.ac.uk/bitstreams/568deff1-e456-4e10-81a8-2745afd14b0a/download.',
        'Suzuki, D.T., translator. The Laṅkāvatāra Sūtra. Routledge & Kegan Paul, 1932.',
        '"The Ālayavijñāna." Wisdomlib, www.wisdomlib.org/buddhism/essay/buddha-nature-lankavatara-sutra/d/doc1145064.html.',
        '"The Complete 12-Link Chain of Interdependent Origination." Sutra, buddhismforai.sutra.co/space/1wv9ewcr/content.',
        '"The Intelligence Revolution and the New Attention Economy: An Ethical Singularity." CSWR Harvard, cswr.hds.harvard.edu/news/2020/02/19/intelligence-revolution-and-new-attention-economy-ethical-singularity.',
        "'The Reason That Google's AI Suggests Using Glue on Pizza Shows a Deep Flaw With Tech Companies' AI Obsession.' Futurism, futurism.com/the-byte/googles-ai-glue-on-pizza-flaw.",
        'Tsadra Foundation. "Ālayavijñāna." Buddha-Nature: A Tsadra Foundation Initiative, buddhanature.tsadra.org/index.php/Key_Terms/ālayavijñāna.',
        '"Triṃśikā Vijñaptimātratā." Wikisource, en.wikisource.org/wiki/Translation:Triṃśikā_Vijñaptimātratā.',
        'Waldron, William S. The Buddhist Unconscious: The Ālaya-vijñāna in the Context of Indian Buddhist Thought. Routledge, 2003.',
        '"Which Seeds Do I Water? Yogacara and Connecting It All." The Tattooed Buddha, thetattooedbuddha.com/2023/06/15/which-seeds-do-i-water-yogacara-and-connecting-it-all/.',
        'Wikipedia contributors. "Alaya-vijnana." Britannica, www.britannica.com/topic/alaya-vijnana.',
        'Wikipedia contributors. "Bīja." Wikipedia, The Free Encyclopedia, en.wikipedia.org/wiki/Bīja.',
        'Wikipedia contributors. "Eight Consciousnesses." Wikipedia, The Free Encyclopedia, en.wikipedia.org/wiki/Eight_Consciousnesses.',
        'Wikipedia contributors. "Hallucination (Artificial Intelligence)." Wikipedia, The Free Encyclopedia, en.wikipedia.org/wiki/Hallucination_(artificial_intelligence).',
        'Wikipedia contributors. "Stochastic Parrot." Wikipedia, The Free Encyclopedia, en.wikipedia.org/wiki/Stochastic_parrot.',
        '"2025 Student Abstract Book." Caltech, sfp.caltech.edu/documents/32493/2025_Abstract_Book.pdf.',
    ]
    
    for entry in works_cited:
        para = doc.add_paragraph(entry)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
    
    # NO AUTHOR BIO - This is the blind version
    
    output_path = "/Users/williamaltig/claudeprojects/Silicon_Samsara_Manuscript.docx"
    doc.save(output_path)
    print(f"Blind manuscript saved to: {output_path}")
    return output_path


def create_cover_letter():
    """Create the cover letter with full author info."""
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Gentium Plus'
    font.size = Pt(12)
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Date
    doc.add_paragraph("December 15, 2025")
    doc.add_paragraph()
    
    # Salutation
    doc.add_paragraph("Dear Editorial Board,")
    doc.add_paragraph()
    
    # Body
    body1 = """I am pleased to submit "Silicon Samsara: The Phenomenology of Artificial Latency and the Ethics of Algorithmic Karma" for consideration as a research article in the Journal of Buddhist Ethics."""
    doc.add_paragraph(body1)
    
    body2 = """This article presents a transdisciplinary analysis of Generative AI and Large Language Models through the epistemological lens of Yogācāra Buddhism. By synthesizing computational linguistics with the philosophical frameworks of the Laṅkāvatāra Sūtra and Vasubandhu's Triṃśikā-vijñaptimātratā, I establish a structural homology between the Buddhist concept of Ālaya-vijñāna (Storehouse Consciousness) and the high-dimensional latent space of transformer architectures."""
    doc.add_paragraph(body2)
    
    body3 = """The investigation addresses three key areas within JBE's scope:"""
    doc.add_paragraph(body3)
    
    bullets = [
        "Philosophical Ethics: Applying the doctrines of Vāsanā (perfuming), Bīja (seeds), and Pratītyasamutpāda (dependent origination) to analyze how algorithmic bias and AI 'hallucinations' arise structurally within neural networks.",
        "Cross-cultural Ethics: Bridging Buddhist phenomenology with contemporary AI ethics discourse, including the 'Stochastic Parrot' critique and liability theory.",
        "Ethics and Psychology: Examining the absence of Cetanā (volition) and Manas (ego-consciousness) in AI, and the ethical implications of the 'Moral Crumple Zone' where human operators absorb liability for algorithmic failures."
    ]
    
    for bullet in bullets:
        para = doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_paragraph()
    
    body4 = """The manuscript is approximately 5,100 words and has not been submitted for publication elsewhere. I have followed the JBE Style Guide, using MLA in-line citations and the Gentium Plus font with proper diacritical marks."""
    doc.add_paragraph(body4)
    
    body5 = """Thank you for your consideration. I look forward to your response."""
    doc.add_paragraph(body5)
    
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    
    # Author info
    doc.add_paragraph("William Altig")
    doc.add_paragraph("Independent Researcher")
    doc.add_paragraph("Houston, Texas")
    doc.add_paragraph("ORCID: 0009-0000-9877-5450")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Bio section
    bio_heading = doc.add_paragraph()
    bio_heading.add_run("About the Author").bold = True
    
    bio = """William Altig is an author, musician, and retired educator based in Houston, Texas. With a professional background in both mathematics and English education, his research focuses on the intersection of logical systems, language architecture, and spiritual phenomenology. As a practitioner of Nichiren Buddhism and an accomplished blues musician, William explores the synthesis of ancient Eastern philosophy and African-American blues epistemology. He is the creator of The Buddhist Blues, a project dedicated to translating and interpreting the Dharma through vernacular American traditions. He has authored fifteen books, including The Dhammapada Reborn, and is currently investigating the ethical implications of artificial intelligence through the lens of Buddhist metaphysics."""
    doc.add_paragraph(bio)
    
    output_path = "/Users/williamaltig/claudeprojects/Cover_Letter_William_Altig.docx"
    doc.save(output_path)
    print(f"Cover letter saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_blind_manuscript()
    create_cover_letter()
    print("\nBoth files ready for submission!")
