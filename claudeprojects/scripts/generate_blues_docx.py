from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

input_file = "madhyama-kasha-shastra/01_TRANSLATIONS/Mūlamadhyamakakārikā_Blues_Complete.md"
output_file = "madhyama-kasha-shastra/01_TRANSLATIONS/Mūlamadhyamakakārikā_Blues_Lyrics_Only.docx"

document = Document()

# Set up styles if needed, or just use direct formatting
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

def clean_text(text):
    # Remove markdown bold/italic markers
    text = text.replace('**', '').replace('*', '').replace('>', '').strip()
    return text

def is_lyric_line(line):
    stripped = line.strip()
    if not stripped: return False
    if re.match(r'^\(Verse', stripped, re.IGNORECASE): return True
    if stripped.startswith("(") and stripped.endswith(")"): return True
    if stripped.startswith(">"): return True
    if "|" in stripped: return False
    if stripped.startswith("#"): return False
    if stripped.startswith("Now,"): return False
    if stripped.startswith("See,"): return False
    if stripped.startswith("Let me"): return False
    if stripped.startswith("Here's"): return False
    if stripped.startswith("That's"): return False
    if stripped.startswith("So "): return False
    if stripped.startswith("But "): return False
    if stripped.startswith("And "): return False
    if stripped.endswith("?"): return False
    if len(stripped) > 80: return False
    return True

with open(input_file, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Title Page
document.add_heading('Mūlamadhyamakakārikā Blues', 0)
document.add_paragraph('The Lyrics', style='Subtitle')
document.add_paragraph('Interpreter: Master John Kim', style='Subtitle')
document.add_page_break()

for line in lines:
    stripped = line.strip()
    
    # 1. Chapters (Heading 1)
    if stripped.startswith("# Madhyamakashāstra Chapter") or stripped.startswith("# Chapter"):
        clean_header = clean_text(stripped.replace("# Madhyamakashāstra Chapter", "").replace("# Chapter", "Chapter"))
        document.add_heading(clean_header, level=1)
        continue

    # 2. Song Titles (Heading 2)
    if stripped.startswith("**Title:") or stripped.startswith("### Title:") or stripped.startswith("## *") or stripped.startswith("##"):
        raw_title = stripped
        if "Title:" in raw_title:
            raw_title = raw_title.split("Title:")[1]
        
        title_text = clean_text(raw_title)
        if title_text:
            h2 = document.add_heading(title_text, level=2)
            h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    # 3. Lyrics & Verse Markers
    if is_lyric_line(line):
        # Check for Verse Marker
        if stripped.startswith("(Verse") or stripped.startswith("(Chorus"):
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)
            runner = p.add_run(stripped)
            runner.bold = True
        else:
            # It's a lyric line
            # Check if it was italicized in markdown (starts/ends with *)
            is_italic = stripped.startswith("*") or stripped.startswith("> *")
            
            final_text = clean_text(stripped)
            
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0) # Keep lines close
            runner = p.add_run(final_text)
            
            if is_italic:
                runner.italic = True
    
    # 4. Handle blank lines for stanza spacing
    if not stripped:
        # If the last element added was a paragraph, ensure there's some breathing room
        # But docx adds space by default, so maybe just pass
        pass

document.save(output_file)
print(f"Created {output_file}")
