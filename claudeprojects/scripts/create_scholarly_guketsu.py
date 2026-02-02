from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_scholarly_manuscript():
    doc = Document()
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title Page
    doc.add_heading('The Great Calming and Contemplation', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('The Guketsu Commentary (Selected Fascicles)\nBy Zhanran (Miao-lo)\nTranslated by Dr. Rajesh Patel').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # FASCICLE 1: THE INTRODUCTION
    doc.add_heading('Fascicle 1: The Introduction', 1)
    
    doc.add_heading('1. The Necessity of the Commentary', 2)
    doc.add_paragraph("Question: 'For what causes and conditions have you presumed to compile this record?'")
    doc.add_paragraph("Answer: 'The matter was unavoidable. I recount the causes for this work, which are ten in total:")
    doc.add_paragraph("1. To demonstrate that there is a Lineage of Transmission, and that this is not based on my own personal speculation or deviation from the Master's mind.")
    doc.add_paragraph("2. For the sake of those who once had the transmission but have Discarded the Root to follow untested views.")
    doc.add_paragraph("3. Because future generations, as the teaching spreads, will give rise to Divergent Interpretations and lose their original basis.")
    doc.add_paragraph("4. For those who have faith in the school but practice in other regions without a master to consult.")
    doc.add_paragraph("5. For those who practice both Doctrine (Yi) and Contemplation (Guan), to ensure their Practice and Understanding are complete.")
    doc.add_paragraph("6. To point out the crucial Articulations (Guanjie - 'Joints') and the profound meaning of the text.")
    doc.add_paragraph("7. To establish the Master's interpretation so it does not fall into ruin, benefiting future ages.")
    doc.add_paragraph("8. To aid my own contemplation and understanding, preventing error and making it easy to research.")
    doc.add_paragraph("9. To expose my own understanding [to scrutiny], fearing I may have forgotten or erred, and seeking correction.")
    doc.add_paragraph("10. To accord with the Buddha's intent, exercising Great Compassion to benefit others.")

    doc.add_heading('2. The Three Versions', 2)
    doc.add_paragraph("'This work exists in three successive versions. The first version was comprised of twenty fascicles. The second version consisted of ten fascicles. Both were originally titled 'The Perfect and Sudden [Calming and Contemplation]' to distinguish them from the biased, small, or indeterminate teachings.")
    doc.add_paragraph("The third version has a slightly different title intent... and begins with the words: 'Calming and Contemplation are Bright and Quiet.' The version I am following now is this third version.")
    doc.add_paragraph("People often pass down the idea that the third version is the 'Abridged' (Lue) text, while the second is the 'Expanded' (Guang) text. Looking at it superficially, they do seem to differ in length. However, upon investigating the text from beginning to end, the page count is nearly identical. Therefore, the third version should be considered the Revised Edition (Zai-zhi ben); it is not necessary to call it 'abridged'.'")

    doc.add_heading('3. The Meaning of MAHA', 2)
    doc.add_paragraph("'There are two reasons why the title was changed to 'Mohe' (Maha/Great):")
    doc.add_paragraph("1. To distinguish it from the 'Small Calming and Contemplation' (Xiao Zhiguan), which Master Zhiyi produced for his lay brother.")
    doc.add_paragraph("2. To preserve the richness of the Sanskrit sound (Fan-yin), which contains inclusive meanings.")
    doc.add_paragraph("The Great Treatise (Da Lun) states: 'The word Maha contains three meanings: Great (Da), Many (Duo), and Superior (Sheng).'")
    doc.add_paragraph("Using these three meanings to interpret the Perfect Threefold Contemplation (Yuan San Guan) fits the title perfectly:")
    doc.add_paragraph("* Great (Da) corresponds to the meaning of Emptiness (Kū).")
    doc.add_paragraph("* Many (Duo) corresponds to the meaning of Provisionality (Ke).")
    doc.add_paragraph("* Superior (Sheng) corresponds to the meaning of the Middle Way (Chū).'")

    doc.add_heading('4. The Three Approaches to Contemplation', 2)
    doc.add_paragraph('1. Gradual (Jian): Compared to "Ladders and Steps" (Ti-Deng). Extreme heights require a ladder; one ascends step by step.')
    doc.add_paragraph('2. Variable (Bu Ding): Compared to a "Diamond Treasure placed in the sun." Its color is indeterminate depending on the angle.')
    doc.add_paragraph('3. Perfect/Sudden (Yuan Dun): Compared to a "Magician Ascending into Space." The space near the ground is identical in substance to the space at ten thousand fathoms; thus, it is called Non-Dual.')

    doc.add_heading('5. The Five Great Meanings', 2)
    doc.add_paragraph('1. Arousing the Great Heart (Bodhicitta): The Self-Practice Cause.')
    doc.add_paragraph('2. Cultivating the Great Practice (The Four Samadhis): The process of the cause.')
    doc.add_paragraph('3. Sensing the Great Fruit (The Three Virtues): The Self-Practice Result.')
    doc.add_paragraph('4. Rending the Great Net (Transforming Others): The power to transform.')
    doc.add_paragraph('5. Returning to the Great Abode (Nirvana): The final destination.')

    # FASCICLE 2: THE FOUR SAMADHIS
    doc.add_heading('Fascicle 2: The Four Samadhis', 1)
    
    doc.add_heading('1. Constant Sitting (One Practice Samadhi)', 2)
    doc.add_paragraph("'It is called 'One Practice' (Yi Xing) based on the teaching it relies on. The name comes from the fact that it simplifies the physical posture, excluding all other activities. It is not called 'One' because of the Principle it contemplates; for if it were based on the Principle, all four Samadhis would be called 'One,' as they all contemplate the True Characteristic.")
    doc.add_paragraph("The methods for this practice must involve the Three Karma (Body, Speech, Mind). Regarding the Body: one may practice in a group in a meditation hall, but a solitary, secluded place is considered the most superior and excellent (Mi Shan). Regarding the Full Lotus (Jie-jia): First the left foot, then the right, aligned with the thighs. The Great Treatise asks: 'There are many ways to sit; why only the Lotus position?' Answer: 'Because it is the most stable and secure (An Wen). It gathers the hands and feet so the mind does not scatter. It causes even the Mara-kings to be terrified.''")

    doc.add_heading('2. Constant Walking (Buddha-Standing Samadhi)', 2)
    doc.add_paragraph("'Constant Walking is called 'Buddha-Standing' (Fo-li). ...Regarding the time: one must not rest even for a 'moment' (Xuyu). A 'moment' is technically one-thirtieth of a day and night, but here it is a general term meaning one must not allow the practice to cease for even a second.")
    doc.add_paragraph("...One must not 'Deceive the Buddhas'. The Great Treatise says: 'If one claims to be a Bodhisattva but is lazy, or craves fame rather than diligently seeking the way, that is self-deception and deception of all Buddhas and Bodhisattvas.''")

    doc.add_heading('3. Half-Walking, Half-Sitting', 2)
    doc.add_paragraph("'Next is Half-Walking, Half-Sitting. First, we explain the Vaipulya (Fang-Deng). It relates to the 'Four Teachings.' ... If one only contemplates the mind internally but neglects the external rituals (Shi-yi), the practice is insufficient.")
    doc.add_paragraph("...One must first seek the Twelve Dream Kings. One is taught to seek these kings; if one sees even one, the 7-day method is granted.")
    doc.add_paragraph("The Lotus Repentance (Fa-hua Chan) is usually performed on fast-days. It should last at least seven days, though the standard is 21 days.'")

    doc.add_heading('4. Neither Walking Nor Sitting (Following Self-Intent)', 2)
    doc.add_paragraph("'Next, regarding the various 'Good' acts... one must first distinguish the Four Phases of Intent (Si Yun). 'Yun' means motion. ...One must encompass the Nine Realms within the Buddha Realm.")
    doc.add_paragraph("If one understands that the Nine Realms are the objects of contemplation, then whether the Bodhisattva realm has arisen or not, it must be observed—how much more so the others? ...The Great Master Wugou (The Pure Master - Fu Dashi) also considered the 'Four Phases' to be the essential pivot of the mind.")
    doc.add_paragraph("Regarding the phrase 'Better to arouse...' [It means] it is better to arouse the mind of the Three Evil Paths (Hell, Ghost, Animal) than to arouse the [narrow] mind of the Two Vehicles. ...Analogy: In the great void of space, there is neither light nor dark. Light and dark depend on matter (rupa) to exclude each other. But in the 'Dharma-nature Great Void,' there is originally Neither Good nor Evil.'")

    # FASCICLE 5-7: THE TEN VEHICLES
    doc.add_heading('Fascicles 5-7: The Ten Vehicles of Contemplation', 1)
    doc.add_paragraph("The Ten Vehicles (Shi Cheng) constitute the core methodology of the Mohe Zhiguan.")
    
    doc.add_heading('1. Viewing the Mind as the Inconceivable Realm', 2)
    doc.add_paragraph("'The 'One Mind' serves to conclude the 'Principle-Realm' (Li Jing). As previously explained, the root lies in the One Mind. Having already expounded the perfectly fused Three Truths... one must know that the Body and the Land are 'Three Thousand [Realms] in a Single Thought' (Ichinen Sanzen).")
    doc.add_paragraph("When one attains the Way, one accords with this original Principle: a single body and a single thought pervade the entire Dharma-realm.")
    doc.add_paragraph("To say 'there is no mind' is simply to highlight that the mind has no [fixed] existence; yet the mind is not 'nothing.' The term 'Jie-er' (Tiny/Momentary) refers to the kshana-thought—the momentary mind that continues without interruption. In a single such kshana, the Three Thousand Realms are fully possessed (Ju-zu).'")

    doc.add_heading('10. Freedom from Dharma-Love (Wu Fa Ai)', 2)
    doc.add_paragraph("'Tenth is Freedom from Dharma-Love (Wu Fa Ai). Having practiced the previous nine matters and passed through internal and external obstacles... one has entered the stage of the Six Purified Senses (The Likeness Stage).")
    doc.add_paragraph("The reason one does not 'Issue the True [Realization]' (Fa Zhen) is due to Abiding and Attachment (Zhu Zhe). Being attached to the 'Likeness Dharma' (Xiang-si Fa) is called 'Falling from the Summit' (Ding Duo).")
    doc.add_paragraph("...In the Small Vehicle, 'Falling from the Summit' often means regressing [to a lower state]. However, in the Perfect Position (Yuan Wei), since Views and Thoughts are already shed, 'Falling' here has only one meaning: Stagnating at the Summit (Zhu Ding).'")

    # FASCICLE 8: SICKNESS
    doc.add_heading('Fascicle 8: The Object of Sickness', 1)
    
    doc.add_heading('1. Diagnostics', 2)
    doc.add_paragraph("'Regarding 'Long Illness' (Chang Bing)... 'Diagnosing' (Zhen) means reading the pulse (Hou Mai).")
    doc.add_paragraph('* The Four Seasons Pulses: Spring is Wiry (Xian), Summer is Surging (Hong), Autumn is Floating (Fu), Winter is Deep (Chen).')
    doc.add_paragraph("* The Five Viscera: Illness is formed when a specific organ's energy is Excessive (Tai Zeng), not merely due to mutual generation or destruction (Xiang Sheng Xiang Ke).'")

    doc.add_heading('2. The Six Breaths', 2)
    doc.add_paragraph("'The text mentions the Six Breaths (Liu Qi): Chui, Hu, Xi, He, Xu, Si.")
    doc.add_paragraph('1. Chui (吹): Cures Cold (Leng).')
    doc.add_paragraph('2. Hu (呼): Cures Heat (Re).')
    doc.add_paragraph('3. Xi (唏): Cures Pain (Tong).')
    doc.add_paragraph('4. He (呵): Cures Annoyance/Agitation (Fan).')
    doc.add_paragraph('5. Xu (噓): Cures Phlegm (Tan).')
    doc.add_paragraph('6. Si (呬): Cures Exhaustion (Lao).\'')

    doc.save('Miao-lo/Guketsu_Project/01_TRANSLATIONS/Guketsu_Scholarly.docx')
    print("Expanded Scholarly Edition created.")

if __name__ == "__main__":
    create_scholarly_manuscript()
