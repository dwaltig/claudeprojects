import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration
BASE_DIR = "/Users/williamaltig/claudeprojects/Tiantai_Teachings_Project/01_TRANSLATIONS/Great_Cessation/Scholarly"
OUTPUT_FILE = "/Users/williamaltig/claudeprojects/Mohe_Zhiguan_Publication_Ready.docx"

def get_chapter_files(chapter_num):
    """Finds and sorts files for a specific chapter."""
    # Pattern: Mohe_Zhiguan_Chapter_01_...SCHOLARLY.md
    pattern = f"Mohe_Zhiguan_Chapter_{chapter_num:02d}_"
    files = [f for f in os.listdir(BASE_DIR) if f.startswith(pattern) and "SCHOLARLY" in f]
    
    # Sort by Part number if present
    def sort_key(filename):
        part_match = re.search(r"Part_(\d+)", filename)
        if part_match:
            return int(part_match.group(1))
        if "Summary" in filename: return 999
        return 0
        
    files.sort(key=sort_key)
    return files

def read_markdown_file(filename):
    path = os.path.join(BASE_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()

def clean_header_text(text):
    """Removes 'Mohe Zhiguan - Chapter X' and other artifacts."""
    # Remove "Mohe Zhiguan - Chapter \d+:" prefix
    text = re.sub(r'^Mohe Zhiguan - Chapter \d+[:]? ', '', text)
    # Remove (Part X) suffix in headers if redundant
    text = re.sub(r'\(Part \d+.*?\)', '', text)
    return text.strip()

def parse_markdown_to_runs(paragraph, text):
    """Parses simple markdown (bold, italic) and adds runs to paragraph."""
    # This is a simplified parser. For full robustness, a library is better, 
    # but this covers *bold* and _italic_ basics.
    
    # Check for bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            font = paragraph.add_run(part[2:-2]).font
            font.bold = True
        else:
            # Check for italic (*text*) inside the non-bold part
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    font = paragraph.add_run(sub_part[1:-1]).font
                    font.italic = True
                else:
                    paragraph.add_run(sub_part)

def create_table_from_lines(doc, table_lines):
    """Creates a Word table from markdown table lines."""
    if not table_lines: return
    
    # Determine columns by pipe count in first line
    header = table_lines[0].strip('|').split('|')
    cols = len(header)
    rows = len(table_lines)
    
    # Skip separator line (second line usually |---|---|)
    valid_rows = []
    for line in table_lines:
        if '---' in line: continue
        cells = line.strip('|').split('|')
        # Normalize cell count
        if len(cells) < cols:
            cells += [''] * (cols - len(cells))
        valid_rows.append(cells[:cols])
        
    table = doc.add_table(rows=len(valid_rows), cols=cols)
    table.style = 'Table Grid'
    
    for r_idx, row_data in enumerate(valid_rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            parse_markdown_to_runs(p, cell_text.strip())

def process_mohe_docx():
    doc = Document()
    
    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    doc.add_heading('Mohe Zhiguan', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('By Master Zhiyi\nTranslated by The Tiantai Teachings Project').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('ISBN: 9798245489018\nImprint: Independently published').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Add Table of Contents ---
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()

    # Buffer for sequential content
    content_buffer = [] 
    
    # --- Add General Introduction ---
    print("Adding General Introduction...")
    content_buffer.append(('heading_1', 'General Introduction'))
    intro_path = "/Users/williamaltig/claudeprojects/MOHE_ZHIGUAN_INTRODUCTION.md"
    if os.path.exists(intro_path):
        with open(intro_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            for line in lines:
                line = line.strip()
                if not line or line.startswith('# '): continue # Skip title in file, we added H1 above
                if line.startswith('* ') or line.startswith('- '):
                     content_buffer.append(('list_bullet', line[2:]))
                else:
                     content_buffer.append(('paragraph', line))
    content_buffer.append(('page_break', '')) 
    
    # Table parsing state
    in_table = False
    table_buffer = []

    # Process Chapters 1-10 (files found for 1-7 currently)
    for i in range(1, 11):
        files = get_chapter_files(i)
        if not files: continue
        
        print(f"Processing Chapter {i}...")
        
        if i > 1:
            content_buffer.append(('page_break', ''))
        
        # Semantic H1 for the Chapter
        content_buffer.append(('heading_1', f"Chapter {i}"))

        for filename in files:
            lines = read_markdown_file(filename)

            for line in lines:
                line = line.strip()
                if not line:
                    if in_table:
                        content_buffer.append(('table', table_buffer))
                        table_buffer = []
                        in_table = False
                    continue

                # --- Metadata Stripping (Robust) ---
                if line.startswith('<!--') or line.startswith('Next:') or line.startswith('**Next') or line.startswith('**Translated by'):
                    continue
                
                # Check for metadata keys (bolded or plain)
                clean_line = line.replace('*', '').strip()
                if clean_line.startswith('Source:') or clean_line.startswith('Author:') or clean_line.startswith('Translator:') or clean_line.startswith('Date:') or clean_line.startswith('CBETA:'):
                    continue
                
                if "End of Section" in clean_line or "End of Chapter" in clean_line: continue
                if line.startswith('---'): continue

                # --- Header Processing ---
                if line.startswith('# '):
                    text = clean_header_text(line[2:])
                    # Demote file title to H2
                    content_buffer.append(('heading_2', text))
                    continue

                if line.startswith('## '):
                    text = line[3:]
                    if "Source Text" in text or "Architect's Translation" in text or "Scholarly Translation" in text:
                        # Keep these as structural markers but maybe H3
                        content_buffer.append(('heading_3', clean_header_text(text)))
                    else:
                        content_buffer.append(('heading_3', clean_header_text(text)))
                    continue
                
                if line.startswith('### '):
                    content_buffer.append(('heading_4', clean_header_text(line[4:])))
                    continue

                # --- Table Parsing ---
                if line.startswith('|'):
                    in_table = True
                    table_buffer.append(line)
                    continue
                elif in_table:
                    content_buffer.append(('table', table_buffer))
                    table_buffer = []
                    in_table = False
                
                # --- List Bullet Handling ---
                if line.startswith('* ') or line.startswith('- '):
                    content = line[2:].strip()
                    content_buffer.append(('list_bullet', content))
                    continue
                
                # --- Blockquote ---
                if line.startswith('> '):
                    content = line[2:].strip()
                    if content.startswith('*') and content.endswith('*'): content = content[1:-1]
                    content_buffer.append(('blockquote', content))
                    continue

                # --- Paragraph ---
                content_buffer.append(('paragraph', line))
            
            if in_table:
                content_buffer.append(('table', table_buffer))
                table_buffer = []
                in_table = False

    # --- Append Methodology Appendix ---
    print("Appending Methodology Appendix...")
    content_buffer.append(('page_break', ''))
    
    methodology_path = "/Users/williamaltig/claudeprojects/TRANSLATION_METHODOLOGY.md"
    if os.path.exists(methodology_path):
        with open(methodology_path, 'r', encoding='utf-8') as f:
            meth_lines = f.readlines()
            for line in meth_lines:
                line = line.strip()
                if not line: continue
                if line.startswith('# '):
                    content_buffer.append(('heading_1', line[2:]))
                elif line.startswith('## '):
                    content_buffer.append(('heading_2', line[3:]))
                elif line.startswith('### '):
                    content_buffer.append(('heading_3', line[4:]))
                elif line.startswith('* ') or line.startswith('- '):
                    content_buffer.append(('list_bullet', line[2:]))
                else:
                    content_buffer.append(('paragraph', line))

    # --- Rendering ---
    print("Rendering Content...")
    for item_type, content in content_buffer:
        if item_type == 'table':
            create_table_from_lines(doc, content)
        elif item_type == 'page_break':
            doc.add_page_break()
        elif item_type.startswith('heading_'):
            level = int(item_type.split('_')[1])
            p = doc.add_heading(level=level)
            parse_markdown_to_runs(p, content)
        elif item_type == 'list_bullet':
            p = doc.add_paragraph(style='List Bullet')
            parse_markdown_to_runs(p, content)
        elif item_type == 'blockquote':
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            parse_markdown_to_runs(p, content)
            for run in p.runs: run.italic = True
        else:
            p = doc.add_paragraph(style='Normal')
            parse_markdown_to_runs(p, content)

    doc.save(OUTPUT_FILE)
    print(f"Successfully created: {OUTPUT_FILE}")

if __name__ == "__main__":
    process_mohe_docx()
