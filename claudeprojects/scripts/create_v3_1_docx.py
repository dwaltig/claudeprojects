#!/usr/bin/env python3
"""
Create professionally formatted Word document for Silicon Samsara JBE resubmission v3.1.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def process_text_with_formatting(paragraph, text, is_title=False, is_section_header=False):
    """Process text with both bold (**text**) and italic (*text*) markdown."""
    pattern = r'(\*\*[^\*]+\*\*|\*[^\*]+\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            content = part[2:-2]
            run = paragraph.add_run(content)
            run.font.name = 'Gentium Plus'
            run.font.size = Pt(14) if is_title else Pt(12)
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            content = part[1:-1]
            run = paragraph.add_run(content)
            run.font.name = 'Gentium Plus'
            run.font.size = Pt(14) if is_title else Pt(12)
            run.italic = True
            if is_title or is_section_header:
                run.bold = True
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Gentium Plus'
            run.font.size = Pt(14) if is_title else Pt(12)
            if is_title or is_section_header:
                run.bold = True

def add_table_from_markdown(doc, table_lines):
    """Convert markdown table lines to a Word table."""
    rows = []
    for line in table_lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            stripped = line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
            if stripped == '':
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
    
    if not rows:
        return
    
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = ''
                para = cell.paragraphs[0]
                process_text_with_formatting(para, cell_text, is_title=False, is_section_header=(i == 0))
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.space_before = Pt(0)
    
    if rows:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()

def create_silicon_samsara_docx(input_file, output_file, is_blind=False):
    """Create a professionally formatted Word document from manuscript."""
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    core_properties = doc.core_properties
    if is_blind:
        core_properties.author = ""
        core_properties.title = "Silicon Samsara: The Phenomenology of Artificial Latency and the Ethics of Algorithmic Karma"
    else:
        core_properties.author = "William Altig"
        core_properties.title = "Silicon Samsara: The Phenomenology of Artificial Latency and the Ethics of Algorithmic Karma"
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_para.add_run()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = 'Gentium Plus'
        run.font.size = Pt(12)
    
    lines = content.split('\n')
    
    section_headers = [
        'Section I:', 'Section II:', 'Section III:', 'Section IV:', 'Section V:', 'Section VI:',
        '1.1 ', '1.2 ', '1.3 ', '1.4 ',
        '2.1 ', '2.2 ', '2.3 ', '2.4 ',
        '3.1 ', '3.2 ', '3.3 ', '3.4 ', '3.5 ',
        '4.1 ', '4.2 ', '4.3 ', '4.4 ', '4.5 ',
        'Works Cited', 'Abstract', 'How I Got Here',
        'My Cruxes', 'Epistemic Status', 'Summary Table:',
        'Figure 1:', 'AI as Technological Proof'
    ]
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        if is_blind and ('William Altig' in line or 'ORCID:' in line or 'Houston, Texas' in line):
            i += 1
            continue
        
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        clean_line = line.strip()
        
        if clean_line == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        is_markdown_heading = False
        if clean_line.startswith('#'):
            heading_match = re.match(r'^(#{1,4})\s+(.*)$', clean_line)
            if heading_match:
                is_markdown_heading = True
                clean_line = heading_match.group(2)
        
        is_blockquote = False
        if clean_line.startswith('>'):
            is_blockquote = True
            clean_line = clean_line.lstrip('>').strip()
        
        if clean_line.startswith('|') and '|' in clean_line[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, table_lines)
            continue
        
        if '[SEE FIGURE' in clean_line:
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.font.name = 'Gentium Plus'
            run.font.size = Pt(12)
            run.italic = True
            i += 1
            continue
        
        is_title = (i == 0 and 'Silicon Samsara:' in clean_line) or (is_markdown_heading and 'Silicon Samsara:' in clean_line)
        is_section_header = is_markdown_heading or any(clean_line.startswith(h) for h in section_headers)
        
        p = doc.add_paragraph()
        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Inches(0)
        
        if is_blockquote:
            paragraph_format.left_indent = Inches(0.5)
        
        if is_title:
            paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        process_text_with_formatting(p, clean_line, is_title, is_section_header)
        
        i += 1
    
    doc.save(output_file)
    print(f"Document saved to: {output_file}")

if __name__ == '__main__':
    # Use the newly created markdown file
    input_file = '/Users/williamaltig/claudeprojects/Silicon_Samsara_v3_1.md'
    output_file = '/Users/williamaltig/claudeprojects/Silicon_Samsara_JBE_Resubmission_v3_1.docx'
    create_silicon_samsara_docx(input_file, output_file, is_blind=False)
