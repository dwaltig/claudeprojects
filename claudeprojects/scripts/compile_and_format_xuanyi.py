import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuration
BASE_DIR = "/Users/williamaltig/claudeprojects/Tiantai_Teachings_Project/01_TRANSLATIONS/The_Profound_Meaning/Scholarly"
OUTPUT_FILE = "/Users/williamaltig/claudeprojects/Fahua_Xuanyi_Publication_Ready.docx"

def get_fascicle_files(fascicle_num):
    """Finds and sorts files for a specific fascicle."""
    pattern = f"Xuanyi_Fascicle_{fascicle_num:02d}_"
    files = [f for f in os.listdir(BASE_DIR) if f.startswith(pattern) and "SCHOLARLY" in f]
    
    # Sort by Part number if present
    def sort_key(filename):
        part_match = re.search(r"Part_(\d+)", filename)
        if part_match:
            return int(part_match.group(1))
        return 0
        
    files.sort(key=sort_key)
    return files

def read_markdown_file(filename):
    path = os.path.join(BASE_DIR, filename)
    with open(path, 'r', encoding='utf-8') as f:
        return f.readlines()

def clean_header_text(text):
    """Removes 'Xuanyi - Fascicle X' and other artifacts."""
    text = re.sub(r'^Xuanyi [-–] Fascicle \d+[:]? ', '', text)
    text = re.sub(r'^Fascicle \d+[:]? ', '', text)
    return text.strip()

def parse_markdown_to_runs(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            font = paragraph.add_run(part[2:-2]).font
            font.bold = True
        else:
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub_part in sub_parts:
                if sub_part.startswith('*') and sub_part.endswith('*'):
                    font = paragraph.add_run(sub_part[1:-1]).font
                    font.italic = True
                else:
                    paragraph.add_run(sub_part)

def create_table_from_lines(doc, table_lines):
    if not table_lines: return
    header = table_lines[0].strip('|').split('|')
    cols = len(header)
    rows = len(table_lines)
    valid_rows = []
    for line in table_lines:
        if '---' in line: continue
        cells = line.strip('|').split('|')
        if len(cells) < cols: cells += [''] * (cols - len(cells))
        valid_rows.append(cells[:cols])
    table = doc.add_table(rows=len(valid_rows), cols=cols)
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(valid_rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            parse_markdown_to_runs(p, cell_text.strip())

def process_xuanyi_docx():
    doc = Document()
    
    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    # Title Page
    doc.add_paragraph('The Profound Meaning of the Lotus Sutra\nBy Master Zhiyi').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('ISBN: 9798245493282\nImprint: Independently published').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Add Table of Contents ---
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()

    # Buffer
    content_buffer = [] 
    in_table = False
    table_buffer = []

    # --- Add General Introduction ---
    print("Adding General Introduction...")
    content_buffer.append(('heading_1', 'General Introduction'))
    intro_path = "/Users/williamaltig/claudeprojects/FAHUA_XUANYI_INTRODUCTION.md"
    if os.path.exists(intro_path):
        with open(intro_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            for line in lines:
                line = line.strip()
                if not line or line.startswith('# '): continue 
                if line.startswith('* ') or line.startswith('- '):
                     content_buffer.append(('list_bullet', line[2:]))
                else:
                     content_buffer.append(('paragraph', line))
    content_buffer.append(('page_break', ''))

    # Process Fascicles 1-10
    for i in range(1, 11):
        files = get_fascicle_files(i)
        if not files: continue
        
        print(f"Processing Fascicle {i}...")
        
        if i > 1:
            content_buffer.append(('page_break', ''))
        
        # Semantic H1
        content_buffer.append(('heading_1', f"Fascicle {i}"))

        for filename in files:
            lines = read_markdown_file(filename)

            for line in lines:
                line = line.strip()
                if not line:
                    if in_table:
                        content_buffer.append(('table', table_buffer))
                        table_buffer = []
                        in_table = False
                    continue

                # --- Metadata Stripping (Robust) ---
                if line.startswith('<!--') or line.startswith('Next:') or line.startswith('**Next') or line.startswith('**Translated by'):
                    continue
                
                clean_line = line.replace('*', '').strip()
                # Metadata Keys
                if clean_line.startswith('Source:') or clean_line.startswith('Author:') or clean_line.startswith('Translator:') or clean_line.startswith('Date:') or clean_line.startswith('CBETA:') or clean_line.startswith('Section:'):
                    continue
                # Specific "The Architect" lines
                if "The Architect" in clean_line or "MAS Agent" in clean_line: continue
                
                # Navigation / Status lines
                if "Translation continues" in clean_line or "complete. Translation continues" in clean_line or "translation complete" in clean_line or "fully translated" in clean_line: continue
                if "End of Section" in clean_line or "End of Fascicle" in clean_line: continue
                if "Colophon" in clean_line: continue
                if "🎉" in clean_line: continue
                if line.startswith('---'): continue

                # Filters for "Footnotes" header if redundant
                if clean_line == "Footnotes" or clean_line == "## Footnotes": continue

                # --- Header Processing ---
                if line.startswith('# '):
                    text = clean_header_text(line[2:])
                    content_buffer.append(('heading_2', text))
                    continue

                if line.startswith('## '):
                    text = line[3:]
                    # Remove "Fascicle X, Part Y:" logic from headers to keep them clean
                    # Regex to remove "Fascicle \d+(, Part \d+)?: " prefix
                    text = re.sub(r'Fascicle \d+(, Part \d+)?: ', '', text).strip()
                    
                    if "Source Text" in text or "Architect's Translation" in text or "Scholarly Translation" in text or "Colophon" in text or "Footnotes" in text:
                        continue # Skip these entirely
                    else:
                        content_buffer.append(('heading_3', clean_header_text(text)))
                    continue
                
                if line.startswith('### '):
                    content_buffer.append(('heading_4', clean_header_text(line[4:])))
                    continue

                # --- Table Parsing ---
                if line.startswith('|'):
                    in_table = True
                    table_buffer.append(line)
                    continue
                elif in_table:
                    content_buffer.append(('table', table_buffer))
                    table_buffer = []
                    in_table = False
                
                # --- List Bullet ---
                if line.startswith('* ') or line.startswith('- '):
                    content = line[2:].strip()
                    content_buffer.append(('list_bullet', content))
                    continue
                
                # --- Blockquote ---
                if line.startswith('> '):
                    content = line[2:].strip()
                    if content.startswith('*') and content.endswith('*'): content = content[1:-1]
                    content_buffer.append(('blockquote', content))
                    continue

                # --- Paragraph ---
                content_buffer.append(('paragraph', line))
            
            if in_table:
                content_buffer.append(('table', table_buffer))
                table_buffer = []
                in_table = False

    # --- Append Methodology Appendix ---
    print("Appending Methodology Appendix...")
    content_buffer.append(('page_break', ''))
    
    methodology_path = "/Users/williamaltig/claudeprojects/TRANSLATION_METHODOLOGY.md"
    if os.path.exists(methodology_path):
        with open(methodology_path, 'r', encoding='utf-8') as f:
            meth_lines = f.readlines()
            for line in meth_lines:
                line = line.strip()
                if not line: continue
                if line.startswith('# '):
                    content_buffer.append(('heading_1', line[2:]))
                elif line.startswith('## '):
                    content_buffer.append(('heading_2', line[3:]))
                elif line.startswith('### '):
                    content_buffer.append(('heading_3', line[4:]))
                elif line.startswith('* ') or line.startswith('- '):
                    content_buffer.append(('list_bullet', line[2:]))
                else:
                    content_buffer.append(('paragraph', line))

    # --- Rendering ---
    print("Rendering Content...")
    for item_type, content in content_buffer:
        if item_type == 'table':
            create_table_from_lines(doc, content)
        elif item_type == 'page_break':
            doc.add_page_break()
        elif item_type.startswith('heading_'):
            level = int(item_type.split('_')[1])
            p = doc.add_heading(level=level)
            parse_markdown_to_runs(p, content)
        elif item_type == 'list_bullet':
            p = doc.add_paragraph(style='List Bullet')
            parse_markdown_to_runs(p, content)
        elif item_type == 'blockquote':
            p = doc.add_paragraph(style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            parse_markdown_to_runs(p, content)
            for run in p.runs: run.italic = True
        else:
            p = doc.add_paragraph(style='Normal')
            parse_markdown_to_runs(p, content)

    doc.save(OUTPUT_FILE)
    print(f"Successfully created: {OUTPUT_FILE}")

if __name__ == "__main__":
    process_xuanyi_docx()
